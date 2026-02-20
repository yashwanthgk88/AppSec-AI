"""
Document Analysis Service - Extracts architecture information from uploaded documents.
Supports PDF, images (PNG, JPG, WEBP), and DOCX files.
Uses AI vision for diagram analysis, OCR, and text extraction for documents.
"""
import base64
import io
import json
import logging
import re
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Constants for OCR quality
MIN_TEXT_THRESHOLD = 100  # Minimum characters to consider PDF as text-based
MAX_PDF_PAGES_FOR_OCR = 20  # Maximum pages to OCR for performance


@dataclass
class ExtractedContent:
    """Content extracted from a single document."""
    filename: str
    file_type: str
    text_content: str = ""
    is_diagram: bool = False
    diagram_analysis: Optional[Dict] = None
    components_found: List[str] = field(default_factory=list)
    data_flows_found: List[str] = field(default_factory=list)


@dataclass
class CombinedArchitecture:
    """Combined architecture from all analyzed documents."""
    project_name: str = ""
    description: str = ""
    components: List[Dict] = field(default_factory=list)
    data_flows: List[Dict] = field(default_factory=list)
    trust_boundaries: List[Dict] = field(default_factory=list)
    technology_stack: List[str] = field(default_factory=list)
    source_documents: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "project_name": self.project_name,
            "description": self.description,
            "components": self.components,
            "data_flows": self.data_flows,
            "trust_boundaries": self.trust_boundaries,
            "technology_stack": self.technology_stack,
            "source_documents": self.source_documents
        }


class DocumentAnalysisService:
    """Service for analyzing uploaded documents to extract architecture information."""

    DIAGRAM_ANALYSIS_PROMPT = """You are an expert software architect analyzing a system architecture diagram.
Analyze this image and extract ALL components, their types, and data flows between them.

For each component, identify:
- Name (what it's labeled as)
- Type (api, database, service, frontend, cache, queue, storage, external_service, etc.)
- Technology (if visible - e.g., PostgreSQL, Redis, AWS S3, etc.)
- Whether it appears to be internet-facing
- Whether it handles sensitive data (based on labels or context)

For each data flow/connection:
- Source component
- Target component
- Protocol (if visible - HTTP, HTTPS, gRPC, TCP, etc.)
- Data types flowing (if labeled)
- Whether encryption is indicated

For trust boundaries:
- Identify any visible boundaries (DMZ, internal network, cloud regions, etc.)
- List which components are in each boundary

Respond with a JSON object in exactly this format:
{
    "system_overview": "Brief description of the overall system",
    "components": [
        {
            "id": "component_1",
            "name": "Component Name",
            "type": "rest_api|database|web_frontend|etc",
            "technology": "Technology if visible",
            "trust_zone": "internet|dmz|internal|restricted",
            "description": "What this component does",
            "internet_facing": true/false,
            "handles_sensitive_data": true/false
        }
    ],
    "data_flows": [
        {
            "id": "flow_1",
            "source_id": "component_1",
            "target_id": "component_2",
            "protocol": "HTTPS",
            "data_types": ["user_data", "auth_tokens"],
            "is_encrypted": true/false,
            "description": "Description of data flow"
        }
    ],
    "trust_boundaries": [
        {
            "id": "boundary_1",
            "name": "Internal Network",
            "zone": "internal",
            "component_ids": ["component_2", "component_3"]
        }
    ],
    "technology_stack": ["React", "Node.js", "PostgreSQL", "Redis"]
}

Be thorough - extract EVERY component and connection visible in the diagram."""

    OCR_PROMPT = """You are a document OCR specialist. Extract ALL text content from this image/document page accurately.

Instructions:
1. Read and transcribe ALL visible text exactly as it appears
2. Preserve the document structure (headings, paragraphs, lists, tables)
3. For tables, format them clearly with separators
4. For diagrams with labels, extract all text labels and annotations
5. Include any captions, footnotes, or annotations
6. If there are multiple columns, read left-to-right, top-to-bottom
7. Preserve technical terms, acronyms, and proper nouns exactly

Output the extracted text in a structured format:
---
[Document/Page Title if visible]

[Main content - preserve paragraphs and structure]

[Tables - format with | separators]

[Diagram Labels/Annotations - list them]

[Footnotes/Notes]
---

Be thorough and accurate. Extract EVERY piece of text visible in the image."""

    COMBINED_ANALYSIS_PROMPT = """You are an expert software architect analyzing a document that contains both text and diagrams.

First, perform OCR to extract ALL text content from this image accurately.
Then, analyze any architectural diagrams, flowcharts, or system diagrams present.

For text extraction:
- Read ALL visible text exactly as it appears
- Preserve structure (headings, paragraphs, lists)
- Extract table contents
- Include all labels and annotations

For architecture analysis:
- Identify system components and their types
- Map data flows and connections
- Identify trust boundaries
- Note security controls mentioned

Respond with a JSON object:
{
    "extracted_text": "Full OCR text from the document...",
    "has_architecture_diagram": true/false,
    "system_overview": "Description based on text and diagrams",
    "components": [
        {
            "id": "component_1",
            "name": "Component Name",
            "type": "rest_api|database|web_frontend|etc",
            "technology": "Technology mentioned or visible",
            "trust_zone": "internet|dmz|internal|restricted",
            "description": "What this component does",
            "internet_facing": true/false,
            "handles_sensitive_data": true/false,
            "security_controls": ["auth", "encryption"]
        }
    ],
    "data_flows": [
        {
            "id": "flow_1",
            "source_id": "component_1",
            "target_id": "component_2",
            "protocol": "HTTPS",
            "data_types": ["user_data"],
            "is_encrypted": true/false,
            "description": "Description"
        }
    ],
    "trust_boundaries": [
        {
            "id": "boundary_1",
            "name": "Boundary Name",
            "zone": "zone_type",
            "component_ids": []
        }
    ],
    "technology_stack": ["list", "of", "technologies"],
    "security_requirements": ["Any security requirements mentioned"],
    "key_findings": ["Important points from the document"]
}

Be thorough - extract ALL text and identify ALL architectural elements."""

    TEXT_ANALYSIS_PROMPT = """You are an expert software architect analyzing technical documentation.
Extract architecture information from this text to understand the system components and their interactions.

From the text, identify:

1. COMPONENTS - Any systems, services, databases, APIs, frontends, or external integrations mentioned
   - What is the component called?
   - What technology does it use?
   - What is its purpose?
   - Is it user-facing or internal?

2. DATA FLOWS - How do components communicate?
   - What protocols are used?
   - What data is exchanged?
   - Is encryption mentioned?

3. SECURITY CONSIDERATIONS - Any security-related information
   - Authentication mechanisms
   - Encryption requirements
   - Trust zones or network segments
   - Compliance requirements

4. TECHNOLOGY STACK - All technologies, frameworks, languages, and tools mentioned

Respond with a JSON object in exactly this format:
{
    "system_overview": "Brief description of the overall system based on the documentation",
    "components": [
        {
            "id": "component_1",
            "name": "Component Name",
            "type": "rest_api|database|web_frontend|etc",
            "technology": "Technology mentioned",
            "trust_zone": "internet|dmz|internal|restricted",
            "description": "What this component does based on the text",
            "internet_facing": true/false,
            "handles_sensitive_data": true/false,
            "security_controls": ["auth_mechanism", "encryption"]
        }
    ],
    "data_flows": [
        {
            "id": "flow_1",
            "source_id": "component_1",
            "target_id": "component_2",
            "protocol": "HTTPS",
            "data_types": ["data_type"],
            "is_encrypted": true/false,
            "description": "Description"
        }
    ],
    "trust_boundaries": [
        {
            "id": "boundary_1",
            "name": "Boundary Name",
            "zone": "zone_type",
            "component_ids": []
        }
    ],
    "technology_stack": ["list", "of", "technologies"],
    "security_notes": "Any important security information from the document"
}

Extract as much detail as possible from the text."""

    def __init__(self, ai_config=None):
        """Initialize the document analysis service."""
        self._ai_client = None
        self.enabled = False
        self.provider = "none"
        self._init_client(ai_config)

    def _init_client(self, ai_config=None):
        """Initialize AI client for document analysis."""
        try:
            from services.ai_client_factory import get_ai_client, get_global_ai_config

            config = ai_config if ai_config else get_global_ai_config()

            if config.api_key:
                self._ai_client = get_ai_client(config)
                self.enabled = self._ai_client.is_configured
                self.provider = config.provider
                logger.info(f"[DocumentAnalysisService] Initialized with {self.provider}")
            else:
                logger.warning("[DocumentAnalysisService] No API key configured")
        except Exception as e:
            logger.warning(f"[DocumentAnalysisService] Failed to initialize AI client: {e}")

    def extract_text_from_pdf(self, pdf_content: bytes) -> Tuple[str, bool]:
        """
        Extract text content from a PDF file.

        Returns:
            Tuple of (extracted_text, is_scanned_pdf)
            is_scanned_pdf indicates if the PDF has minimal text (likely scanned/image-based)
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=pdf_content, filetype="pdf")
            text_parts = []
            total_chars = 0

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    total_chars += len(text.strip())

            doc.close()

            full_text = "\n\n".join(text_parts)
            # If less than threshold chars per page on average, likely a scanned PDF
            avg_chars_per_page = total_chars / max(len(doc), 1) if text_parts else 0
            is_scanned = avg_chars_per_page < MIN_TEXT_THRESHOLD

            return full_text, is_scanned
        except ImportError:
            logger.warning("PyMuPDF not installed. PDF text extraction unavailable.")
            return "", True
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            return "", True

    def convert_pdf_pages_to_images(self, pdf_content: bytes, max_pages: int = None) -> List[Tuple[bytes, str]]:
        """
        Convert PDF pages to images for OCR analysis.

        Args:
            pdf_content: Raw PDF file bytes
            max_pages: Maximum number of pages to convert (defaults to MAX_PDF_PAGES_FOR_OCR)

        Returns:
            List of (image_bytes, media_type) tuples
        """
        if max_pages is None:
            max_pages = MAX_PDF_PAGES_FOR_OCR

        images = []
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=pdf_content, filetype="pdf")
            pages_to_convert = min(len(doc), max_pages)

            for page_num in range(pages_to_convert):
                page = doc[page_num]
                # Render at 2x zoom for better OCR quality
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat, alpha=False)

                # Convert to PNG bytes
                img_bytes = pix.tobytes("png")
                images.append((img_bytes, "image/png"))

                logger.debug(f"Converted PDF page {page_num + 1} to image")

            doc.close()
            logger.info(f"Converted {len(images)} PDF pages to images for OCR")
            return images

        except ImportError:
            logger.warning("PyMuPDF not installed. PDF to image conversion unavailable.")
            return []
        except Exception as e:
            logger.error(f"Failed to convert PDF pages to images: {e}")
            return []

    def perform_ocr_on_image(self, image_content: bytes, media_type: str) -> str:
        """
        Perform OCR on an image using AI vision.

        Args:
            image_content: Raw image bytes
            media_type: MIME type of the image

        Returns:
            Extracted text from the image
        """
        if not self.enabled or not self._ai_client:
            logger.warning("AI client not available for OCR")
            return ""

        try:
            base64_image = base64.standard_b64encode(image_content).decode('utf-8')

            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_image
                            }
                        },
                        {
                            "type": "text",
                            "text": self.OCR_PROMPT
                        }
                    ]
                }
            ]

            response = self._ai_client.chat_completion(
                messages=messages,
                max_tokens=4000
            )

            return response.get('content', '').strip()

        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""

    def perform_combined_analysis(self, image_content: bytes, media_type: str) -> Dict[str, Any]:
        """
        Perform combined OCR and architecture analysis on an image.
        Useful for documents that contain both text and diagrams.

        Args:
            image_content: Raw image bytes
            media_type: MIME type of the image

        Returns:
            Dictionary with extracted text and architecture analysis
        """
        if not self.enabled or not self._ai_client:
            logger.warning("AI client not available for combined analysis")
            return {}

        try:
            base64_image = base64.standard_b64encode(image_content).decode('utf-8')

            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_image
                            }
                        },
                        {
                            "type": "text",
                            "text": self.COMBINED_ANALYSIS_PROMPT
                        }
                    ]
                }
            ]

            response = self._ai_client.chat_completion(
                messages=messages,
                max_tokens=4000
            )

            response_text = response.get('content', '')

            # Extract JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                return json.loads(json_match.group())

            return {}

        except Exception as e:
            logger.error(f"Combined analysis failed: {e}")
            return {}

    def analyze_scanned_pdf(self, pdf_content: bytes) -> Dict[str, Any]:
        """
        Analyze a scanned PDF by converting pages to images and performing OCR.

        Args:
            pdf_content: Raw PDF file bytes

        Returns:
            Combined analysis results from all pages
        """
        page_images = self.convert_pdf_pages_to_images(pdf_content)

        if not page_images:
            logger.warning("No images extracted from PDF for OCR")
            return {}

        all_text_parts = []
        all_components = {}
        all_flows = {}
        all_boundaries = {}
        all_tech = set()
        descriptions = []
        has_diagram = False

        for i, (img_bytes, media_type) in enumerate(page_images):
            logger.info(f"Performing combined analysis on page {i + 1}")

            # Use combined analysis to get both OCR text and architecture info
            analysis = self.perform_combined_analysis(img_bytes, media_type)

            if analysis:
                # Collect extracted text
                if analysis.get('extracted_text'):
                    all_text_parts.append(f"--- Page {i + 1} ---\n{analysis['extracted_text']}")

                if analysis.get('has_architecture_diagram'):
                    has_diagram = True

                if analysis.get('system_overview'):
                    descriptions.append(analysis['system_overview'])

                # Merge components
                for comp in analysis.get('components', []):
                    name = comp.get('name', '')
                    if name and name not in all_components:
                        comp['id'] = f"comp_{len(all_components) + 1}"
                        all_components[name] = comp

                # Merge flows
                for flow in analysis.get('data_flows', []):
                    flow_key = f"{flow.get('source_id', '')}->{flow.get('target_id', '')}"
                    if flow_key not in all_flows:
                        flow['id'] = f"flow_{len(all_flows) + 1}"
                        all_flows[flow_key] = flow

                # Merge boundaries
                for boundary in analysis.get('trust_boundaries', []):
                    name = boundary.get('name', '')
                    if name and name not in all_boundaries:
                        boundary['id'] = f"boundary_{len(all_boundaries) + 1}"
                        all_boundaries[name] = boundary

                # Merge tech stack
                for tech in analysis.get('technology_stack', []):
                    all_tech.add(tech)

        return {
            'extracted_text': "\n\n".join(all_text_parts),
            'system_overview': " ".join(descriptions) if descriptions else "",
            'has_architecture_diagram': has_diagram,
            'components': list(all_components.values()),
            'data_flows': list(all_flows.values()),
            'trust_boundaries': list(all_boundaries.values()),
            'technology_stack': list(all_tech)
        }

    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text content from a DOCX file."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(docx_content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except ImportError:
            logger.warning("python-docx not installed. DOCX text extraction unavailable.")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""

    def extract_images_from_docx(self, docx_content: bytes) -> List[Tuple[bytes, str]]:
        """
        Extract embedded images from a DOCX file.

        Args:
            docx_content: Raw DOCX file bytes

        Returns:
            List of (image_bytes, media_type) tuples
        """
        images = []
        try:
            import zipfile

            # DOCX is a ZIP archive
            with zipfile.ZipFile(io.BytesIO(docx_content)) as zf:
                for name in zf.namelist():
                    # Images are typically in word/media/
                    if name.startswith('word/media/'):
                        img_data = zf.read(name)
                        # Determine media type from extension
                        if name.lower().endswith('.png'):
                            media_type = 'image/png'
                        elif name.lower().endswith(('.jpg', '.jpeg')):
                            media_type = 'image/jpeg'
                        elif name.lower().endswith('.gif'):
                            media_type = 'image/gif'
                        elif name.lower().endswith('.webp'):
                            media_type = 'image/webp'
                        else:
                            # Skip unsupported formats
                            continue

                        images.append((img_data, media_type))
                        logger.debug(f"Extracted image from DOCX: {name}")

            logger.info(f"Extracted {len(images)} images from DOCX")
            return images

        except Exception as e:
            logger.error(f"Failed to extract images from DOCX: {e}")
            return []

    def analyze_docx_with_images(self, docx_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Analyze a DOCX file including both text and embedded images.

        Args:
            docx_content: Raw DOCX file bytes

        Returns:
            Tuple of (extracted_text, combined_analysis)
        """
        # Extract text
        text = self.extract_text_from_docx(docx_content)

        # Extract and analyze embedded images
        images = self.extract_images_from_docx(docx_content)

        combined_analysis = {
            'system_overview': '',
            'components': [],
            'data_flows': [],
            'trust_boundaries': [],
            'technology_stack': [],
            'has_embedded_diagrams': len(images) > 0
        }

        all_components = {}
        all_flows = {}
        all_boundaries = {}
        all_tech = set()
        descriptions = []

        # Analyze each embedded image
        for i, (img_bytes, media_type) in enumerate(images):
            logger.info(f"Analyzing embedded image {i + 1} from DOCX")
            analysis = self.perform_combined_analysis(img_bytes, media_type)

            if analysis:
                if analysis.get('system_overview'):
                    descriptions.append(analysis['system_overview'])

                # Merge components
                for comp in analysis.get('components', []):
                    name = comp.get('name', '')
                    if name and name not in all_components:
                        comp['id'] = f"comp_{len(all_components) + 1}"
                        all_components[name] = comp

                # Merge flows
                for flow in analysis.get('data_flows', []):
                    flow_key = f"{flow.get('source_id', '')}->{flow.get('target_id', '')}"
                    if flow_key not in all_flows:
                        flow['id'] = f"flow_{len(all_flows) + 1}"
                        all_flows[flow_key] = flow

                # Merge boundaries
                for boundary in analysis.get('trust_boundaries', []):
                    name = boundary.get('name', '')
                    if name and name not in all_boundaries:
                        boundary['id'] = f"boundary_{len(all_boundaries) + 1}"
                        all_boundaries[name] = boundary

                # Merge tech stack
                for tech in analysis.get('technology_stack', []):
                    all_tech.add(tech)

        combined_analysis['system_overview'] = " ".join(descriptions) if descriptions else ""
        combined_analysis['components'] = list(all_components.values())
        combined_analysis['data_flows'] = list(all_flows.values())
        combined_analysis['trust_boundaries'] = list(all_boundaries.values())
        combined_analysis['technology_stack'] = list(all_tech)

        return text, combined_analysis

    def _merge_analyses(self, text_analysis: Dict, image_analysis: Dict) -> Dict:
        """
        Merge analysis results from text and image analysis.
        Deduplicates components and flows while preserving all unique findings.
        """
        merged = {
            'system_overview': '',
            'components': [],
            'data_flows': [],
            'trust_boundaries': [],
            'technology_stack': []
        }

        # Merge system overviews
        overviews = []
        if text_analysis.get('system_overview'):
            overviews.append(text_analysis['system_overview'])
        if image_analysis.get('system_overview'):
            overviews.append(image_analysis['system_overview'])
        merged['system_overview'] = " ".join(overviews)

        # Merge and dedupe components by name
        all_components = {}
        for comp in text_analysis.get('components', []):
            name = comp.get('name', '')
            if name:
                all_components[name.lower()] = comp
        for comp in image_analysis.get('components', []):
            name = comp.get('name', '')
            if name and name.lower() not in all_components:
                all_components[name.lower()] = comp

        # Re-assign IDs
        merged['components'] = []
        for i, comp in enumerate(all_components.values()):
            comp['id'] = f"comp_{i + 1}"
            merged['components'].append(comp)

        # Merge and dedupe flows
        all_flows = {}
        for flow in text_analysis.get('data_flows', []):
            key = f"{flow.get('source_id', '')}->{flow.get('target_id', '')}"
            all_flows[key] = flow
        for flow in image_analysis.get('data_flows', []):
            key = f"{flow.get('source_id', '')}->{flow.get('target_id', '')}"
            if key not in all_flows:
                all_flows[key] = flow

        merged['data_flows'] = []
        for i, flow in enumerate(all_flows.values()):
            flow['id'] = f"flow_{i + 1}"
            merged['data_flows'].append(flow)

        # Merge boundaries
        all_boundaries = {}
        for boundary in text_analysis.get('trust_boundaries', []):
            name = boundary.get('name', '')
            if name:
                all_boundaries[name.lower()] = boundary
        for boundary in image_analysis.get('trust_boundaries', []):
            name = boundary.get('name', '')
            if name and name.lower() not in all_boundaries:
                all_boundaries[name.lower()] = boundary

        merged['trust_boundaries'] = []
        for i, boundary in enumerate(all_boundaries.values()):
            boundary['id'] = f"boundary_{i + 1}"
            merged['trust_boundaries'].append(boundary)

        # Merge tech stack
        tech_set = set()
        for tech in text_analysis.get('technology_stack', []):
            tech_set.add(tech)
        for tech in image_analysis.get('technology_stack', []):
            tech_set.add(tech)
        merged['technology_stack'] = list(tech_set)

        return merged

    def analyze_image_with_ai(self, image_content: bytes, media_type: str) -> Dict[str, Any]:
        """Analyze an image/diagram using AI vision capabilities."""
        if not self.enabled or not self._ai_client:
            logger.warning("AI client not available for image analysis")
            return {}

        try:
            # Encode image to base64
            base64_image = base64.standard_b64encode(image_content).decode('utf-8')

            # Build message with image
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_image
                            }
                        },
                        {
                            "type": "text",
                            "text": self.DIAGRAM_ANALYSIS_PROMPT
                        }
                    ]
                }
            ]

            response = self._ai_client.chat_completion(
                messages=messages,
                max_tokens=4000
            )

            response_text = response.get('content', '')

            # Extract JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                return json.loads(json_match.group())

            return {}
        except Exception as e:
            logger.error(f"Failed to analyze image with AI: {e}")
            return {}

    def analyze_text_with_ai(self, text_content: str) -> Dict[str, Any]:
        """Analyze text content using AI to extract architecture information."""
        if not self.enabled or not self._ai_client:
            logger.warning("AI client not available for text analysis")
            return {}

        if not text_content.strip():
            return {}

        try:
            # Truncate very long text
            max_chars = 15000
            if len(text_content) > max_chars:
                text_content = text_content[:max_chars] + "\n...[truncated]..."

            messages = [
                {
                    "role": "user",
                    "content": f"{self.TEXT_ANALYSIS_PROMPT}\n\n--- DOCUMENT TEXT ---\n{text_content}"
                }
            ]

            response = self._ai_client.chat_completion(
                messages=messages,
                max_tokens=4000
            )

            response_text = response.get('content', '')

            # Extract JSON from response
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                return json.loads(json_match.group())

            return {}
        except Exception as e:
            logger.error(f"Failed to analyze text with AI: {e}")
            return {}

    def analyze_document(
        self,
        filename: str,
        content: bytes,
        content_type: str
    ) -> ExtractedContent:
        """Analyze a single document and extract architecture information."""
        extracted = ExtractedContent(
            filename=filename,
            file_type=content_type
        )

        # Determine document type and process accordingly
        if content_type.startswith('image/'):
            # It's an image - use combined analysis for both OCR and diagram detection
            logger.info(f"Analyzing image: {filename}")
            extracted.is_diagram = True

            # Use combined analysis to get both text and architecture info
            analysis = self.perform_combined_analysis(content, content_type)

            if analysis:
                # Store extracted text if present
                if analysis.get('extracted_text'):
                    extracted.text_content = analysis['extracted_text']

                extracted.diagram_analysis = analysis
                extracted.components_found = [
                    c.get('name', '') for c in analysis.get('components', [])
                ]
                extracted.data_flows_found = [
                    f"{f.get('source_id', '')} -> {f.get('target_id', '')}"
                    for f in analysis.get('data_flows', [])
                ]
            else:
                # Fallback to diagram-only analysis
                analysis = self.analyze_image_with_ai(content, content_type)
                extracted.diagram_analysis = analysis
                if analysis:
                    extracted.components_found = [
                        c.get('name', '') for c in analysis.get('components', [])
                    ]
                    extracted.data_flows_found = [
                        f"{f.get('source_id', '')} -> {f.get('target_id', '')}"
                        for f in analysis.get('data_flows', [])
                    ]

        elif content_type == 'application/pdf':
            # Extract text from PDF and check if it's scanned
            logger.info(f"Analyzing PDF: {filename}")
            text, is_scanned = self.extract_text_from_pdf(content)
            extracted.text_content = text

            if is_scanned:
                # Scanned PDF - use OCR via AI vision
                logger.info(f"PDF appears to be scanned, performing OCR: {filename}")
                extracted.is_diagram = True

                ocr_analysis = self.analyze_scanned_pdf(content)

                if ocr_analysis:
                    # Combine OCR text with any extracted text
                    if ocr_analysis.get('extracted_text'):
                        if extracted.text_content:
                            extracted.text_content += "\n\n--- OCR Results ---\n" + ocr_analysis['extracted_text']
                        else:
                            extracted.text_content = ocr_analysis['extracted_text']

                    extracted.diagram_analysis = ocr_analysis
                    extracted.components_found = [
                        c.get('name', '') for c in ocr_analysis.get('components', [])
                    ]
                    extracted.data_flows_found = [
                        f"{f.get('source_id', '')} -> {f.get('target_id', '')}"
                        for f in ocr_analysis.get('data_flows', [])
                    ]
            else:
                # Text-based PDF - analyze text content
                if text:
                    analysis = self.analyze_text_with_ai(text)
                    extracted.diagram_analysis = analysis
                    if analysis:
                        extracted.components_found = [
                            c.get('name', '') for c in analysis.get('components', [])
                        ]
                        extracted.data_flows_found = [
                            f"{f.get('source_id', '')} -> {f.get('target_id', '')}"
                            for f in analysis.get('data_flows', [])
                        ]

        elif 'wordprocessingml' in content_type or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Extract text and embedded images from DOCX
            logger.info(f"Analyzing DOCX: {filename}")

            # Use enhanced DOCX analysis that handles embedded images
            text, image_analysis = self.analyze_docx_with_images(content)
            extracted.text_content = text

            # Analyze text content
            text_analysis = {}
            if text:
                text_analysis = self.analyze_text_with_ai(text)

            # Merge text analysis with image analysis
            if image_analysis.get('has_embedded_diagrams') and image_analysis.get('components'):
                extracted.is_diagram = True
                # Prefer image analysis for components if available
                merged_analysis = self._merge_analyses(text_analysis, image_analysis)
                extracted.diagram_analysis = merged_analysis
            elif text_analysis:
                extracted.diagram_analysis = text_analysis
            else:
                extracted.diagram_analysis = image_analysis

            if extracted.diagram_analysis:
                extracted.components_found = [
                    c.get('name', '') for c in extracted.diagram_analysis.get('components', [])
                ]
                extracted.data_flows_found = [
                    f"{f.get('source_id', '')} -> {f.get('target_id', '')}"
                    for f in extracted.diagram_analysis.get('data_flows', [])
                ]

        else:
            logger.warning(f"Unsupported content type: {content_type} for file: {filename}")

        return extracted

    def merge_extracted_content(
        self,
        extractions: List[ExtractedContent]
    ) -> CombinedArchitecture:
        """Merge extracted content from multiple documents into a single architecture."""
        combined = CombinedArchitecture()

        all_components = {}
        all_flows = {}
        all_boundaries = {}
        all_tech = set()
        descriptions = []

        for extraction in extractions:
            combined.source_documents.append(extraction.filename)

            if extraction.diagram_analysis:
                analysis = extraction.diagram_analysis

                # Add system overview
                if analysis.get('system_overview'):
                    descriptions.append(analysis['system_overview'])

                # Merge components (dedupe by name)
                for comp in analysis.get('components', []):
                    name = comp.get('name', '')
                    if name and name not in all_components:
                        # Generate stable ID
                        comp['id'] = f"comp_{len(all_components) + 1}"
                        all_components[name] = comp

                # Merge data flows
                for flow in analysis.get('data_flows', []):
                    flow_key = f"{flow.get('source_id', '')}->{flow.get('target_id', '')}"
                    if flow_key not in all_flows:
                        flow['id'] = f"flow_{len(all_flows) + 1}"
                        all_flows[flow_key] = flow

                # Merge trust boundaries
                for boundary in analysis.get('trust_boundaries', []):
                    name = boundary.get('name', '')
                    if name and name not in all_boundaries:
                        boundary['id'] = f"boundary_{len(all_boundaries) + 1}"
                        all_boundaries[name] = boundary

                # Merge technology stack
                for tech in analysis.get('technology_stack', []):
                    all_tech.add(tech)

        # Build final architecture
        combined.description = " ".join(descriptions) if descriptions else "Architecture extracted from uploaded documents."
        combined.components = list(all_components.values())
        combined.data_flows = list(all_flows.values())
        combined.trust_boundaries = list(all_boundaries.values())
        combined.technology_stack = list(all_tech)

        # Infer project name from components or documents
        if combined.components:
            # Try to find a main component for naming
            main_types = ['web_frontend', 'api_gateway', 'rest_api']
            for comp in combined.components:
                if comp.get('type') in main_types:
                    combined.project_name = f"{comp.get('name', 'Project')} System"
                    break
            if not combined.project_name:
                combined.project_name = f"{combined.components[0].get('name', 'Extracted')} Architecture"
        else:
            combined.project_name = "Extracted Architecture"

        return combined

    async def analyze_documents(
        self,
        files: List[Tuple[str, bytes, str]]
    ) -> Dict[str, Any]:
        """
        Analyze multiple uploaded documents and extract architecture.

        Args:
            files: List of (filename, content, content_type) tuples

        Returns:
            Dictionary with success status and extracted architecture
        """
        if not files:
            return {"success": False, "error": "No files provided"}

        extractions = []

        for filename, content, content_type in files:
            try:
                logger.info(f"Analyzing document: {filename} ({content_type})")
                extraction = self.analyze_document(filename, content, content_type)
                extractions.append(extraction)
            except Exception as e:
                logger.error(f"Failed to analyze {filename}: {e}")
                continue

        if not extractions:
            return {"success": False, "error": "Failed to analyze any documents"}

        # Merge all extractions
        combined = self.merge_extracted_content(extractions)

        return {
            "success": True,
            "architecture": combined.to_dict(),
            "documents_analyzed": len(extractions),
            "components_found": len(combined.components),
            "data_flows_found": len(combined.data_flows)
        }
