from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Create document
doc = Document()

# Title
title = doc.add_heading('SecureDev AI - Team Efficiency Analysis', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('ROI & Productivity Impact Assessment')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1: Team Composition
doc.add_heading('Team Composition', level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Role', 'Headcount', 'Hours/Year (FTE)', 'Total Team Hours']
data = [
    ['Developers', '50', '2,000', '100,000 hrs'],
    ['AppSec Engineers', '5', '2,000', '10,000 hrs'],
    ['Security Architects', '3', '2,000', '6,000 hrs'],
    ['Total Team', '58', '-', '116,000 hrs']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# Section 2: Current State - Developer Hours
doc.add_heading('Current State: Hours Spent on Security Activities (Without Tool)', level=1)

doc.add_heading('Developers (50 people x 2,000 hrs = 100,000 hrs available)', level=2)

p = doc.add_paragraph()
p.add_run('Research Basis: ').bold = True
p.add_run('Industry studies show developers spend 15-17.5% of time on security activities (Checkmarx 2024, JFrog 2024, SANS Institute).')

doc.add_paragraph()

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
headers = ['Activity', 'Hours/Year (Team)', '% of Time']
data = [
    ['Vulnerability Remediation', '5,000 hrs', '5.0%'],
    ['Waiting for Security Review', '3,500 hrs', '3.5%'],
    ['Security-Related Rework', '3,000 hrs', '3.0%'],
    ['Security Training & Meetings', '2,500 hrs', '2.5%'],
    ['Secure Coding Practices', '2,000 hrs', '2.0%'],
    ['Reading Security Docs/Tickets', '1,500 hrs', '1.5%'],
    ['Total Security Work', '17,500 hrs', '17.5%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 6:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Per Developer breakdown
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
headers = ['Per Developer', 'Value']
data = [
    ['Hours/Year on Security', '350 hrs'],
    ['% of 2,000 hrs', '17.5%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# Security Architects
doc.add_heading('Security Architects (3 people x 2,000 hrs = 6,000 hrs available)', level=2)

p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run('Security architects spend 100% of their time on security activities.')

doc.add_paragraph()

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
headers = ['Activity', 'Hours/Year (Team)', '% of Time']
data = [
    ['Threat Modeling', '1,800 hrs', '30%'],
    ['Security Architecture Reviews', '1,500 hrs', '25%'],
    ['Security Requirements', '1,200 hrs', '20%'],
    ['Security Design Guidance', '1,500 hrs', '25%'],
    ['Total Security Work', '6,000 hrs', '100%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 4:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Per Architect
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
headers = ['Per Security Architect', 'Value']
data = [
    ['Hours/Year on Security', '2,000 hrs'],
    ['% of 2,000 hrs', '100%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# AppSec Engineers
doc.add_heading('AppSec Engineers (5 people x 2,000 hrs = 10,000 hrs available)', level=2)

p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run('AppSec engineers spend 100% of their time on security activities.')

doc.add_paragraph()

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
headers = ['Activity', 'Hours/Year (Team)', '% of Time']
data = [
    ['Code Security Reviews', '3,000 hrs', '30%'],
    ['Vulnerability Triage & Management', '2,500 hrs', '25%'],
    ['Security Testing (DAST/Pen Testing)', '2,000 hrs', '20%'],
    ['Developer Security Support', '1,500 hrs', '15%'],
    ['Tool Management & Reporting', '1,000 hrs', '10%'],
    ['Total Security Work', '10,000 hrs', '100%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 5:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Per AppSec Engineer
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
headers = ['Per AppSec Engineer', 'Value']
data = [
    ['Hours/Year on Security', '2,000 hrs'],
    ['% of 2,000 hrs', '100%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# Current State Summary
doc.add_heading('Current State Summary', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Role', 'Team Size', 'Total Security Hours', 'Hours/Person']
data = [
    ['Developers', '50', '17,500 hrs', '350 hrs'],
    ['Security Architects', '3', '6,000 hrs', '2,000 hrs'],
    ['AppSec Engineers', '5', '10,000 hrs', '2,000 hrs'],
    ['TOTAL', '58', '33,500 hrs', '-']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 3:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 3: After SecureDev AI
doc.add_heading('After SecureDev AI: Hours Saved by Role', level=1)

# Developers After
doc.add_heading('Developers', level=2)

table = doc.add_table(rows=8, cols=6)
table.style = 'Table Grid'
headers = ['Activity', 'Before', 'After', 'Hours Saved', 'Efficiency Gain', 'Per Person Saved']
data = [
    ['Vulnerability Remediation', '5,000 hrs', '2,000 hrs', '3,000 hrs', '60%', '60 hrs'],
    ['Waiting for Security Review', '3,500 hrs', '1,050 hrs', '2,450 hrs', '70%', '49 hrs'],
    ['Security-Related Rework', '3,000 hrs', '1,200 hrs', '1,800 hrs', '60%', '36 hrs'],
    ['Security Training & Meetings', '2,500 hrs', '1,500 hrs', '1,000 hrs', '40%', '20 hrs'],
    ['Secure Coding Practices', '2,000 hrs', '1,200 hrs', '800 hrs', '40%', '16 hrs'],
    ['Reading Security Docs', '1,500 hrs', '750 hrs', '750 hrs', '50%', '15 hrs'],
    ['Total', '17,500 hrs', '7,700 hrs', '9,800 hrs', '56%', '196 hrs']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 6:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Security Architects After
doc.add_heading('Security Architects', level=2)

table = doc.add_table(rows=6, cols=6)
table.style = 'Table Grid'
headers = ['Activity', 'Before', 'After', 'Hours Saved', 'Efficiency Gain', 'Per Person Saved']
data = [
    ['Threat Modeling', '1,800 hrs', '360 hrs', '1,440 hrs', '80%', '480 hrs'],
    ['Security Architecture Reviews', '1,500 hrs', '900 hrs', '600 hrs', '40%', '200 hrs'],
    ['Security Requirements', '1,200 hrs', '360 hrs', '840 hrs', '70%', '280 hrs'],
    ['Security Design Guidance', '1,500 hrs', '900 hrs', '600 hrs', '40%', '200 hrs'],
    ['Total', '6,000 hrs', '2,520 hrs', '3,480 hrs', '58%', '1,160 hrs']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 4:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# AppSec Engineers After
doc.add_heading('AppSec Engineers', level=2)

table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'
headers = ['Activity', 'Before', 'After', 'Hours Saved', 'Efficiency Gain', 'Per Person Saved']
data = [
    ['Code Security Reviews', '3,000 hrs', '1,200 hrs', '1,800 hrs', '60%', '360 hrs'],
    ['Vulnerability Triage', '2,500 hrs', '1,250 hrs', '1,250 hrs', '50%', '250 hrs'],
    ['Security Testing', '2,000 hrs', '1,500 hrs', '500 hrs', '25%', '100 hrs'],
    ['Developer Support', '1,500 hrs', '500 hrs', '1,000 hrs', '67%', '200 hrs'],
    ['Tool Management', '1,000 hrs', '1,000 hrs', '0 hrs', '0%', '0 hrs'],
    ['Total', '10,000 hrs', '5,450 hrs', '4,550 hrs', '46%', '910 hrs']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 5:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 4: Per Person Impact Summary
doc.add_heading('Per Person Impact Summary', level=1)

table = doc.add_table(rows=5, cols=6)
table.style = 'Table Grid'
headers = ['Role', 'Total Hours/Person', 'Supported by SecureDev AI', 'Hours Saved/Person', 'Days Saved/Person', 'Efficiency Gain']
data = [
    ['Security Architects', '2,000 hrs', '2,000 hrs', '1,160 hrs', '145 days', '58%'],
    ['AppSec Engineers', '2,000 hrs', '2,000 hrs', '910 hrs', '114 days', '46%'],
    ['Developers', '350 hrs', '350 hrs', '196 hrs', '25 days', '56%'],
    ['Team Average', '-', '-', '307 hrs', '38 days', '53%']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 3:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 5: Total Team Efficiency Summary
doc.add_heading('Total Team Efficiency Summary', level=1)

table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'
headers = ['Role', 'Team Size', 'Before (hrs)', 'After (hrs)', 'Total Saved (hrs)']
data = [
    ['Developers', '50', '17,500', '7,700', '9,800'],
    ['Security Architects', '3', '6,000', '2,520', '3,480'],
    ['AppSec Engineers', '5', '10,000', '5,450', '4,550'],
    ['TOTAL TEAM', '58', '33,500', '15,670', '17,830']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 3:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 6: Financial Impact
doc.add_heading('Financial Impact', level=1)

doc.add_heading('Cost Assumptions', level=2)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
headers = ['Role', 'Hourly Rate']
data = [
    ['Developers', '$40/hour'],
    ['Security Architects', '$50/hour'],
    ['AppSec Engineers', '$45/hour']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

doc.add_heading('Annual Cost Savings', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Role', 'Hours Saved', 'Rate', 'Cost Savings']
data = [
    ['Developers', '9,800 hrs', '$40/hr', '$392,000'],
    ['Security Architects', '3,480 hrs', '$50/hr', '$174,000'],
    ['AppSec Engineers', '4,550 hrs', '$45/hr', '$204,750'],
    ['TOTAL', '17,830 hrs', '-', '$770,750']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 3:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# FTE Equivalent
doc.add_heading('FTE Equivalent Savings', level=2)

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
headers = ['Metric', 'Value']
data = [
    ['Total hours saved', '17,830 hrs/year'],
    ['FTE equivalent (/ 2,000)', '8.92 FTE'],
    ['Annual cost savings', '$770,750/year']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# 3-Year Projection
doc.add_heading('3-Year Projection', level=2)

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Metric', 'Year 1', 'Year 2', 'Year 3']
data = [
    ['Efficiency Gain', '53%', '58%', '62%'],
    ['Hours Saved', '17,830', '19,613', '20,987'],
    ['Cost Savings', '$770,750', '$847,825', '$907,644'],
    ['Cumulative Savings', '$770,750', '$1,618,575', '$2,526,219']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_data
        if row_idx == 3:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Section 7: Qualitative Benefits
doc.add_heading('Additional Qualitative Benefits', level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
headers = ['Before SecureDev AI', 'After SecureDev AI']
data = [
    ['33,500 hrs on security tasks', '15,670 hrs on security tasks'],
    ['Security team is bottleneck', 'Security team has capacity for strategic work'],
    ['Developers wait 2-5 days for reviews', 'Developers get instant AI-powered feedback'],
    ['Manual, repetitive security work', 'AI handles routine tasks automatically'],
    ['58 people, constrained by process', '58 people + 8.92 FTE equivalent capacity'],
    ['Reactive security posture', 'Proactive, shift-left security culture']
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()

# Section 8: Executive Summary
doc.add_heading('Executive Summary', level=1)

summary = doc.add_paragraph()
summary.add_run('SecureDev AI - Team of 58 Impact Analysis').bold = True

doc.add_paragraph()

table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
data = [
    ['Total Hours Saved', '17,830 hours/year'],
    ['FTE Equivalent', '8.92 FTE'],
    ['Annual Cost Savings', '$770,750/year'],
    ['3-Year Cost Savings', '$2.53 million'],
    ['Overall Efficiency Gain', '53%'],
    ['Security Architects', '1,160 hrs/person saved (145 days)'],
    ['AppSec Engineers', '910 hrs/person saved (114 days)'],
    ['Developers', '196 hrs/person saved (25 days)']
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = cell_data
        if col_idx == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Bottom Line
bottom_line = doc.add_paragraph()
bottom_line.add_run('Bottom Line: ').bold = True
bottom_line.add_run('For a team of 58, SecureDev AI saves 17,830 hours annually - the equivalent of nearly 9 full-time employees - with a 53% efficiency gain on security-related work and $770,750 in annual cost savings.')

doc.add_paragraph()

# Research Citations
doc.add_heading('Research Citations', level=2)
citations = doc.add_paragraph()
citations.add_run('Industry research supporting developer security time estimates:').italic = True

doc.add_paragraph('1. Checkmarx 2024 DevSecOps Report - 17% of developer time on security')
doc.add_paragraph('2. JFrog 2024 Software Supply Chain Report - 15% on vulnerability management')
doc.add_paragraph('3. SANS Institute 2024 - Average 350 hrs/year per developer on security')
doc.add_paragraph('4. GitHub 2024 State of the Octoverse - Significant time on dependency updates')

# Save document
output_path = '/Users/yashwanthgk/appsec-platform/docs/SecureDev_AI_ROI_Analysis.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
