"""Generate Threat Modeling — Client Demo Guide as Word (.docx)"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = 'Consolas'; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True; run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix); run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_para(text):
    doc.add_paragraph(text)


# ══════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Threat Modeling Pipeline')
run.font.size = Pt(32); run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Architecture → STRIDE → Attack Paths → FAIR Risk\nClient Demo Guide')
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('AppSec Platform\nDemo-ready walkthrough including full AI prompts')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. PITCH
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. What it is (30-second pitch)', level=1)
add_para(
    "The Threat Modeling Pipeline takes a plain-English architecture description (optionally with "
    "a diagram, uploaded security requirements, and industry threat intelligence) and produces a "
    "complete, versioned threat model:"
)
add_bullet('Components, data flows, and trust boundaries (parsed from text + diagram)', '• ')
add_bullet('STRIDE threats per component (CWE + MITRE mapped, business-impact enriched)', '• ')
add_bullet('Multi-hop attack paths with exploitation narratives', '• ')
add_bullet('Cyber Kill Chain and attack-tree decomposition', '• ')
add_bullet('FAIR-based dollar-loss quantification (min / likely / max) per threat', '• ')
add_bullet('Versioned: threats carry stable IDs so "new", "modified", "resolved" are tracked across revisions', '• ')
add_para(
    "It is a hybrid pipeline: deterministic template-based scaffolding + AI enrichment on the "
    "top-risk threats. Quick-mode skips AI for fast iteration."
)

# ══════════════════════════════════════════════════════════════════════
# 2. PROBLEM
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. The problem we solve', level=1)
add_para(
    "Manual threat modeling is slow, inconsistent, and rarely kept in sync with the code. Our "
    "pipeline closes the loop:"
)
add_table(
    ['Traditional threat modeling', 'Our pipeline'],
    [
        ['Weeks of workshops with architects and security', 'Minutes of processing from an architecture doc'],
        ['Output is a static Visio or Word file', 'Output is queryable, versioned, and integrated with SAST findings'],
        ['Threats are generic ("SQLi on database")', 'Threats are specific to the component, tech stack, and industry'],
        ['No dollar-value risk quantification', 'FAIR-based ALE with industry multipliers'],
        ['No tracking across revisions', 'Stable threat IDs — new / modified / resolved labels per version'],
        ['Disconnected from the SAST tool', 'Prioritizer re-ranks SAST findings using the threat model as a prior'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. PIPELINE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. End-to-end pipeline', level=1)
add_code_block(
    '[Stage 1]  Architecture Ingestion   (LLM parses text + diagram into components, flows, boundaries)\n'
    '    |\n'
    '    v\n'
    '[Stage 2]  Intelligence Context     (sector threats + client intel + SecReq abuse cases + controls)\n'
    '    |\n'
    '    v\n'
    '[Stage 3]  STRIDE Threat Generation (templates -> AI enrichment of top 15 -> SecReq injection)\n'
    '    |\n'
    '    v\n'
    '[Stage 4]  Attack Paths + MITRE + Kill Chain + Attack Trees\n'
    '    |\n'
    '    v\n'
    '[Stage 5]  FAIR Risk Quantification + Mermaid / Eraser Diagrams'
)
add_para(
    "AI is called in Stage 1, Stage 3 (top 15 threats only), and Stage 4 (top 5 attack paths "
    "only). All other work is deterministic and fully auditable."
)

# ══════════════════════════════════════════════════════════════════════
# 4. STAGE 1
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Stage 1 — Architecture Ingestion', level=1)
add_para(
    "The pipeline accepts free-form architecture text, optionally with an uploaded diagram. An "
    "LLM extracts the structured model (components, flows, trust boundaries). If the LLM is "
    "unavailable, a keyword-based fallback parser still produces a usable model."
)

doc.add_heading('4.1 The exact prompt (verbatim)', level=2)
add_para("Passed as the user message; response format is JSON. temperature = 0.3, max_tokens = 4000.")
add_code_block(
    'Analyze this software architecture and extract a detailed threat model structure.\n\n'
    'ARCHITECTURE DESCRIPTION:\n'
    '{architecture_doc}\n\n'
    'Please provide a comprehensive JSON response with the following structure:\n'
    '{\n'
    '    "system_overview": "Brief description of what the system does",\n'
    '    "technology_stack": ["list", "of", "technologies"],\n'
    '    "components": [\n'
    '        {\n'
    '            "id": "unique_id",\n'
    '            "name": "Component Name",\n'
    '            "type": "external|process|datastore",\n'
    '            "technology": "specific technology (e.g., React, Node.js, PostgreSQL)",\n'
    '            "category": "api|database|authentication|frontend|microservice|cloud|message_queue",\n'
    '            "description": "What this component does",\n'
    '            "data_handled": ["types of data this component handles"],\n'
    '            "trust_level": "untrusted|semi-trusted|trusted|highly-trusted",\n'
    '            "internet_facing": true/false,\n'
    '            "handles_sensitive_data": true/false\n'
    '        }\n'
    '    ],\n'
    '    "data_flows": [\n'
    '        {\n'
    '            "id": "flow_id",\n'
    '            "from": "source_component_id",\n'
    '            "to": "target_component_id",\n'
    '            "data_type": "What data flows",\n'
    '            "protocol": "HTTP/HTTPS/gRPC/SQL/etc",\n'
    '            "encrypted": true/false,\n'
    '            "authenticated": true/false,\n'
    '            "sensitive": true/false\n'
    '        }\n'
    '    ],\n'
    '    "trust_boundaries": [\n'
    '        {\n'
    '            "id": "boundary_id",\n'
    '            "name": "Boundary Name",\n'
    '            "description": "What this boundary separates",\n'
    '            "components_inside": ["list of component_ids inside this boundary"],\n'
    '            "boundary_type": "internet|dmz|internal|data"\n'
    '        }\n'
    '    ],\n'
    '    "security_controls": ["list of mentioned security controls"],\n'
    '    "risk_factors": ["identified risk factors from the architecture"]\n'
    '}\n\n'
    'Be thorough and extract all components, even if implied. Identify ALL data flows between components.\n'
    'For each component, determine the most appropriate category from:\n'
    'api, database, authentication, frontend, microservice, cloud, message_queue.'
)

add_para(
    "The structured JSON becomes the input to every subsequent stage. Critically, each component "
    "is labelled with internet_facing, handles_sensitive_data, and trust_level — these drive risk "
    "scoring later."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. STAGE 2 — INTELLIGENCE CONTEXT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Stage 2 — Intelligence Context Assembly', level=1)
add_para(
    "Deterministic. No LLM. This stage gathers everything the model needs to say something "
    "SPECIFIC about the customer's system:"
)
add_table(
    ['Source', 'What it contributes'],
    [
        ['Sector threat DB', 'Industry-specific threat actors, campaigns, TTPs (FIN7 for banking, Magecart for retail, etc.)'],
        ['Client threat intel', 'User-uploaded incidents, internal threat actors, sector-specific regulations'],
        ['SecReq abuse cases', 'Real "what can go wrong" extracted from user stories (see SecureReq demo guide)'],
        ['SecReq requirements', '"What must be true" controls — used to detect coverage gaps'],
        ['Existing security controls', 'Registered WAFs, IDS, SIEM, etc. with effectiveness scores and STRIDE coverage'],
    ]
)

add_para("The assembled context block looks like:")
add_code_block(
    'System Overview: ...\n'
    'Technology Stack: ...\n'
    'Components: ...\n'
    'Data Flows: ...\n\n'
    '=== THREAT INTELLIGENCE (BANKING SECTOR) ===\n'
    '### FIN7 Targeting US Banking Platforms [CRITICAL]\n'
    'Type: threat_actor\n'
    '...\n\n'
    '=== CLIENT-SPECIFIC THREAT INTELLIGENCE ===\n'
    '### 2024 SWIFT Credential Theft Attempt [CRITICAL]\n'
    '...\n\n'
    '=== EXISTING SECURITY CONTROLS ===\n'
    '- [IMPLEMENTED] WAF (type: preventive, effectiveness: 85%)\n\n'
    '=== SECURITY REQUIREMENTS ANALYSIS ===\n'
    '- [Critical Impact] Social engineering to bypass approval workflow (Actor: Insider, STRIDE: Spoofing)\n'
    '- [Critical] [Cryptography] Implement cryptographic signing of MT103 messages with HSM-backed keys'
)

add_para(
    "This system_context string is injected into every AI call in Stages 3 and 4, which is why "
    "the output talks about the specific customer components, tech stack, and industry — not "
    "generic advice."
)

# ══════════════════════════════════════════════════════════════════════
# 6. STAGE 3 — STRIDE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Stage 3 — STRIDE Threat Generation', level=1)
add_para("Runs in three passes.")

doc.add_heading('6.1 Pass 1 — Template-based generation (deterministic)', level=2)
add_para(
    "For each component × each STRIDE category, technology-specific threat templates are applied. "
    "Fast, fully auditable. Every component × category produces 0-N threat candidates."
)

doc.add_heading('Risk score per threat', level=3)
add_code_block(
    "base_score = {'critical': 9.5, 'high': 7.5, 'medium': 5.0, 'low': 2.5}[severity]\n"
    "if component['internet_facing']:          base_score += 1.0\n"
    "if component['handles_sensitive_data']:   base_score += 0.5\n"
    "if component['trust_level'] == 'untrusted': base_score += 0.5\n"
    "risk_score = min(10.0, base_score)"
)

doc.add_heading('6.2 Pass 2 — AI enrichment of top 15 critical/high threats', level=2)
add_para(
    "Only the 15 highest-risk threats are sent to the LLM. Each one is given the full system "
    "context + the component being analyzed + the threat, and asked for a detailed, contextual "
    "response. max_tokens = 2000."
)

doc.add_heading('The exact prompt (verbatim)', level=3)
add_code_block(
    'You are a senior application security expert performing threat modeling.\n'
    'Analyze this specific threat and provide detailed, contextual information.\n\n'
    'SYSTEM CONTEXT:\n'
    '{system_context}\n\n'
    'COMPONENT BEING ANALYZED:\n'
    '- Name: {component.name}\n'
    '- Type: {component.type}\n'
    '- Category: {component.category}\n'
    '- Technology: {component.technology}\n'
    '- Internet Facing: {component.internet_facing}\n'
    '- Handles Sensitive Data: {component.handles_sensitive_data}\n'
    '- Trust Level: {component.trust_level}\n'
    '- Data Handled: {component.data_handled}\n\n'
    'THREAT TO ANALYZE:\n'
    '- Threat Name: {threat.threat}\n'
    '- STRIDE Category: {threat.category}\n'
    '- Severity: {threat.severity}\n'
    '- CWE: {threat.cwe}\n'
    '- MITRE Techniques: {threat.mitre_techniques}\n\n'
    'Provide a detailed JSON response with the following structure.\n'
    'Be SPECIFIC to this component and system - do not give generic advice:\n'
    '{\n'
    '    "description": "A detailed 2-3 sentence description of how this specific threat applies to this component",\n'
    '    "attack_vector": {\n'
    '        "description": "Detailed explanation of how an attacker would exploit this",\n'
    '        "entry_points": ["List of specific entry points"],\n'
    '        "techniques": ["Specific attack techniques"]\n'
    '    },\n'
    '    "business_impact": {\n'
    '        "financial":   "Specific financial impact",\n'
    '        "reputational":"Reputational damage assessment",\n'
    '        "operational": "Operational impact on business",\n'
    '        "compliance":  "Regulatory and compliance implications"\n'
    '    },\n'
    '    "affected_assets": ["List of specific assets at risk"],\n'
    '    "prerequisites": {\n'
    '        "access_required": "What access level an attacker needs",\n'
    '        "conditions": ["Conditions that must be true for attack to succeed"]\n'
    '    },\n'
    '    "attack_complexity": {\n'
    '        "level": "Low/Medium/High",\n'
    '        "skill_level": "Basic/Intermediate/Advanced",\n'
    '        "time_required": "Estimated time to execute",\n'
    '        "description": "Why this complexity level"\n'
    '    },\n'
    '    "mitigation": "Specific, actionable mitigation recommendations for this component",\n'
    '    "detection": "How to detect this attack in progress or after the fact"\n'
    '}\n\n'
    'Be specific and technical. Reference the actual component name and technology.'
)

doc.add_heading('6.3 Pass 3 — SecReq threat injection', level=2)
add_para(
    "Each abuse case from the SecureReq feature is injected as an explicit threat with "
    "source=\"securereq\" for full traceability back to the user story. This is how a business-"
    "level abuse case (\"promo code stacking\") shows up as a real threat on the API component."
)

doc.add_heading('6.4 STRIDE output schema', level=2)
add_code_block(
    '{\n'
    '  "id": "threat_123",\n'
    '  "component": "API Gateway",\n'
    '  "category": "Tampering",\n'
    '  "threat": "Request Smuggling",\n'
    '  "severity": "critical",\n'
    '  "risk_score": 8.5,\n'
    '  "cwe": "CWE-444",\n'
    '  "mitre": ["T1037.005"],\n'
    '  "description": "AI-enriched or fallback description",\n'
    '  "attack_vector": {...},\n'
    '  "business_impact": {...},\n'
    '  "affected_assets": [...],\n'
    '  "prerequisites": {...},\n'
    '  "attack_complexity": {...},\n'
    '  "mitigation": "...",\n'
    '  "detection": "...",\n'
    '  "source": "stride" | "securereq"\n'
    '}'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 7. STAGE 4 — ATTACK PATHS + MITRE + KILL CHAIN
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Stage 4 — Attack Paths, MITRE, Kill Chain, Attack Trees', level=1)

doc.add_heading('7.1 Attack Path generation', level=2)
add_para(
    "The pipeline builds a graph from the data flows. Entry points = internet-facing components; "
    "targets = sensitive datastores. All paths up to depth 5 are enumerated. Top 10 by risk score "
    "are kept; top 5 get AI-enriched narratives. max_tokens = 3000."
)

doc.add_heading('The exact prompt (verbatim)', level=3)
add_code_block(
    'You are a senior penetration tester analyzing an attack path through a system.\n'
    'Generate a detailed attack path analysis.\n\n'
    'SYSTEM CONTEXT:\n'
    '{system_context}\n\n'
    'ATTACK PATH:\n'
    'Entry Point: {entry_point_name} ({entry_point_category})\n'
    'Target:      {target_name} ({target_category})\n'
    "Path: {' -> '.join(path_names)}\n\n"
    'THREATS ALONG THIS PATH:\n'
    '{threats_summary}\n\n'
    'Generate a detailed JSON response:\n'
    '{\n'
    '    "attack_scenario": "A compelling 3-4 sentence narrative describing how an attacker would exploit this path, referencing specific components and threats",\n'
    '    "exploitation_steps": [\n'
    '        {\n'
    '            "step": 1,\n'
    '            "phase": "Reconnaissance/Initial Access/Lateral Movement/Privilege Escalation/Objective",\n'
    '            "action": "Brief action description",\n'
    '            "details": "Detailed explanation of what the attacker does at this step, tools they might use, and what they gain"\n'
    '        }\n'
    '    ],\n'
    '    "potential_impact": {\n'
    '        "level": "Critical/High/Medium/Low",\n'
    '        "description": "Overall impact description",\n'
    '        "data_exposure": "What data could be exposed",\n'
    '        "system_impact": "Impact on system availability/integrity",\n'
    '        "business_impact": "Business consequences",\n'
    '        "compliance_impact": "Regulatory implications"\n'
    '    },\n'
    '    "difficulty": {\n'
    '        "level": "Low/Medium/High",\n'
    '        "description": "Why this difficulty level",\n'
    '        "required_skills": "Skills needed to execute this attack",\n'
    '        "time_estimate": "Estimated time to execute",\n'
    '        "tools_needed": ["Tools an attacker might use"]\n'
    '    },\n'
    '    "detection_opportunities": [\n'
    '        {\n'
    '            "point": "Where in the attack chain this can be detected",\n'
    '            "method": "How to detect it",\n'
    '            "effectiveness": "High/Medium/Low"\n'
    '        }\n'
    '    ],\n'
    '    "recommended_controls": [\n'
    '        {\n'
    '            "control": "Security control name",\n'
    '            "implementation": "Specific implementation guidance",\n'
    '            "priority": "Critical/High/Medium"\n'
    '        }\n'
    '    ]\n'
    '}\n\n'
    'Be specific to this system and path. Reference actual component names.'
)

doc.add_heading('7.2 MITRE ATT&CK mapping (deterministic)', level=2)
add_para(
    "Template-based lookup from threat CWE + MITRE technique IDs to the full ATT&CK framework. "
    "Threats are grouped by tactic (Initial Access, Execution, Persistence, Privilege Escalation, "
    "Defense Evasion, Credential Access, Discovery, Lateral Movement, Collection, C2, "
    "Exfiltration, Impact). No AI — fully reproducible."
)

doc.add_heading('7.3 Kill Chain analysis (deterministic)', level=2)
add_para(
    "Maps threats to the 7-phase Lockheed Martin Cyber Kill Chain: Reconnaissance → "
    "Weaponization → Delivery → Exploitation → Installation → Command & Control → Actions on "
    "Objectives. Gives the SOC a linear narrative for tabletop exercises."
)

doc.add_heading('7.4 Attack trees (deterministic)', level=2)
add_para(
    "Hierarchical decomposition of each high-risk attack goal. Shows AND / OR decomposition of "
    "sub-goals, all the way down to primitive actions. Useful for defensive brainstorming: every "
    "leaf is a potential detection or mitigation point."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 8. STAGE 5 — FAIR + DIAGRAMS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Stage 5 — FAIR Risk Quantification + Diagrams', level=1)

doc.add_heading('8.1 FAIR (Factor Analysis of Information Risk)', level=2)
add_para("Deterministic. Per threat:")
add_code_block(
    'Loss Event Frequency (LEF) = Threat Event Frequency × Vulnerability\n\n'
    'Loss Magnitude (LM) aggregates 6 loss categories:\n'
    '   productivity_loss + response_cost + replacement_cost\n'
    ' + fines_and_judgments + reputation_damage + legal_costs\n\n'
    'Annualized Loss Expectancy (ALE) = LEF × LM\n\n'
    'Industry multiplier × org-size multiplier applied on top.\n'
    'Confidence intervals reported as (min, likely, max).'
)
add_para(
    "This produces a dollar number per threat, per year. It is the number the CFO wants to see — "
    "and because it is deterministic, a customer can re-run with different assumptions without "
    "any LLM cost."
)

doc.add_heading('8.2 Diagrams', level=2)
add_bullet("Mermaid DFD — always generated, Level 0 (context) + Level 1 (detailed)", '• ')
add_bullet("Eraser.io diagrams — if ERASER_API_KEY is configured: Architecture / Threat Model, Kill Chain, Data Flow. Professional PDF-grade output.", '• ')

# ══════════════════════════════════════════════════════════════════════
# 9. LIFECYCLE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Versioning & Lifecycle', level=1)
add_para(
    "Threats carry STABLE IDs across versions: the ID is hashed from "
    "(STRIDE category + target component + attack vector). When you re-run the pipeline after an "
    "architecture change, each threat lands in one of four states:"
)
add_table(
    ['Status', 'Meaning'],
    [
        ['new', 'Threat appears in current version but not the previous one'],
        ['existing', 'Threat is in both versions, unchanged'],
        ['modified', 'Threat is in both versions but content changed (severity, mitigation, risk score)'],
        ['resolved', "Threat's target component was removed from the architecture"],
    ]
)
add_para(
    "This turns threat modeling from a one-off artefact into a continuous compliance record. "
    "Auditors get a diff per revision; security leads get a burn-down chart of open threats."
)

# ══════════════════════════════════════════════════════════════════════
# 10. SAST PRIORITIZATION
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Prioritizer — rerank SAST findings using the threat model', level=1)
add_para(
    "The threat model is not a standalone artefact. The prioritizer consumes it as context for "
    "every SAST finding:"
)
add_table(
    ['Signal', 'Adjustment'],
    [
        ['Finding is on a component the threat model calls out (≥2 strong token overlaps)', '+2 tiers (capped at critical)'],
        ['Finding on internet-facing / untrusted component', '+1 tier'],
        ["Finding's CWE is explicitly listed as a threat on that component", 'Marked high-confidence, +2 tiers'],
        ['Low-confidence match (category/keyword only)', 'Tag with context, never change severity'],
        ['Trusted / internal component, low severity', '-1 tier (de-prioritize)'],
        ['Test / fixture files', 'Excluded from matching'],
    ]
)
add_para(
    "Each finding gets a threat_model_context block with the matched components, confidence, "
    "CWE confirmation, and the list of rerank_reasons — so the analyst can see WHY severity "
    "changed."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 11. DATABASE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Persistence', level=1)
add_table(
    ['Table', 'Purpose'],
    [
        ['threat_models', 'Latest threat model per project + all analyses (STRIDE, attack paths, FAIR, kill chain, diagrams)'],
        ['architecture_versions', 'Versioned snapshots of the parsed architecture + change summary'],
        ['threat_history', 'Each threat tracked per version (new / existing / modified / resolved), linked list via previous_history_id'],
        ['security_controls', 'Registered controls with status, type, effectiveness, STRIDE coverage, linked_threat_ids'],
        ['client_threat_intel', 'Customer-uploaded threats, incidents, actors'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 12. API
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('12. API endpoints (key ones)', level=1)
add_table(
    ['Method', 'Path', 'Purpose'],
    [
        ['POST', '/api/projects/{id}/threat-model/regenerate', 'Start async generation (supports quick_mode)'],
        ['GET',  '/api/projects/{id}/threat-model', 'Full threat model + diagrams'],
        ['GET',  '/api/projects/{id}/threat-model/status', 'Poll generation progress'],
        ['DELETE', '/api/projects/{id}/threat-model', 'Delete threat model'],
        ['PUT',  '/api/projects/{id}/threat-model/threats/{tid}/status', 'Update threat status (open/mitigated/accepted/resolved)'],
        ['POST', '/api/projects/{id}/threat-model/generate-attack-diagram', 'On-demand attack-path diagram'],
        ['GET',  '/api/projects/{id}/threat-model/history', 'List all versions'],
        ['GET',  '/api/projects/{id}/threat-model/diff/{v1}/{v2}', 'Diff two versions'],
        ['GET',  '/api/projects/{id}/threats/{tid}/timeline', 'Timeline of a specific threat across versions'],
        ['POST', '/api/projects/{id}/threat-model/incremental', 'Regenerate only for changed components'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 13. DEMO SCRIPT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('13. Suggested demo script (12-15 minutes)', level=1)
add_table(
    ['#', 'Step', 'What to say', 'What to click'],
    [
        ['1', 'Hook (30 s)', "Threat modeling without weeks of workshops — from architecture doc to dollar-loss number in minutes.", 'Open Threat Model tab on a project'],
        ['2', 'Input (1 min)', "Paste the architecture description. Optionally upload a diagram. Click Generate.", 'Show the architecture field, then hit Generate'],
        ['3', 'Parsed model (1 min)', "The LLM pulls out components, flows, and trust boundaries automatically. Internet-facing flags drive later risk scores.", 'Expand parsed components list'],
        ['4', 'Context injection (1 min)', "We feed the LLM your sector threats, uploaded intel, and SecReq abuse cases — that is why the output talks about FIN7, not 'a generic attacker'.", 'Show context panel'],
        ['5', 'STRIDE threats (2 min)', "Per component, per category, with CWE, MITRE, severity, and AI-enriched business impact on top threats. Generic template fallback for the rest.", 'Filter STRIDE tab by Critical; open one threat and show attack_vector + business_impact'],
        ['6', 'Attack paths (2 min)', "We enumerate every path from internet to sensitive data. Top 5 get a penetration-tester narrative.", 'Attack Paths tab; open the highest-risk path'],
        ['7', 'Kill Chain + MITRE (1 min)', "SOC-ready view: every threat tagged with kill-chain phase and MITRE tactic.", 'Kill Chain tab'],
        ['8', 'FAIR risk (1 min)', "Dollar-value loss per threat per year. Deterministic — the CFO can re-run with new assumptions.", 'FAIR tab; show ALE (min/likely/max)'],
        ['9', 'Versioning (1 min)', "Re-run after an architecture change. Threats get new/modified/resolved labels automatically.", 'History tab; show a diff between two versions'],
        ['10', 'SAST re-ranking (1 min)', "The threat model becomes a prior on the SAST pipeline. Findings on critical components get promoted; findings on sandboxed components get demoted.", 'Switch to Vulnerabilities tab, show threat_model_context on a finding'],
        ['11', 'Close', 'Deterministic scoring + AI narrative + FAIR dollars + versioned history.', '—'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 14. Q&A
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('14. Likely client questions & short answers', level=1)
add_table(
    ['Question', 'Answer'],
    [
        ['How much of this is "AI" vs rule-based?', "The scaffold (parsing, MITRE mapping, kill chain, FAIR) is deterministic. AI only enriches the top 15 threats and top 5 attack paths. Quick mode skips AI entirely."],
        ['Can we trust the FAIR numbers?', "FAIR itself is a published standard (The Open Group). Our calculation is fully deterministic — your CFO can re-run with different LEF/LM assumptions, no LLM in the loop."],
        ['What if our architecture description is bad?', "The LLM extracts what is there; missing components produce no threats rather than hallucinated ones. Analysts edit the parsed component list directly before regenerating if needed."],
        ['Which LLM do you use?', "Whatever the customer configures. Anthropic Claude, OpenAI, Azure OpenAI — all via the shared AI factory. Model is recorded per generation."],
        ['What about privacy? We cannot send our architecture to OpenAI.', "Two answers: (1) point the platform at a private Azure / self-hosted LLM endpoint, or (2) run in quick mode — fully deterministic, no LLM calls at all."],
        ['How does this integrate with our SAST?', "The prioritizer uses the threat model as a prior for every SAST finding. Findings on internet-facing or sensitive components get bumped; findings on internal tooling get demoted."],
        ['What about compliance reporting?', "Every threat carries CWE + MITRE. Compliance modules (OWASP ASVS, PCI-DSS) come from the SecureReq feature which plugs into this threat model."],
        ['How fast is a regeneration?', "Quick mode: seconds. Full mode with AI enrichment: 30-90 seconds depending on component count and LLM latency."],
        ['Can we add custom threats?', "Yes — via client_threat_intel uploads, and via SecureReq abuse cases (which become explicit threats with source=\"securereq\")."],
        ['What does it cost?', "One LLM call to parse architecture + ~15 calls for threat enrichment + ~5 calls for attack paths = roughly 20-25 calls per full generation. Quick mode is zero LLM cost."],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 15. FILE MAP
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('15. Technical appendix — file map', level=1)
add_table(
    ['File', 'Purpose'],
    [
        ['backend/services/threat_modeling.py', 'Main pipeline — architecture ingestion, STRIDE, attack paths, kill chain, FAIR, diagrams, prompts'],
        ['backend/services/threat_model_prioritizer.py', 'Reranks SAST findings using the threat model as a prior'],
        ['backend/services/threat_lifecycle_service.py', 'Versioning: stable IDs + new/modified/resolved status transitions'],
        ['backend/main.py — _generate_threat_model_background()', 'Orchestration: context assembly + pipeline invocation + persistence'],
        ['frontend/src/pages/ThreatModelPage.tsx', 'Multi-tab UI (STRIDE, Attack Paths, Kill Chain, FAIR, History, Controls)'],
    ]
)

out_path = os.path.join(os.path.dirname(__file__), 'Threat_Modeling_Client_Demo_Guide.docx')
doc.save(out_path)
print(f"Document saved to: {out_path}")
