"""Generate Threat Modeling Pipeline Technical Document as Word (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text):
    """Add a formatted code block"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    # Add shading
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)
    return p


def add_table(headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)
    return p


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SecureDev AI')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Threat Modeling Pipeline')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Technical Document')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x64, 0x64, 0x80)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('How Threat Intelligence, Security Requirements,\nand AI-Powered Analysis Combine to Generate\nComprehensive STRIDE Threat Models')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x64, 0x80)
run.italic = True

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Confidential').bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1.', 'Pipeline Overview'),
    ('2.', 'Stage 1: Architecture Ingestion'),
    ('3.', 'Stage 2: Intelligence Context Assembly'),
    ('  3.1', 'Source A — Sector Threat Intelligence'),
    ('  3.2', 'Source B — Client-Uploaded Threat Intelligence'),
    ('  3.3', 'Source C — SecReq Abuse Cases & Requirements'),
    ('  3.4', 'Source D — Existing Security Controls'),
    ('  3.5', 'Final Assembled Context'),
    ('4.', 'Stage 3: STRIDE Threat Generation'),
    ('  4.1', 'Pass 1 — Template-Based Generation'),
    ('  4.2', 'Pass 2 — AI Enrichment (Top 15 Threats)'),
    ('  4.3', 'Pass 3 — SecReq Threat Injection'),
    ('5.', 'Stage 4: Attack Path, MITRE & Kill Chain'),
    ('6.', 'Stage 5: Risk Quantification & Diagrams'),
    ('7.', 'How Threat Intel & SecReq Are Consumed'),
    ('8.', 'Value Comparison: With vs Without Intelligence'),
    ('9.', 'Recommended Workflow'),
    ('10.', 'Appendix: Key Files Reference'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)
    if not num.startswith(' '):
        run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. PIPELINE OVERVIEW
# ══════════════════════════════════════════════════════════════════

doc.add_heading('1. Pipeline Overview', level=1)

doc.add_paragraph(
    'The threat model is generated through a 5-stage pipeline. Each stage builds on the previous one. '
    'Threat Intelligence and Security Requirements (SecReq) are injected at Stage 2 as context that '
    'shapes the AI\'s analysis throughout all subsequent stages.'
)

add_code_block(
    '┌─────────────────────────────────────────────────────────────────────┐\n'
    '│                    THREAT MODELING PIPELINE                         │\n'
    '│                                                                     │\n'
    '│  ┌──────────┐   ┌──────────────┐   ┌─────────┐   ┌──────────────┐ │\n'
    '│  │ Stage 1  │──>│   Stage 2    │──>│ Stage 3 │──>│   Stage 4    │ │\n'
    '│  │Arch Parse│   │Intel Assembly│   │ STRIDE  │   │Attack Paths  │ │\n'
    '│  └──────────┘   └──────────────┘   └─────────┘   │MITRE Mapping │ │\n'
    '│                       ▲                           │Kill Chain    │ │\n'
    '│                       │                           └──────────────┘ │\n'
    '│              ┌────────┼────────┐          │                        │\n'
    '│              │        │        │          ▼                        │\n'
    '│         ┌────┴──┐ ┌───┴───┐ ┌─┴────┐ ┌──────────────┐            │\n'
    '│         │Sector │ │Client │ │SecReq│ │   Stage 5    │            │\n'
    '│         │Threat │ │Threat │ │Abuse │ │FAIR Risk     │            │\n'
    '│         │Intel  │ │Intel  │ │Cases │ │Eraser Diagrams│           │\n'
    '│         └───────┘ └───────┘ └──────┘ └──────────────┘            │\n'
    '└─────────────────────────────────────────────────────────────────────┘'
)

add_table(
    ['Stage', 'Purpose', 'AI Calls', 'Intel Used'],
    [
        ['1. Architecture Parse', 'Extract components, data flows, trust boundaries', '1 call', 'None'],
        ['2. Intel Assembly', 'Gather sector intel, client intel, SecReq, controls', '0 calls', 'All sources gathered'],
        ['3. STRIDE Generation', 'Generate threats per component × STRIDE category', 'Up to 15 calls', 'Full context in each AI call'],
        ['4. Attack Paths + MITRE', 'Find attack chains, map to MITRE ATT&CK', 'Up to 5 calls', 'Full context in each AI call'],
        ['5. FAIR Risk + Diagrams', 'Dollar-based risk estimates, professional diagrams', '3 Eraser API calls', 'Indirect (uses threat data)'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. STAGE 1
# ══════════════════════════════════════════════════════════════════

doc.add_heading('2. Stage 1: Architecture Ingestion', level=1)

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'Parse the user-provided architecture (text description, uploaded diagrams, or Architecture Builder input) '
    'into a structured representation of components, data flows, and trust boundaries.'
)

doc.add_heading('How It Works', level=2)
doc.add_paragraph(
    'The architecture document is sent to the AI with a structured extraction prompt. The AI returns a JSON object containing:'
)
add_bullet('Components — each service, database, frontend, external integration')
add_bullet('Data Flows — how components communicate, encryption and authentication status')
add_bullet('Trust Boundaries — where trust levels change (e.g., Internet ↔ DMZ)')
add_bullet('Component Properties — internet-facing, handles sensitive data, trust level, technology stack')

doc.add_heading('AI Prompt Used', level=2)
p = doc.add_paragraph()
run = p.add_run('File: ')
run.bold = True
p.add_run('backend/services/threat_modeling.py, line 508')
p = doc.add_paragraph()
run = p.add_run('Function: ')
run.bold = True
p.add_run('analyze_architecture_with_ai()')
p = doc.add_paragraph()
run = p.add_run('AI Parameters: ')
run.bold = True
p.add_run('max_tokens=4000, temperature=0.3')

add_code_block(
    'Analyze this software architecture and extract a detailed threat model structure.\n\n'
    'ARCHITECTURE DESCRIPTION:\n'
    '{architecture_doc}\n\n'
    'Please provide a comprehensive JSON response with the following structure:\n'
    '{\n'
    '    "system_overview": "Brief description of what the system does",\n'
    '    "technology_stack": ["list", "of", "technologies"],\n'
    '    "components": [\n'
    '        {\n'
    '            "id": "unique_id",\n'
    '            "name": "Component Name",\n'
    '            "type": "external|process|datastore",\n'
    '            "technology": "specific technology (e.g., React, PostgreSQL)",\n'
    '            "category": "api|database|authentication|frontend|microservice|cloud|message_queue",\n'
    '            "description": "What this component does",\n'
    '            "data_handled": ["types of data this component handles"],\n'
    '            "trust_level": "untrusted|semi-trusted|trusted|highly-trusted",\n'
    '            "internet_facing": true/false,\n'
    '            "handles_sensitive_data": true/false\n'
    '        }\n'
    '    ],\n'
    '    "data_flows": [\n'
    '        {\n'
    '            "id": "flow_id",\n'
    '            "from": "source_component_id",\n'
    '            "to": "target_component_id",\n'
    '            "data_type": "What data flows",\n'
    '            "protocol": "HTTP/HTTPS/gRPC/SQL/etc",\n'
    '            "encrypted": true/false,\n'
    '            "authenticated": true/false,\n'
    '            "sensitive": true/false\n'
    '        }\n'
    '    ],\n'
    '    "trust_boundaries": [\n'
    '        {\n'
    '            "id": "boundary_id",\n'
    '            "name": "Boundary Name",\n'
    '            "description": "What this boundary separates",\n'
    '            "components_inside": ["list of component_ids"],\n'
    '            "boundary_type": "internet|dmz|internal|data"\n'
    '        }\n'
    '    ],\n'
    '    "security_controls": ["list of mentioned security controls"],\n'
    '    "risk_factors": ["identified risk factors from the architecture"]\n'
    '}\n\n'
    'Be thorough and extract all components, even if implied.\n'
    'Identify ALL data flows between components.'
)

doc.add_heading('Fallback', level=2)
doc.add_paragraph(
    'If AI is unavailable, a keyword-based parser (parse_architecture_basic) extracts components by matching '
    'terms like "api", "database", "auth", "frontend", "kafka", "redis", etc.'
)

doc.add_heading('Output', level=2)
doc.add_paragraph(
    'A parsed architecture dictionary used by all subsequent stages. For Apex Banking, this produced '
    '14+ components across 6 trust boundaries.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. STAGE 2
# ══════════════════════════════════════════════════════════════════

doc.add_heading('3. Stage 2: Intelligence Context Assembly', level=1)

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'Before generating threats, gather all available intelligence to create a rich context document '
    'that accompanies every AI call. This is where Threat Intel and SecReq are consumed.'
)

doc.add_heading('How It Works', level=2)
doc.add_paragraph(
    'The system assembles context from 4 sources, each described below. All sources are concatenated '
    'into a single system_context string that is passed to every AI prompt in Stages 3 and 4.'
)

# Source A
doc.add_heading('3.1 Source A: Sector Threat Intelligence', level=2)
doc.add_paragraph(
    'Pulls industry-specific threats based on the project\'s industry_sector field (e.g., "banking", '
    '"healthcare", "retail"). Uses built-in sector threat databases covering known attack patterns, '
    'threat actors, and regulatory requirements for each industry.'
)

add_code_block(
    '# Code: main.py, lines 2459-2462\n'
    'sector_threats = get_sector_threats(industry)  # e.g., "banking"\n'
    'if sector_threats:\n'
    '    threat_intel_context = format_intel_for_prompt(sector_threats, max_entries=12)'
)

doc.add_paragraph('Each entry is formatted as:')
add_code_block(
    '### FIN7 Targeting US Banking Platforms [CRITICAL]\n'
    'Type: threat_actor\n'
    'Description: FIN7 actively targets banking platforms using spear-phishing...\n'
    'MITRE ATT&CK: T1566.001, T1059.007, T1055\n'
    'Regulatory: PCI-DSS, SOX\n'
    '  - Network segmentation for payment systems\n'
    '  - Application whitelisting'
)

# Source B
doc.add_heading('3.2 Source B: Client-Uploaded Threat Intelligence', level=2)
doc.add_paragraph(
    'Your organization\'s custom intel entries from the Threat Intel page. These are entries your team '
    'uploads: threat actors, past incidents, pentest findings, regulations, and attack scenarios. '
    'Each entry carries MITRE techniques, severity, regulatory impact, and recommended controls.'
)

add_code_block(
    '# Code: main.py, lines 2466-2491\n'
    'cursor.execute(\n'
    '    "SELECT * FROM client_threat_intel WHERE project_id = ? AND active = 1 "\n'
    '    "ORDER BY severity DESC LIMIT 15",\n'
    '    (project_id,)\n'
    ')\n'
    'client_rows = [dict(r) for r in cursor.fetchall()]\n\n'
    'if client_rows:\n'
    '    client_context = format_intel_for_prompt(client_intel, max_entries=10)\n'
    '    threat_intel_context += "\\n\\n=== CLIENT-SPECIFIC THREAT INTELLIGENCE ===\\n"\n'
    '    threat_intel_context += client_context'
)

doc.add_paragraph(
    'For Apex Banking, 8 client entries were loaded including: FIN7 targeting, Magecart skimming, '
    'credential stuffing, SIM swap fraud, and BSA/AML regulation requirements.'
)

# Source C
doc.add_heading('3.3 Source C: SecReq Abuse Cases & Requirements', level=2)
doc.add_paragraph(
    'Pulled from analyzed user stories in the SecReq module. Only user stories that have been run through '
    'the SecReq AI analysis contribute data. Unanalyzed stories are not included.'
)

add_bullet('Abuse cases = "what can go wrong" for each user story', bold_prefix='Abuse Cases:')
add_bullet('Security requirements = "what must be true" for secure implementation', bold_prefix='Requirements:')

add_code_block(
    '# Code: main.py, lines 2498-2538\n'
    'analyses = (\n'
    '    db.query(SecurityAnalysis)\n'
    '    .join(UserStory)\n'
    '    .filter(UserStory.project_id == project_id)\n'
    '    .order_by(SecurityAnalysis.id.desc())\n'
    '    .all()\n'
    ')\n\n'
    'for analysis in analyses:\n'
    '    # Abuse cases → become explicit threats in Stage 3\n'
    '    if analysis.abuse_cases:\n'
    '        for ac in analysis.abuse_cases[:5]:\n'
    '            securereq_abuse_cases.append(ac)\n'
    '            # Format: "- [Critical] Social engineering bypass (Actor: Insider, STRIDE: Spoofing)"\n\n'
    '    # Security requirements → used for coverage matrix\n'
    '    if analysis.security_requirements:\n'
    '        for req in analysis.security_requirements[:5]:\n'
    '            securereq_requirements.append(req)\n'
    '            # Format: "- [Critical] [Cryptography] Implement HSM-backed signing"'
)

p = doc.add_paragraph()
run = p.add_run('Important: ')
run.bold = True
p.add_run(
    'User stories must be individually analyzed in the SecReq page before they contribute to threat modeling. '
    'Simply adding stories without analyzing them will not generate abuse cases or requirements.'
)

# Source D
doc.add_heading('3.4 Source D: Existing Security Controls', level=2)
doc.add_paragraph(
    'Pulls already-registered controls from the Controls tab. This tells the AI what mitigations are '
    'already in place so it recommends what\'s missing rather than duplicating existing controls.'
)

add_code_block(
    '# Code: main.py, lines 2544-2561\n'
    'controls = db.query(SecurityControl).filter(\n'
    '    SecurityControl.project_id == project_id\n'
    ').all()\n\n'
    '# Format: "- [IMPLEMENTED] WAF (type: preventive, effectiveness: 85%, covers: Tampering)"'
)

# Final assembled context
doc.add_heading('3.5 Final Assembled Context', level=2)
doc.add_paragraph(
    'All four sources are concatenated into a single system context document. This is passed to every '
    'AI call in Stages 3 and 4:'
)

add_code_block(
    'System Overview: Enterprise digital banking platform...\n'
    'Technology Stack: Spring Boot, React, PostgreSQL, Kafka, Redis...\n\n'
    'Components (14):\n'
    '- API Gateway (api): Internet-facing, Handles sensitive data\n'
    '- Auth Service (authentication): Internal, Handles sensitive data\n'
    '- Fund Transfer Service (microservice): Internal, Handles sensitive data\n'
    '...\n\n'
    'Data Flows (20):\n'
    '- API Gateway → Auth Service: Encrypted, Authenticated\n'
    '- Transfer Service → SWIFT Network: Encrypted, Authenticated\n'
    '...\n\n'
    '=== THREAT INTELLIGENCE (BANKING SECTOR) ===\n\n'
    '### FIN7 Targeting US Banking Platforms [CRITICAL]\n'
    'Type: threat_actor\n'
    'Description: FIN7 actively targets banking platforms...\n'
    'MITRE ATT&CK: T1566.001, T1059.007, T1055\n\n'
    '=== CLIENT-SPECIFIC THREAT INTELLIGENCE ===\n\n'
    '### 2024 SWIFT Credential Theft Attempt [CRITICAL]\n'
    'Type: incident\n'
    'Regulatory: PCI-DSS Req 7.1, SOX Section 404\n\n'
    '=== EXISTING SECURITY CONTROLS ===\n'
    '- [IMPLEMENTED] WAF (type: preventive, effectiveness: 85%)\n\n'
    '=== SECURITY REQUIREMENTS ANALYSIS ===\n'
    '- [Critical] Social engineering to bypass approvals (STRIDE: Spoofing)\n'
    '- [Critical] [Cryptography] Implement HSM-backed signing for MT103 messages'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. STAGE 3
# ══════════════════════════════════════════════════════════════════

doc.add_heading('4. Stage 3: STRIDE Threat Generation', level=1)

doc.add_heading('Purpose', level=2)
doc.add_paragraph(
    'Generate comprehensive STRIDE threats for every component in the architecture, enriched with '
    'AI-powered analysis for the highest-risk threats. This stage produces the core deliverable of the '
    'threat model.'
)

# Pass 1
doc.add_heading('4.1 Pass 1: Template-Based Generation', level=2)
doc.add_paragraph(
    'For each component × each STRIDE category, the system matches technology-specific threat templates. '
    'This produces the bulk of threats (500+ for complex architectures) without requiring AI API calls.'
)

add_code_block(
    'for component in parsed_arch[\'components\']:\n'
    '    category = component[\'category\']  # e.g., "api", "authentication"\n'
    '    tech_threats = TECHNOLOGY_THREATS.get(category, {})\n\n'
    '    for stride_cat, threats in tech_threats.items():\n'
    '        for threat_template in threats:\n'
    '            # Calculate risk score\n'
    '            base_score = {\'critical\': 9.5, \'high\': 7.5, \'medium\': 5.0, \'low\': 2.5}[severity]\n\n'
    '            # Adjust for component properties\n'
    '            if component[\'internet_facing\']:        base_score += 1.0\n'
    '            if component[\'handles_sensitive_data\']: base_score += 0.5\n'
    '            if component[\'trust_level\'] == \'untrusted\': base_score += 0.5\n\n'
    '            risk_score = min(10.0, base_score)'
)

doc.add_paragraph(
    'Each threat includes: component name, STRIDE category, severity, risk score, CWE mapping, '
    'MITRE technique IDs, template description, and template mitigation.'
)

# Pass 2
doc.add_heading('4.2 Pass 2: AI Enrichment (Top 15 Threats)', level=2)
doc.add_paragraph(
    'The 15 highest-risk threats are sent to Claude/GPT with the full intelligence context from Stage 2. '
    'The AI generates context-aware descriptions, attack vectors, business impact, and mitigations that '
    'reference your specific architecture, threat actors, and existing controls.'
)

p = doc.add_paragraph()
run = p.add_run('AI Parameters: ')
run.bold = True
p.add_run('max_tokens=2000')

doc.add_paragraph('AI Enrichment Prompt:')

add_code_block(
    'You are a senior application security expert performing threat modeling.\n'
    'Analyze this specific threat and provide detailed, contextual information.\n\n'
    'SYSTEM CONTEXT:\n'
    '{system_context}\n'
    '<-- Full assembled context from Stage 2 including:\n'
    '    - Architecture components and data flows\n'
    '    - Sector threat intel (banking threats)\n'
    '    - Client-uploaded threat intel (FIN7, Magecart, incidents)\n'
    '    - SecReq abuse cases and requirements\n'
    '    - Existing security controls\n\n'
    'COMPONENT BEING ANALYZED:\n'
    '- Name: {component_name}\n'
    '- Type: {component_type}\n'
    '- Technology: {technology}\n'
    '- Internet Facing: {internet_facing}\n'
    '- Handles Sensitive Data: {handles_sensitive_data}\n'
    '- Trust Level: {trust_level}\n\n'
    'THREAT TO ANALYZE:\n'
    '- Threat Name: {threat_name}\n'
    '- STRIDE Category: {stride_category}\n'
    '- Severity: {severity}\n'
    '- CWE: {cwe}\n'
    '- MITRE Techniques: {mitre_techniques}\n\n'
    'Provide a detailed JSON response:\n'
    '{\n'
    '    "description": "2-3 sentence description specific to this component",\n'
    '    "attack_vector": {\n'
    '        "description": "How an attacker would exploit this",\n'
    '        "entry_points": ["Specific entry points"],\n'
    '        "techniques": ["Specific attack techniques"]\n'
    '    },\n'
    '    "business_impact": {\n'
    '        "financial": "Dollar-based financial impact",\n'
    '        "reputational": "Reputational damage assessment",\n'
    '        "operational": "Operational impact",\n'
    '        "compliance": "Regulatory implications"\n'
    '    },\n'
    '    "affected_assets": ["Assets at risk"],\n'
    '    "prerequisites": {\n'
    '        "access_required": "Access level needed",\n'
    '        "conditions": ["Conditions for attack success"]\n'
    '    },\n'
    '    "attack_complexity": {\n'
    '        "level": "Low/Medium/High",\n'
    '        "skill_level": "Basic/Intermediate/Advanced",\n'
    '        "time_required": "Estimated time"\n'
    '    },\n'
    '    "mitigation": "Specific, actionable recommendations",\n'
    '    "detection": "How to detect this attack"\n'
    '}\n\n'
    'Be specific and technical. Reference actual component names and technology.'
)

# Pass 3
doc.add_heading('4.3 Pass 3: SecReq Threat Injection', level=2)
doc.add_paragraph(
    'Each analyzed abuse case from SecReq is injected as a dedicated STRIDE threat tagged with '
    'source="securereq" for traceability. A coverage matrix is also built showing which security '
    'requirements are covered by generated threats and which gaps remain.'
)

add_code_block(
    '# Code: threat_modeling.py, line 772 — _inject_securereq_threats()\n\n'
    'for ac in abuse_cases:\n'
    '    threat_obj = {\n'
    '        "id": f"securereq_{ac_id}",\n'
    '        "component": "System-wide",\n'
    '        "category": stride_category,  # Mapped from abuse case\n'
    '        "threat": ac[\'title\'],\n'
    '        "description": ac[\'description\'],\n'
    '        "severity": severity,  # Derived from impact text\n'
    '        "source": "securereq",  # Tagged for traceability\n'
    '        "abuse_case_id": ac_id,\n'
    '        "mitre_techniques": ac.get(\'mitre_techniques\', []),\n'
    '        "threat_actor": ac.get(\'actor\', \'\'),\n'
    '    }\n'
    '    stride_analysis[stride_cat].append(threat_obj)'
)

doc.add_paragraph('Coverage matrix output:')

add_code_block(
    'coverage = {\n'
    '    "summary": {\n'
    '        "total_requirements": 5,\n'
    '        "covered": 3,\n'
    '        "uncovered": 2,\n'
    '        "coverage_percentage": 60.0\n'
    '    },\n'
    '    "requirements": [\n'
    '        {\n'
    '            "id": "SR-001",\n'
    '            "requirement": "Implement HSM-backed signing for MT103 messages",\n'
    '            "coverage_status": "uncovered"\n'
    '        }\n'
    '    ]\n'
    '}'
)

doc.add_heading('Stage 3 Output', level=2)
doc.add_paragraph(
    'A complete STRIDE analysis dictionary with threats categorized under: Spoofing, Tampering, '
    'Repudiation, Information Disclosure, Denial of Service, Elevation of Privilege. '
    'For Apex Banking: 541 threats (536 template-based + 5 SecReq-injected, with 15 AI-enriched).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. STAGE 4
# ══════════════════════════════════════════════════════════════════

doc.add_heading('5. Stage 4: Attack Path, MITRE & Kill Chain', level=1)

doc.add_heading('5.1 Attack Path Generation', level=2)

doc.add_paragraph('How it works:')
add_bullet('Identify entry points (internet-facing components) and targets (sensitive data stores)')
add_bullet('Build an adjacency graph from data flows')
add_bullet('Find all paths from entry points to targets using BFS (max depth 5)')
add_bullet('Score each path based on threats along the route')
add_bullet('Top 10 paths selected by risk score')
add_bullet('Top 5 paths are AI-enriched with the full system context (including threat intel)')

doc.add_paragraph('Attack Path AI Prompt:')

add_code_block(
    'You are a senior penetration tester analyzing an attack path.\n\n'
    'SYSTEM CONTEXT:\n'
    '{system_context}  <-- Same full context including threat intel\n\n'
    'ATTACK PATH:\n'
    'Entry Point: API Gateway (api)\n'
    'Target: Core Banking DB (database)\n'
    'Path: API Gateway → Auth Service → Transfer Service → Core Banking DB\n\n'
    'THREATS ALONG THIS PATH:\n'
    '- Request Smuggling (critical) at API Gateway\n'
    '- JWT Token Forgery (high) at Auth Service\n'
    '- SQL Injection (critical) at Core Banking DB\n\n'
    'Generate a detailed JSON response:\n'
    '{\n'
    '    "attack_scenario": "3-4 sentence narrative of the attack",\n'
    '    "exploitation_steps": [\n'
    '        {"step": 1, "phase": "Initial Access", "action": "...", "details": "..."}\n'
    '    ],\n'
    '    "potential_impact": {\n'
    '        "level": "Critical", "data_exposure": "...",\n'
    '        "business_impact": "...", "compliance_impact": "..."\n'
    '    },\n'
    '    "difficulty": {"level": "Medium", "tools_needed": ["..."]},\n'
    '    "detection_opportunities": [\n'
    '        {"point": "...", "method": "...", "effectiveness": "High"}\n'
    '    ],\n'
    '    "recommended_controls": [\n'
    '        {"control": "...", "implementation": "...", "priority": "Critical"}\n'
    '    ]\n'
    '}'
)

doc.add_heading('5.2 MITRE ATT&CK Mapping', level=2)
doc.add_paragraph(
    'Template-based (no AI calls). Maps each threat\'s CWE and MITRE technique IDs to the full '
    'MITRE ATT&CK framework. Groups techniques by tactic: Initial Access, Execution, Persistence, '
    'Privilege Escalation, Defense Evasion, Credential Access, Discovery, Lateral Movement, '
    'Collection, Exfiltration, Command & Control, Impact.'
)

doc.add_heading('5.3 Kill Chain Analysis', level=2)
doc.add_paragraph(
    'Maps STRIDE threats to Cyber Kill Chain phases: Reconnaissance → Weaponization → Delivery → '
    'Exploitation → Installation → Command & Control → Actions on Objectives.'
)

doc.add_heading('5.4 Attack Trees', level=2)
doc.add_paragraph(
    'Hierarchical decomposition of high-risk goals showing AND/OR relationships between attack sub-goals.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. STAGE 5
# ══════════════════════════════════════════════════════════════════

doc.add_heading('6. Stage 5: Risk Quantification & Diagrams', level=1)

doc.add_heading('6.1 FAIR Risk Quantification', level=2)
doc.add_paragraph(
    'Uses the FAIR (Factor Analysis of Information Risk) model to calculate dollar-based loss estimates '
    'with confidence intervals. No AI calls — pure calculation.'
)

add_table(
    ['Parameter', 'Purpose', 'Example'],
    [
        ['organization_size', 'Scale multiplier (small=0.3x, enterprise=5.0x)', 'medium (1.0x)'],
        ['industry', 'Risk multiplier (finance=1.8x, healthcare=1.5x)', 'banking (1.8x)'],
        ['annual_revenue', 'Base for financial impact calculations', '$10,000,000'],
        ['customer_count', 'Base for data breach impact calculations', '10,000'],
    ]
)

doc.add_paragraph('For each threat, FAIR calculates:')
add_bullet('Loss Event Frequency (LEF) = Threat Event Frequency × Vulnerability')
add_bullet('Loss Magnitude (LM) across 6 categories: Productivity, Response, Replacement, Fines, Reputation, Legal')
add_bullet('Annualized Loss Expectancy (ALE) = LEF × LM, with industry and size multipliers')
add_bullet('Confidence intervals: minimum, likely, and maximum estimates')

doc.add_heading('6.2 Eraser AI Diagrams', level=2)
doc.add_paragraph('Generates 3 professional diagrams via the Eraser.io API (requires ERASER_API_KEY):')
add_bullet('Architecture/Threat Model Diagram — Components with top threats overlaid')
add_bullet('Kill Chain Diagram — Visual kill chain with threats mapped to phases')
add_bullet('Data Flow Diagram — Components, data flows, and trust boundaries')
doc.add_paragraph('Attack path diagrams are generated on-demand via the UI.')

doc.add_heading('6.3 Mermaid DFDs', level=2)
doc.add_paragraph(
    'Level 0 (high-level) and Level 1 (detailed) Data Flow Diagrams rendered as Mermaid syntax in the browser. '
    'Always available as fallback when Eraser is not configured.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. HOW THREAT INTEL & SECUREREQ ARE CONSUMED
# ══════════════════════════════════════════════════════════════════

doc.add_heading('7. How Threat Intel & SecReq Are Consumed', level=1)

doc.add_heading('Data Flow', level=2)

add_code_block(
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│                        THREAT INTEL                              │\n'
    '│                                                                  │\n'
    '│  ┌─────────────┐     ┌──────────────┐     ┌──────────────┐     │\n'
    '│  │Sector Intel │     │Client Upload │     │Live Feeds    │     │\n'
    '│  │(built-in DB)│     │(your entries)│     │(CISA/NVD)    │     │\n'
    '│  └──────┬──────┘     └──────┬──────┘     │[NOT USED YET]│     │\n'
    '│         │                   │             └──────────────┘     │\n'
    '│         ▼                   ▼                                   │\n'
    '│  ┌──────────────────────────────────┐                          │\n'
    '│  │   format_intel_for_prompt()      │                          │\n'
    '│  └──────────────┬───────────────────┘                          │\n'
    '│                 ▼                                               │\n'
    '│  ┌──────────────────────────────────┐                          │\n'
    '│  │      threat_intel_context        │                          │\n'
    '│  └──────────────┬───────────────────┘                          │\n'
    '└─────────────────┼──────────────────────────────────────────────┘\n'
    '                  │\n'
    '┌─────────────────┼──────────────────────────────────────────────┐\n'
    '│                 ▼        SECUREREQ                              │\n'
    '│  ┌──────────────┐     ┌──────────────┐                        │\n'
    '│  │ User Stories │────>│  AI Analysis │                        │\n'
    '│  └──────────────┘     └──────┬───────┘                        │\n'
    '│                    ┌─────────┴─────────┐                      │\n'
    '│                    ▼                   ▼                       │\n'
    '│           ┌──────────────┐   ┌──────────────┐                │\n'
    '│           │ Abuse Cases  │   │ Requirements │                │\n'
    '│           └──────┬───────┘   └──────┬───────┘                │\n'
    '│                  ▼                  ▼                          │\n'
    '│  ┌──────────────────────────────────────┐                    │\n'
    '│  │      securereq_context               │                    │\n'
    '│  └──────────────┬───────────────────────┘                    │\n'
    '└─────────────────┼──────────────────────────────────────────────┘\n'
    '                  ▼\n'
    '┌──────────────────────────────────────────────────────────────────┐\n'
    '│                    SYSTEM CONTEXT                                 │\n'
    '│  (architecture + threat intel + SecReq + controls)               │\n'
    '│                                                                  │\n'
    '│  Used by:                                                        │\n'
    '│  ├── AI Enrichment     → 15 critical/high threats               │\n'
    '│  ├── Attack Path AI    → Top 5 attack paths                     │\n'
    '│  └── SecReq Injection  → Direct threats + Coverage matrix       │\n'
    '└──────────────────────────────────────────────────────────────────┘'
)

doc.add_heading('Impact by Source', level=2)

add_table(
    ['Source', 'What It Influences', 'How'],
    [
        ['Sector Intel', 'AI-enriched descriptions, mitigations, business impact',
         'Passed as context — AI references industry-specific attack patterns'],
        ['Client Intel', 'AI-enriched descriptions; threat actor awareness',
         'AI sees your threat actors and generates threats mentioning their TTPs'],
        ['SecReq Abuse Cases', 'Direct STRIDE threats + coverage matrix',
         'Each abuse case becomes a tagged threat with traceability'],
        ['SecReq Requirements', 'Coverage matrix showing gaps',
         'Requirements compared against threats to show % coverage'],
        ['Existing Controls', 'AI-enriched mitigations',
         'AI knows what you have and recommends what\'s missing'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. VALUE COMPARISON
# ══════════════════════════════════════════════════════════════════

doc.add_heading('8. Value Comparison: With vs Without Intelligence', level=1)

add_table(
    ['Aspect', 'Without Intel/SecReq', 'With Intel/SecReq'],
    [
        ['Descriptions',
         '"Potential SQL Injection vulnerability in Database"',
         '"SQL Injection in Core Banking Oracle DB could enable unauthorized SWIFT transaction modification, similar to FIN7 techniques (T1059.007)"'],
        ['Mitigations',
         '"Use parameterized queries"',
         '"Implement parameterized queries with Oracle bind variables. Given existing WAF (85% effective), add SQL injection rules for MT103 parsing endpoints"'],
        ['Business Impact',
         '"Financial impact possible"',
         '"$500K-$10M based on banking industry breach data. Regulatory exposure under PCI-DSS Req 6.5.1 and SOX Section 404"'],
        ['Threat Actors',
         'Not considered',
         '"FIN7 actively targets banking platforms using T1566.001 spear-phishing to gain initial access to payment processing"'],
        ['Coverage',
         'No traceability to requirements',
         '"60% of security requirements covered. 2 critical gaps: HSM-backed signing, Dynamic MFA for approvals"'],
        ['Controls',
         'Generic recommendations',
         'Recommendations that account for existing controls and fill gaps'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. RECOMMENDED WORKFLOW
# ══════════════════════════════════════════════════════════════════

doc.add_heading('9. Recommended Workflow', level=1)

doc.add_paragraph('For maximum value from the threat modeling pipeline, follow these steps in order:')

steps = [
    ('Step 1: SecureReq',
     'Add user stories (manually or import from Jira). Analyze each story to generate abuse cases '
     'and security requirements. These feed directly into the threat model as explicit threats and '
     'a coverage matrix.'),
    ('Step 2: Threat Intel',
     'Upload client-specific intelligence: known threat actors, past incidents, pentest findings, '
     'regulatory requirements, and industry-specific attack scenarios. These inform the AI\'s threat '
     'descriptions, mitigations, and impact assessments.'),
    ('Step 3: Architecture',
     'Define the system architecture via text description, diagram upload, or the Architecture Builder. '
     'Include: components, data flows, trust boundaries, and technology stack.'),
    ('Step 4: Generate Threat Model',
     'The pipeline automatically ingests all of the above. Full mode provides AI enrichment for top 15 '
     'threats and top 5 attack paths. Quick mode uses templates only (faster but less contextual).'),
    ('Step 5: Review & Triage',
     'Review threats using severity, component, STRIDE, and status filters with pagination. '
     'AI-enriched threats contain banking-specific context. SecReq-derived threats are traceable to '
     'user stories. The coverage matrix highlights requirement gaps.'),
    ('Step 6: Register Controls',
     'Add security controls in the Controls tab (inside the Threat Model page). Link controls to '
     'threats. Track coverage and effectiveness percentages.'),
    ('Step 7: Iterate',
     'When architecture or intelligence changes, regenerate the threat model. The new model accounts '
     'for updated context. Incremental mode preserves existing threat status and analyst notes.'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. APPENDIX
# ══════════════════════════════════════════════════════════════════

doc.add_heading('10. Appendix: Key Files Reference', level=1)

add_table(
    ['File', 'Line', 'Function', 'Purpose'],
    [
        ['main.py', '2423', '_generate_threat_model_background()', 'Orchestrates pipeline, assembles intel context'],
        ['threat_modeling.py', '2314', 'generate_threat_model()', 'Core generation logic for all stages'],
        ['threat_modeling.py', '499', 'analyze_architecture_with_ai()', 'Stage 1: Architecture parsing'],
        ['threat_modeling.py', '87', '_enrich_threat_with_ai()', 'Stage 3: AI enrichment of top threats'],
        ['threat_modeling.py', '949', 'generate_stride_analysis()', 'Stage 3: STRIDE threat generation'],
        ['threat_modeling.py', '772', '_inject_securereq_threats()', 'Stage 3: SecReq threat injection'],
        ['threat_modeling.py', '1928', 'generate_attack_paths()', 'Stage 4: Attack path analysis'],
        ['threat_modeling.py', '175', '_generate_attack_path_with_ai()', 'Stage 4: AI-enriched attack paths'],
        ['threat_modeling.py', '2584', 'calculate_fair_risk()', 'Stage 5: FAIR risk quantification'],
        ['threat_modeling.py', '1752', 'generate_eraser_diagrams()', 'Stage 5: Professional diagrams'],
        ['sector_threat_intel.py', '2345', 'format_intel_for_prompt()', 'Formats intel for AI prompts'],
        ['routers/threat_intel.py', '-', 'router', 'Client threat intel CRUD + feeds'],
        ['routers/securereq.py', '-', 'router', 'User stories + security analysis'],
    ]
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Threat_Modeling_Pipeline_Technical_Document.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
