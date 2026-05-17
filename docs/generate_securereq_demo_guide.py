"""Generate SecureReq (Security Requirements) — Client Demo Guide as Word (.docx)"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2); section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = 'Consolas'; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5'
    })
    shading.append(shd)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True; run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]; cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix); run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_para(text):
    doc.add_paragraph(text)


# ══════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SecureReq')
run.font.size = Pt(34); run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI-Powered Security Requirements from User Stories\nClient Demo Guide')
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('AppSec Platform\nDemo-ready walkthrough including full AI prompts')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. PITCH
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. What it is (30-second pitch)', level=1)
add_para(
    "SecureReq takes an ordinary user story (from Jira, Azure DevOps, ServiceNow, or manually "
    "entered) and automatically produces:"
)
add_bullet('5-7 realistic abuse cases — business-level, not technical ("promo code stacking", not "SQLi")', '• ')
add_bullet('6+ STRIDE-categorised threat scenarios', '• ')
add_bullet('10-15 actionable security requirements with testable acceptance criteria', '• ')
add_bullet('A risk score (0-100) for backlog prioritisation', '• ')
add_bullet('Compliance mappings to OWASP ASVS v1-14 and PCI-DSS requirements', '• ')
add_bullet('Optional insider-threat mode — generates additional insider-specific abuse cases and controls', '• ')
add_para(
    "The analysis can be published back to the originating ticket (Jira / ADO / SNOW) in the "
    "correct custom fields so developers see the security work items alongside their story."
)

# ══════════════════════════════════════════════════════════════════════
# 2. PROBLEM
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. The problem we solve', level=1)
add_para(
    "Product and security teams have different languages. SecureReq translates:"
)
add_table(
    ['Before', 'After'],
    [
        ['"As a user I can upload a profile picture"', '7 abuse cases (malicious file upload, phishing avatars, content-based attacks, EXIF leaks, ...) + 12 security requirements with acceptance criteria'],
        ['Security requirements added late in a sprint as "hardening"', 'Security requirements attached to the user story BEFORE development starts'],
        ['Threat modeling workshops per feature', 'Minutes of LLM analysis, with in-context learning from your team\'s positive / negative feedback'],
        ['Generic compliance boilerplate', 'Requirements mapped per-item to OWASP ASVS V1-V14 and PCI-DSS Req 1-12 with relevance scores'],
        ['No audit trail', 'Every analysis versioned, stored, and round-tripped into the originating ticket'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. PIPELINE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. End-to-end pipeline', level=1)
add_code_block(
    '[Stage 1]  Story Ingestion          (manual entry or Jira/ADO/SNOW sync)\n'
    '    |\n'
    '    v\n'
    '[Stage 2]  AI Provider Init         (Anthropic Claude preferred, OpenAI fallback)\n'
    '    |\n'
    '    v\n'
    '[Stage 3]  Feedback Fetch           (pull positive/negative examples for in-context learning)\n'
    '    |\n'
    '    v\n'
    '[Stage 4]  Prompt Construction      (main prompt + insider-threat extensions if enabled)\n'
    '    |\n'
    '    v\n'
    '[Stage 5]  LLM Invocation           (JSON output, 8192 max_tokens, retry on network error)\n'
    '    |\n'
    '    v\n'
    '[Stage 6]  JSON Parse + Normalise   (field aliasing, defaults, validation)\n'
    '    |\n'
    '    v\n'
    '[Stage 7]  Risk Scoring             (0-100 from counts + risk factors)\n'
    '    |\n'
    '    v\n'
    '[Stage 8]  Persistence              (security_analyses table, versioned per story)\n'
    '    |\n'
    '    v\n'
    '[Stage 9]  Compliance Mapping       (OWASP ASVS + PCI-DSS, keyword + category scoring)\n'
    '    |\n'
    '    v\n'
    '[Stage 10] Optional Publish         (back to Jira/ADO/SNOW with auto-resolved custom fields)'
)
add_para(
    "Only Stage 5 is LLM-based. Everything else is deterministic code. That means 90% of the "
    "pipeline is reproducible even with the LLM down."
)

# ══════════════════════════════════════════════════════════════════════
# 4. INPUT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Input — what the user provides', level=1)
add_table(
    ['Field', 'Required', 'Notes'],
    [
        ['title', 'yes', 'Story title'],
        ['description', 'yes', 'Detailed description'],
        ['acceptance_criteria', 'no', 'If omitted, the LLM is instructed to work from description alone'],
        ['insider_threat', 'no (bool)', 'Enables insider-threat mode (adds a second wave of abuse cases + requirements)'],
        ['source', 'auto', '"manual", "jira", "ado", "github", "snow"'],
        ['external_id / external_url', 'auto', 'e.g., PROJ-123 — set when synced from Jira/ADO'],
    ]
)
add_para(
    "AI provider and API key come from the user's profile — each customer picks their own "
    "provider (Anthropic Claude default, OpenAI fallback, Azure OpenAI supported)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. THE PROMPTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. The exact prompts (verbatim)', level=1)
add_para(
    "All of Stage 5 happens in one LLM call. The prompt is built dynamically from: the user "
    "story + recent positive/negative feedback examples + the insider-threat toggle."
)
add_para("Model parameters: max_tokens = 8192. Response is JSON. Retries up to 3× on network errors.")

# 5.1
doc.add_heading('5.1 System prompt (OpenAI only; Anthropic uses user-message only)', level=2)
add_code_block(
    'You are a security architect. Generate concise, actionable security analysis. Return only valid JSON.'
)

# 5.2
doc.add_heading('5.2 Main user prompt (non-insider-threat mode)', level=2)
add_code_block(
    'You are an expert application security analyst. Analyze the following user story for\n'
    'security threats, abuse cases, and generate security requirements.\n\n'
    '**User Story Title:** {title}\n'
    '**Description:** {description}{acceptance_criteria_section}\n\n'
    'Generate 5-7 realistic abuse cases. Each abuse case must have:\n'
    '- id: Unique identifier (AC-001, AC-002, etc.)\n'
    '- threat: Clear title of the abuse scenario\n'
    '- actor: Who would do this (Malicious User, Disgruntled Employee, Competitor, Fraudster, Bot)\n'
    '- description: Realistic scenario describing how this abuse would occur through normal application use\n'
    '- impact: Critical/High/Medium/Low\n'
    '- likelihood: High/Medium/Low\n'
    '- attack_vector: How the abuse is carried out\n'
    '- stride_category: Spoofing/Tampering/Repudiation/Information Disclosure/Denial of Service/Elevation of Privilege\n\n'
    'Focus on REALISTIC business abuse scenarios, not technical hacking attacks. Examples:\n'
    '- Account sharing to avoid subscription fees\n'
    '- Promotional code abuse and stacking\n'
    '- Refund/chargeback fraud\n'
    '- Data scraping by competitors\n'
    '- Insider data theft before resignation\n'
    '- Fake reviews or ratings manipulation\n\n'
    'Generate 10-15 actionable security requirements. Each requirement must have:\n'
    '- id: Unique identifier (SR-001, SR-002, etc.)\n'
    '- requirement: Clear, actionable security control statement\n'
    '- priority: Critical/High/Medium\n'
    '- category: Authentication/Authorization/Input Validation/Cryptography/Logging/Rate Limiting/\n'
    '            API Security/Data Protection/Session Management/Error Handling\n'
    '- rationale: Why this requirement is needed and implementation guidance\n'
    '- acceptance_criteria: Bullet-pointed testable criteria (use \\n for line breaks, start each with •)\n\n'
    'Requirements should be SPECIFIC to the user story functionality, not generic security controls.\n'
    'Map requirements to relevant compliance standards (OWASP, CWE, PCI-DSS, GDPR) where applicable.\n\n'
    '{feedback_section}\n\n'
    'Return ONLY valid JSON with this exact structure:\n'
    '{\n'
    '  "abuse_cases": [\n'
    '    {"id": "AC-001", "threat": "...", "actor": "...", "description": "...",\n'
    '     "impact": "High", "likelihood": "Medium", "attack_vector": "...",\n'
    '     "stride_category": "Information Disclosure"}\n'
    '  ],\n'
    '  "stride_threats": [\n'
    '    {"category": "Spoofing", "threat": "...", "description": "...", "risk_level": "High"}\n'
    '  ],\n'
    '  "security_requirements": [\n'
    '    {"id": "SR-001", "requirement": "...", "priority": "Critical", "category": "Authentication",\n'
    '     "rationale": "...", "acceptance_criteria": "• criterion 1\\n• criterion 2\\n• criterion 3"}\n'
    '  ],\n'
    '  "risk_score": 65\n'
    '}\n\n'
    'Generate at least 5 abuse cases, 6 STRIDE threats, and 10 security requirements.\n'
    'Be SPECIFIC to this user story, not generic.'
)

# 5.3
doc.add_heading('5.3 Insider-threat extension — appended when insider_threat=true', level=2)

doc.add_heading('Context note (added to preamble)', level=3)
add_code_block(
    'INSIDER THREAT MODE ENABLED: Generate BOTH regular abuse cases AND insider threat abuse cases.\n'
    'For each item, set "insider_threat": true for insider-threat-specific items,\n'
    'and "insider_threat": false for regular items.'
)

doc.add_heading('Additional abuse-case instructions (appended)', level=3)
add_code_block(
    'Additionally, generate insider threat abuse cases:\n\n'
    'Generate 5-7 realistic INSIDER THREAT abuse cases focused on privileged users, employees,\n'
    'contractors, and trusted insiders. Each abuse case must have:\n'
    '- id: Unique identifier (AC-001, AC-002, etc.)\n'
    '- threat: Clear title of the insider abuse scenario\n'
    '- actor: Disgruntled Employee / Privileged Admin / Contractor / Departing Employee /\n'
    '         Malicious Insider / Negligent Insider\n'
    '- description: Realistic scenario of insider abuse using legitimate access and credentials\n'
    '- impact: Critical/High/Medium/Low\n'
    '- likelihood: High/Medium/Low\n'
    '- attack_vector: How the insider carries out the abuse using their privileged access\n'
    '- stride_category: Spoofing/Tampering/Repudiation/Information Disclosure/\n'
    '                   Denial of Service/Elevation of Privilege\n\n'
    'Focus on INSIDER-SPECIFIC scenarios. Examples:\n'
    '- Privileged admin accessing customer financial data outside their role\n'
    '- Departing employee exfiltrating sensitive data before resignation\n'
    '- Developer embedding backdoor logic in production code\n'
    '- IT staff tampering with audit logs to cover unauthorized access\n'
    '- Contractor abusing temporary elevated permissions beyond scope\n'
    '- Insider selling customer PII to competitors\n'
    '- Employee bypassing approval workflows for unauthorized transactions'
)

doc.add_heading('Additional security-requirement instructions (appended)', level=3)
add_code_block(
    'Additionally, generate insider threat security requirements:\n\n'
    'Generate 10-15 actionable security requirements specifically addressing insider threat risks.\n'
    'Each requirement must have:\n'
    '- id: Unique identifier (SR-001, SR-002, etc.)\n'
    '- requirement: Clear, actionable security control to prevent or detect insider threats\n'
    '- priority: Critical/High/Medium\n'
    '- category: Access Control / Audit Logging / Data Loss Prevention /\n'
    '            Privileged Access Management / Separation of Duties / Behavioral Analytics /\n'
    '            Data Classification / Session Management / Monitoring / Cryptography\n'
    '- rationale: Why this control is critical for insider threat prevention\n'
    '             and how it detects or limits insider abuse\n'
    '- acceptance_criteria: Bullet-pointed testable criteria (use \\n, start each with •)\n\n'
    'Requirements must address: least-privilege access, comprehensive audit trails,\n'
    'anomaly detection, data exfiltration prevention, privileged access governance,\n'
    'and separation of duties.\n\n'
    'Map requirements to relevant compliance standards (NIST SP 800-53, ISO 27001, SOC2, PCI-DSS)\n'
    'where applicable.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. IN-CONTEXT LEARNING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. In-context learning — the {feedback_section} block', level=1)
add_para(
    "SecureReq has a thumbs-up / thumbs-down feedback UI. Every time an analyst rates a "
    "generated abuse case or security requirement, we persist it in the prompt_feedback table."
)
add_para(
    "On each subsequent analysis, the latest 5 positive + 3 negative examples per type are "
    "injected into the prompt as few-shot examples:"
)
add_code_block(
    '{feedback_section} block example:\n\n'
    'Here are examples of abuse cases rated by our team:\n\n'
    'POSITIVE EXAMPLES (generate similar style):\n'
    '  - "Promo Code Stacking": Bargain Hunter / High / ...\n'
    '  - "Refund Fraud via Chargeback": Fraudster / Critical / ...\n'
    '  ...\n\n'
    'NEGATIVE EXAMPLES (avoid these):\n'
    '  - "Generic SQL injection attack": ... (too generic, not tied to story)\n'
    '  - "DDOS the server": ... (not a business abuse case)\n'
    '  ...'
)
add_para(
    "Over time the model output converges on the style your security team finds useful. This is "
    "how the platform 'learns' without any fine-tuning."
)

# ══════════════════════════════════════════════════════════════════════
# 7. OUTPUT SCHEMA
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Output schema', level=1)
add_para("The LLM returns one JSON object; it is stored as a versioned SecurityAnalysis record:")
add_code_block(
    '{\n'
    '  "id": 42,\n'
    '  "user_story_id": 5,\n'
    '  "version": 1,\n'
    '  "abuse_cases": [\n'
    '    {\n'
    '      "id": "AC-001",\n'
    '      "threat": "Promo Code Stacking and Abuse",\n'
    '      "actor": "Bargain Hunter / Deal Community",\n'
    '      "description": "Customer discovers multiple promo codes can be combined at checkout...",\n'
    '      "impact": "High", "likelihood": "High",\n'
    '      "attack_vector": "Share working codes on deal forums",\n'
    '      "stride_category": "Tampering",\n'
    '      "mitigations": ["Enforce single promo code per order rule"],\n'
    '      "insider_threat": false\n'
    '    }\n'
    '  ],\n'
    '  "stride_threats": [\n'
    '    {"category": "Tampering",\n'
    '     "threat": "Transaction amount manipulation via request tampering",\n'
    '     "description": "Attacker modifies cart total before submission",\n'
    '     "risk_level": "High"}\n'
    '  ],\n'
    '  "security_requirements": [\n'
    '    {\n'
    '      "id": "SR-001",\n'
    '      "requirement": "Implement adaptive rate limiting on all authentication endpoints",\n'
    '      "priority": "Critical",\n'
    '      "category": "Authentication",\n'
    '      "rationale": "Prevents brute force attacks... [CWE-307]",\n'
    '      "acceptance_criteria": "• 5 failed attempts per minute per IP/username\\n• Exponential backoff...",\n'
    '      "insider_threat": false\n'
    '    }\n'
    '  ],\n'
    '  "risk_score": 65,\n'
    '  "risk_factors": [\n'
    '    {"factor": "Authentication", "score": 25, "description": "Feature involves user authentication"},\n'
    '    {"factor": "Financial Data", "score": 30, "description": "Feature handles payment card data"}\n'
    '  ],\n'
    '  "ai_model_used": "claude-sonnet-4-20250514",\n'
    '  "analysis_duration_ms": 3450,\n'
    '  "created_at": "2026-04-23T14:22:30Z"\n'
    '}'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 8. COMPLIANCE MAPPING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Compliance Mapping (deterministic)', level=1)
add_para(
    "After the LLM returns, each security requirement is mapped to every compliance control it "
    "could satisfy. No LLM involved — pure keyword + category scoring."
)

doc.add_heading('8.1 Frameworks covered', level=2)
add_table(
    ['Framework', 'Controls'],
    [
        ['OWASP ASVS', 'V1 Architecture | V2 Authentication | V3 Session Mgmt | V4 Access Control | V5 Validation/Sanitization/Encoding | V6 Cryptography | V7 Error Handling/Logging | V8 Data Protection | V9 Communication | V10 Malicious Code | V11 Business Logic | V12 Files & Resources | V13 API/Web Service | V14 Configuration'],
        ['PCI-DSS', 'Req 1-12 (network security, configuration, data protection, cryptography, malware, secure dev, access control, authentication, physical, logging, testing, policies)'],
        ['CWE', 'Embedded in rationale fields by the LLM (e.g., CWE-307, CWE-89, CWE-79)'],
        ['STRIDE', 'Built-in 6 categories — every abuse case and threat is STRIDE-tagged'],
    ]
)

doc.add_heading('8.2 Relevance scoring', level=2)
add_code_block(
    'relevance_score = 0.0\n'
    'if requirement.category matches control category:   relevance_score += 0.4\n'
    'for each keyword in requirement.text:\n'
    '    if keyword appears in control.title:            relevance_score += 0.1\n'
    'final = min(1.0, relevance_score)\n\n'
    '# Mappings with relevance_score < 0.3 are dropped.'
)
add_para(
    "The output is a compliance_mappings row per (requirement × control) pair above threshold, "
    "with the rationale string explaining the match."
)

# ══════════════════════════════════════════════════════════════════════
# 9. JIRA / ADO / SNOW
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Ticket system integration', level=1)

doc.add_heading('9.1 Inbound — Sync stories FROM Jira', level=2)
add_para("POST /api/securereq/projects/{id}/sync/jira")
add_bullet('JQL: project = {id} AND issuetype IN (...) ORDER BY created DESC', '• ')
add_bullet('For each issue: extract summary + description (from ADF), create or update UserStory with source=JIRA, external_id, external_url.', '• ')
add_bullet('Same pattern exists for Azure DevOps and ServiceNow.', '• ')

doc.add_heading('9.2 Outbound — Publish analysis BACK to the ticket', level=2)
add_para("POST /api/securereq/stories/{story_id}/publish")
add_para(
    "The analysis is posted back into the originating issue in two custom fields (Abuse Cases, "
    "Security Requirements). Content is formatted as Atlassian Document Format (ADF) with "
    "headings, bullets, and emphasis; plain-text fallback is used if ADF fails."
)

doc.add_heading('9.3 Recent robustness improvements', level=2)
add_para(
    "Jira custom-field IDs vary between team-managed and company-managed projects. Earlier "
    "versions used static IDs and broke when projects differed. The current implementation:"
)
add_bullet('Calls Jira editmeta per issue to discover which custom fields are editable', '• ')
add_bullet('Falls back to fuzzy name matching (case-insensitive) when the configured ID does not match', '• ')
add_bullet('Uses overrideScreenSecurity to bypass screen-level restrictions for service accounts', '• ')
add_bullet('Retries with plain-text content when ADF is rejected', '• ')
add_para(
    "Result: the publish flow works across every project configuration we have tested without "
    "per-project tuning."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 10. PERSISTENCE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Persistence', level=1)
add_table(
    ['Table', 'Purpose'],
    [
        ['user_stories', 'One row per story. Tracks source (manual/Jira/ADO/SNOW), external_id, external_url, is_analyzed, risk_score, threat_count, requirement_count.'],
        ['security_analyses', "Versioned analyses per story. Stores abuse_cases, stride_threats, security_requirements (all JSON), risk_score, risk_factors, ai_model_used, analysis_duration_ms."],
        ['compliance_mappings', 'One row per (requirement × control) pair above relevance threshold. OWASP ASVS + PCI-DSS.'],
        ['prompt_feedback', 'Thumbs up/down on individual abuse cases or requirements. Feeds in-context learning.'],
        ['integration_settings', 'Per-user Jira/ADO/SNOW credentials, base URL, custom field IDs, connection status.'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 11. API
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('11. API endpoints', level=1)
add_table(
    ['Method', 'Path', 'Purpose'],
    [
        ['GET / POST', '/api/securereq/projects/{id}/stories', 'List / create stories'],
        ['GET / PUT / DELETE', '/api/securereq/stories/{id}', 'Story CRUD'],
        ['POST', '/api/securereq/stories/{id}/analyze', 'Run analysis (accepts insider_threat flag)'],
        ['GET', '/api/securereq/stories/{id}/analyses', 'List all analyses for story'],
        ['GET', '/api/securereq/analyses/{id}', 'Get specific analysis version'],
        ['GET', '/api/securereq/analyses/{id}/compliance', 'Compliance mappings for analysis'],
        ['GET', '/api/securereq/analyses/{id}/compliance/summary', 'Summary by standard'],
        ['GET', '/api/securereq/projects/{id}/summary', 'Project-level stats (threats, risk, coverage)'],
        ['POST', '/api/securereq/projects/{id}/analyze-all', 'Batch analyze all unanalyzed stories'],
        ['POST', '/api/securereq/projects/{id}/sync/jira', 'Sync stories from Jira'],
        ['POST', '/api/securereq/projects/{id}/sync/ado', 'Sync stories from Azure DevOps'],
        ['POST', '/api/securereq/projects/{id}/sync/snow', 'Sync stories from ServiceNow'],
        ['POST', '/api/securereq/stories/{id}/publish', 'Publish analysis back to the ticket'],
        ['POST / GET / DELETE', '/api/securereq/feedback', 'Thumbs up/down feedback for in-context learning'],
        ['POST', '/api/securereq/from-threat', 'Convert a threat-model threat into a synthetic requirement story'],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 12. DEMO SCRIPT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('12. Suggested demo script (10-12 minutes)', level=1)
add_table(
    ['#', 'Step', 'What to say', 'What to click'],
    [
        ['1', 'Hook (30 s)', 'Security requirements are usually written too late. SecureReq writes them BEFORE development, from the user story itself.', 'Open SecureReq tab'],
        ['2', 'Jira sync (1 min)', 'Pull stories straight from Jira. Source, ticket ID, and URL are all retained.', 'Sync from Jira; show the imported list'],
        ['3', 'Run analysis (1 min)', 'One click. Under 5 seconds later you have abuse cases, threats, requirements.', "Pick a feature story, click Analyze"],
        ['4', 'Abuse cases (2 min)', "Notice these are business abuse cases — promo code stacking, refund fraud, scraping — not 'SQL injection'. That is what Product teams actually need to think about.", 'Expand abuse cases; show STRIDE category + actor + impact'],
        ['5', 'Security requirements (2 min)', "Each one has a priority, category, rationale with CWE, and testable acceptance criteria. Drop these straight into your sprint backlog.", 'Expand a requirement; show acceptance criteria bullet list'],
        ['6', 'Insider-threat mode (2 min)', 'Toggle insider-threat mode and re-run. Now we get 10-14 more items focused on privileged-user risk.', 'Toggle insider-threat; re-analyse; show the new items tagged insider_threat=true'],
        ['7', 'Compliance mapping (1 min)', 'Every requirement auto-maps to OWASP ASVS and PCI-DSS. This is the compliance-audit evidence, already generated.', 'Compliance tab'],
        ['8', 'Feedback loop (1 min)', 'Thumbs up what is useful, thumbs down what is generic. The next analysis picks up your style automatically — no fine-tuning.', 'Thumbs up / down on a couple items'],
        ['9', 'Publish back (1 min)', 'Push the analysis into the Jira ticket as custom fields. Developers see security work items in the tool they already use.', 'Click Publish to Jira; open the Jira issue'],
        ['10', 'Close', 'Bidirectional Jira integration + STRIDE + OWASP ASVS + PCI-DSS + feedback learning, all from one LLM call.', '—'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 13. Q&A
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('13. Likely client questions & short answers', level=1)
add_table(
    ['Question', 'Answer'],
    [
        ['Which LLM do you use?', "Anthropic Claude Sonnet 4 by default, OpenAI GPT-4o fallback. Azure OpenAI also supported. Customer picks per-user in their profile."],
        ['What about data privacy — can we keep stories on-prem?', "Yes. Point the platform at an internal Azure OpenAI or self-hosted LLM endpoint; the story never leaves your network."],
        ['How good are the outputs vs a human security architect?', "On first-pass quality, comparable to a junior analyst. The thumbs up/down feedback loop quickly biases the output toward your team's style — a form of in-context fine-tuning without a training pipeline."],
        ['What frameworks does it cover?', "STRIDE (every output), OWASP ASVS v1-v14 (mapping), PCI-DSS Req 1-12 (mapping), CWE (embedded in rationales), NIST/ISO/SOC2 in insider-threat mode."],
        ['Does this replace our threat modeling?', "No — it feeds it. Abuse cases and requirements flow into the Threat Modeling feature as explicit threats (source=\"securereq\"), closing the loop between product backlog and architectural threat model."],
        ['Can we do bulk analysis?', "Yes. POST /projects/{id}/analyze-all runs analyses in a background task for every unanalyzed story in the project."],
        ['Is it deterministic?', "The LLM call is not, but parsing, risk scoring, compliance mapping, and persistence are. Stage 5 is the only non-deterministic step. We cache the result so re-opening an analysis shows exactly what was generated."],
        ['How does Jira publishing work in edge cases (team-managed projects)?', "We auto-detect the editable custom fields per issue via the Jira editmeta endpoint, with fuzzy-name fallback. Recent fix (April 2026) uses overrideScreenSecurity for service accounts and retries with plain-text if ADF formatting is rejected."],
        ['What does it cost per story?', "One LLM call, ~3-5 K tokens out, ~1-2 K tokens in. At Claude Sonnet pricing ~$0.01-$0.03 per analysis."],
        ['Can we ban certain categories?', "Yes — via the feedback mechanism (thumbs down templates eliminate them over time), or via custom prompt overrides per user."],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 14. FILE MAP
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('14. Technical appendix — file map', level=1)
add_table(
    ['File', 'Purpose'],
    [
        ['backend/services/security_requirements_analyzer.py', 'Main service — prompts, LLM invocation, parsing, compliance mapping'],
        ['backend/routers/securereq.py', 'All API endpoints — stories, analyses, sync, publish, feedback'],
        ['frontend/src/pages/SecurityRequirementsPage.tsx', 'List + bulk-analyse UI'],
        ['frontend/src/pages/StoryAnalysisPage.tsx', 'Per-story analysis view — abuse cases, requirements, compliance, publish button'],
    ]
)

add_para(
    "Database tables: user_stories, security_analyses, compliance_mappings, prompt_feedback, "
    "integration_settings."
)

out_path = os.path.join(os.path.dirname(__file__), 'SecureReq_Security_Requirements_Client_Demo_Guide.docx')
doc.save(out_path)
print(f"Document saved to: {out_path}")
