#!/usr/bin/env python3
"""
Generate a comprehensive API Security Tools Comparison Word Document for SKF.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page setup (A4 Landscape for tables) ──
for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Helper: shade a cell
def shade_cell(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, size=Pt(9), color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, bold) in enumerate(cells_data):
        if i < len(row.cells):
            set_cell_text(row.cells[i], text, bold=bold, size=Pt(8) if not header else Pt(9))
            if header:
                shade_cell(row.cells[i], "1F4E79")
                row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("API Security Platform\nComprehensive Evaluation Report")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph("")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Prepared for SKF — Confidential")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\nClassification: CONFIDENTIAL — For Internal Use Only")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary & Market Context",
    "2. Why API Security Matters — The Business Case",
    "3. Evaluation Methodology & Parameters",
    "4. Vendor Deep-Dive Profiles",
    "   4.1 Akamai API Security (formerly Noname Security)",
    "   4.2 Salt Security",
    "   4.3 Traceable AI (Harness)",
    "   4.4 Wallarm",
    "   4.5 Imperva API Security (Thales)",
    "   4.6 Cequence Security",
    "   4.7 Cloudflare API Shield",
    "5. Head-to-Head Comparison Matrix",
    "6. Pricing & Licensing Deep-Dive",
    "7. Integration & Deployment Capabilities",
    "8. Analyst Recognition & Industry Awards",
    "9. Risk Assessment & Vendor Viability",
    "10. Recommendation Framework for SKF",
    "Appendix A: OWASP API Security Top 10 (2023) Coverage Map",
    "Appendix B: Sources & References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary & Market Context", level=1)

doc.add_paragraph(
    "This report provides an exhaustive, parameter-by-parameter evaluation of seven leading API security platforms "
    "for SKF's consideration. The evaluation covers technical capabilities, deployment models, integration ecosystems, "
    "pricing structures, licensing models, analyst recognition, and vendor viability."
)

doc.add_heading("Market Snapshot", level=2)
market_stats = [
    ("Market Size (2025)", "$1.25 billion"),
    ("Projected Market Size (2030)", "$4.6 billion"),
    ("CAGR (2025–2030)", "29.66%"),
    ("Organizations with API security incidents (past year)", "99%"),
    ("API-related breaches as % of web-based attacks", ">90%"),
    ("Annual cost of API security incidents to enterprises", "Up to $186 billion globally"),
    ("Average cost of a single data breach", "$4.45 million"),
    ("BOLA attacks as % of all API attacks", "~40%"),
    ("Growth in API attacks (year-over-year)", "400% increase"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Metric", bold=True, size=Pt(10))
set_cell_text(hdr[1], "Value", bold=True, size=Pt(10))
shade_cell(hdr[0], "1F4E79")
shade_cell(hdr[1], "1F4E79")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for metric, value in market_stats:
    row = table.add_row()
    set_cell_text(row.cells[0], metric, size=Pt(9))
    set_cell_text(row.cells[1], value, bold=True, size=Pt(9))

doc.add_paragraph("")
doc.add_paragraph(
    "Key Gartner Prediction: By 2026, 60% of organizations will add more granular firewalls and enforce "
    "least privilege across millions of machine identities. Gartner identified APIs as the #1 attack vector, "
    "a prediction that has been validated by real-world breach data."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. WHY API SECURITY MATTERS
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Why API Security Matters — The Business Case for SKF", level=1)

business_case_points = [
    ("Regulatory Compliance", "APIs handling PII, financial data, or health records must comply with GDPR, PCI-DSS, HIPAA, and SOC 2. API security platforms provide automated compliance mapping, audit trails, and policy enforcement that manual processes cannot match."),
    ("Shadow API Risk", "Research shows that organizations typically have 3x more APIs than they are aware of. Shadow APIs — undocumented endpoints deployed by development teams — represent the single largest blind spot in most security programs."),
    ("Supply Chain & Third-Party APIs", "Modern applications consume dozens of third-party APIs. A breach in any one of these can cascade into your environment. API security platforms monitor third-party API behavior for anomalies and data exfiltration."),
    ("AI/GenAI Attack Surface", "The proliferation of AI agents, LLM-powered features, and Model Context Protocol (MCP) servers is creating a new API attack surface that traditional WAFs cannot address. Purpose-built API security is now essential."),
    ("Cost of Inaction", "With 99% of organizations reporting at least one API security incident in the past year, and the average breach costing $4.45M, the ROI on API security tooling is compelling — most vendors report 3-5x ROI within the first year of deployment."),
]

for title_text, detail in business_case_points:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_text}: ")
    run.bold = True
    p.add_run(detail)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. EVALUATION METHODOLOGY
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Evaluation Methodology & Parameters", level=1)

doc.add_paragraph(
    "Each vendor was evaluated across 40+ parameters grouped into the following categories:"
)

eval_categories = [
    ("API Discovery & Inventory", "Automated discovery, protocol support (REST, GraphQL, gRPC, SOAP, WebSocket), shadow/zombie API detection, API cataloging, data classification"),
    ("Security Posture Management", "Risk scoring, conformance analysis, sensitive data detection, spec drift monitoring, compliance mapping (OWASP, PCI-DSS, HIPAA, GDPR)"),
    ("Runtime Threat Protection", "Real-time attack detection, behavioral analytics, ML/AI-powered anomaly detection, BOLA/IDOR detection, inline vs. out-of-band blocking"),
    ("Security Testing", "Pre-production testing, CI/CD integration, DAST capabilities, number of test cases, shift-left support"),
    ("AI/GenAI Security", "LLM/AI agent protection, MCP discovery, prompt injection defense, AI gateway capabilities"),
    ("Integration Ecosystem", "SIEM/SOAR, WAF, API gateways, CI/CD pipelines, ticketing (Jira, ServiceNow), communication (Slack, PagerDuty)"),
    ("Deployment Flexibility", "SaaS, on-premises, hybrid, multi-cloud, edge deployment, onboarding time"),
    ("Pricing & Licensing", "Pricing model, estimated annual costs, licensing structure, multi-year discounts, included support"),
    ("Vendor Viability", "Funding, acquisitions, analyst recognition, customer base, Fortune 500 penetration"),
]

for cat, desc in eval_categories:
    p = doc.add_paragraph()
    run = p.add_run(f"{cat}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. VENDOR DEEP-DIVE PROFILES
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. Vendor Deep-Dive Profiles", level=1)

# ── 4.1 AKAMAI ──
doc.add_heading("4.1 Akamai API Security (formerly Noname Security)", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Akamai Technologies (NASDAQ: AKAM) completed its acquisition of Noname Security for approximately $450 million "
    "in June 2024. Noname was founded in 2020 and achieved unicorn status within one year of stealth. "
    "The combined entity serves 20% of the Fortune 500. Akamai's global edge network (over 4,100 PoPs in 130+ countries) "
    "provides unmatched reach for API traffic analysis."
)

# Capabilities table
akamai_params = [
    ("API Discovery", "Automated discovery of ALL API types — internal, external, shadow, rogue, legacy. Protocol support: HTTP, REST, GraphQL, SOAP, XML-RPC, JSON-RPC, gRPC. Discovers APIs not managed by gateways. Continuous inventory with auto-classification."),
    ("Posture Management", "Continuous risk scoring with sensitive data classification. PII/PHI/PCI data detection in API payloads. Compliance mapping to PCI-DSS, HIPAA, GDPR, SOC 2. Spec drift monitoring against OpenAPI definitions. Authentication gap identification."),
    ("Runtime Protection", "Behavioral-based anomaly detection using ML models. Real-time traffic analysis correlating events across the attack kill chain. Automated full or partial remediation based on alert type and recurrence frequency. Does NOT natively block inline — operates out-of-band."),
    ("Active Testing", "150+ pre-built dynamic security tests simulating malicious traffic. Tests cover the full OWASP API Security Top 10. Can be scheduled at desired intervals. Runs on-demand or integrated into CI/CD pipelines. Highly configurable test suites."),
    ("External Scanning (Recon)", "Scans external API attack surface at regular intervals. Discovers public vulnerabilities and attack paths visible to external threat actors. Provides external attacker's view of the API surface."),
    ("AI/GenAI Security", "Detects Model Context Protocol (MCP) servers in code. Supports AI-powered infrastructure security assessment. Added in Q4 2025 enhancements."),
    ("Inline Blocking", "Does NOT block inline in the request path. Operates out-of-band for zero latency impact. Enforcement is achieved through integration with WAFs, API gateways, and CDNs."),
    ("SIEM/SOAR Integration", "Native SIEM API for exporting security events to Splunk, QRadar, ArcSight, etc. ServiceNow integration added Q4 2025. ITSM and SOAR workflow support for automated remediation."),
    ("CI/CD Integration", "Direct pipeline integration via built-in connectors for Jenkins, GitHub Actions, GitLab CI, Azure DevOps. 150+ tests run automatically in pipeline."),
    ("Deployment", "Fully platform-agnostic: SaaS, hybrid, on-premises. Works in complex environments with multiple CDNs, WAFs, and gateways. Connects with 40+ traffic sources including third-party WAFs, CDNs, and cloud environments."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Akamai API Security — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "1F4E79")
shade_cell(hdr[1], "1F4E79")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in akamai_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Annual subscription model priced by number of API endpoints + monthly request volume. "
    "Typical range for mid-sized deployments (10–50 endpoints, moderate traffic): $30,000–$70,000/year. "
    "Larger deployments with 100+ endpoints: $75,000–$150,000+/year. "
    "Fixed number of APIs per licensed bundle — additional fees if exceeded. "
    "Multi-year contracts (2–3 years) unlock 15–30% discount. "
    "Support services are priced separately and perceived as premium but high-value."
)

doc.add_page_break()

# ── 4.2 SALT SECURITY ──
doc.add_heading("4.2 Salt Security", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Salt Security is a pioneer in the API security space, founded in 2016 in Palo Alto, CA. "
    "The company has raised over $271 million in funding. Salt's patented AI/ML engine powers its proprietary "
    "API Context Engine (ACE), which provides the foundation for discovery, threat detection, and remediation. "
    "Recently expanded into AI agent security and MCP discovery."
)

salt_params = [
    ("API Discovery", "Continuous discovery of all APIs — internal, shadow, third-party — without relying on traffic replay or manual effort. Auto-generates and maintains API inventory. Identifies undocumented and deprecated endpoints. Supports REST, GraphQL, gRPC, SOAP."),
    ("Posture Management", "ACE engine provides pre-production design analysis. Conformance checks against API specifications. Risk assessment with sensitive data exposure identification. Spec violation detection. Authentication and authorization gap analysis."),
    ("Runtime Protection", "Patented AI/ML builds behavioral models per API over time. Detects anomalies in real-time by correlating events across the attack kill chain (reconnaissance → exploitation → data exfiltration). Identifies low-and-slow attacks that WAFs miss."),
    ("Shift-Left Testing", "Pre-production design analysis via ACE engine. Integrates security insights into development pipelines. API security posture feedback to developers during build phase."),
    ("AI Agent Security", "MCP discovery platform — discovers and secures APIs powering AI agents. Extends coverage to mobile apps, SaaS platforms, and microservices. Industry-leading in agentic AI API discovery."),
    ("Inline Blocking", "Operates OUT-OF-BAND with zero latency impact on API traffic. Achieves blocking through integrations with enforcement points — AWS WAF, API gateways (Kong, Apigee, MuleSoft), load balancers. Can automatically configure WAF rules and manage IP blocklists."),
    ("SIEM/SOAR Integration", "Integrates with major SOAR platforms for automated incident response. Pushes alerts and context to SIEM systems. Bidirectional WAF integration for enforcement."),
    ("CI/CD Integration", "Shift-left integrations for pre-production testing. API intelligence fed back into development workflows. Less mature than Akamai's 150+ automated tests."),
    ("Ticketing & Communication", "Native integrations with Jira, Slack, PagerDuty, ServiceNow. Automated ticket creation with full attack context and remediation guidance."),
    ("Deployment", "Flexible: inline or out-of-band deployment. Integrates with API gateways and load balancers without performance impact. SaaS and hybrid options. However, architecture drives high compute/storage costs and requires complex mirroring deployments."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Salt Security — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "2E75B6")
shade_cell(hdr[1], "2E75B6")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in salt_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Custom subscription-based pricing tailored per organization. Pricing depends on: API traffic volume, "
    "number of API deployments, ecosystem complexity, selected modules (discovery, runtime detection, attack prevention). "
    "Vendor benchmark data (Vendr) indicates a redline threshold of approximately $100,000/year. "
    "Large enterprise deployments with full module selection are estimated at $150,000–$300,000+/year. "
    "License includes: customer support + success teams, full onboarding support, deployment assistance, "
    "technical integration for alerts, enforcement, remediation tickets. "
    "No public pricing — requires direct vendor engagement."
)

doc.add_page_break()

# ── 4.3 TRACEABLE AI ──
doc.add_heading("4.3 Traceable AI (merged with Harness)", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Traceable AI merged with Harness in February 2025 (completed March 4, 2025), creating what they describe as "
    "the world's most advanced AI-native DevSecOps platform. Both companies originated from BIG Labs, "
    "a San Francisco-based startup studio founded by Jyoti Bansal. The combined entity has ~1,100 employees, "
    "~$250 million in expected 2025 annualized revenue, and a valuation of approximately $5 billion. "
    "This merger gives Traceable deep native integration with Harness CI/CD — a significant differentiator."
)

traceable_params = [
    ("API Discovery", "Real-time, risk-ranked catalog of all APIs. Discovers shadow and orphaned APIs using deep telemetry married to source-code analysis. Conformance analysis against OpenAPI specifications. Identifies zombie APIs (deployed but no longer maintained)."),
    ("Posture Management", "Comprehensive risk-ranked API catalog with continuous assessment. Identifies sensitive data flows, PII exposure, authentication gaps. Conformance analysis against OpenAPI specs. Maps data flows end-to-end."),
    ("Runtime Protection", "Observes user-level transactions. ML algorithms discover anomalous transactions. Supports inline blocking of malicious requests (differentiator vs. Salt/Akamai). Behavioral analytics with adaptive threat intelligence feeds."),
    ("Security Testing (DAST)", "Purpose-built API DAST — now branded as Traceable API DAST by Harness. Automated API security tests run within CI pipelines throughout the SDLC. Shift-left testing with native Harness platform integration. Tests against OWASP API Top 10."),
    ("Threat Intelligence", "Extended runtime protection with adaptive threat intelligence. Partial shift-left integrations making it more adaptive for API-first development teams."),
    ("Data Flow Mapping", "Deep telemetry combined with source-code analysis for comprehensive visibility into how data moves through APIs. Identifies data leakage paths and unauthorized data access patterns."),
    ("Inline Blocking", "SUPPORTS inline blocking of malicious API requests — a key differentiator versus Salt Security and Akamai, which operate out-of-band."),
    ("SIEM/SOAR Integration", "Integration with SIEM platforms included. Role-based access management (RBAC) built-in."),
    ("CI/CD Integration", "Native integration via Harness platform (post-merger advantage). Also supports Jenkins, GitHub Actions, GitLab CI. Deepest shift-left capability in the comparison."),
    ("Deployment", "SaaS, on-premises, multi-cloud — flexible deployment options. Docker/Kubernetes-native."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Traceable AI — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "548235")
shade_cell(hdr[1], "548235")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in traceable_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Modular licensing model — pay per module (discovery, runtime protection, testing). "
    "AWS Marketplace indicative pricing (requires vendor confirmation): "
    "Discovery tier ~$20,000/year for 250 API endpoints; Protection tier ~$70,000/year for 50M API calls/month. "
    "CAUTION: Per-module licensing + integration effort + infrastructure overhead can make total cost of ownership "
    "unpredictable. Enterprise customers should request a comprehensive TCO estimate. "
    "Free trial available (no credit card required). "
    "Target customers: Large Enterprises, Non-Profits, Public Administrations."
)

doc.add_page_break()

# ── 4.4 WALLARM ──
doc.add_heading("4.4 Wallarm", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Wallarm was named API Security Platform of the Year 2025 by CyberSecurity Breakthrough Awards (October 2025). "
    "Raised $55 million in Series C funding in July 2025, led by Toba Capital. "
    "Unveiled an industry-first API Revenue Protection capability and launched its next-gen Security Edge offering. "
    "Positioned as a cost-effective alternative to Salt and Traceable with native inline blocking."
)

wallarm_params = [
    ("API Discovery", "Complete inventory of APIs, AI applications, and agentic AI endpoints. Automated discovery across all environments. Covers traditional APIs and emerging AI/LLM endpoints."),
    ("Threat Detection", "Patented AI/ML API abuse detection. Automated threat detection, vulnerability scanning, real-time traffic analysis. Identifies injection attacks, API abuse, unauthorized access, credential stuffing, and business logic abuse."),
    ("Inline Blocking", "NATIVE INLINE BLOCKING in the request path — blocks attacks on day zero without requiring separate WAF integration. This is a KEY DIFFERENTIATOR vs. Salt Security and Akamai, which operate out-of-band. Wallarm functions as both detection AND enforcement in a single solution."),
    ("Session Analysis", "Advanced session-based analysis — ties API sessions to specific users and roles. Real-time session blocking capability. Context-aware blocking based on user behavior patterns."),
    ("Revenue Protection", "Industry-first API Revenue Protection capability — protects business logic and revenue-generating APIs from abuse, scraping, and unauthorized automation. Unique in the market."),
    ("Security Edge", "Next-gen Security Edge offering for edge-based API protection. Extends protection to CDN edge locations."),
    ("API SOC-as-a-Service", "Managed security operations included in higher-tier plans. Expert monitoring, threat hunting, and incident response for API attacks."),
    ("SIEM/SOAR Integration", "Integrates with SIEM/SOAR platforms. Alert forwarding with full context."),
    ("CI/CD Integration", "Integrates with development pipelines. Supports shift-left workflows."),
    ("Deployment", "Cloud-native, on-premises, hybrid. Functions as an NG WAF — does NOT require a separate WAF product."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Wallarm — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "BF8F00")
shade_cell(hdr[1], "BF8F00")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in wallarm_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Subscription-based with tiered plans based on usage metrics. "
    "Estimated range: $40,000–$150,000/year depending on tier and traffic volume. "
    "Positioned as cost-effective compared to Salt Security and Traceable. "
    "Higher tiers include API SOC-as-a-Service (managed security operations). "
    "Free trial available. No detailed public pricing — requires vendor engagement."
)

doc.add_page_break()

# ── 4.5 IMPERVA ──
doc.add_heading("4.5 Imperva API Security (Thales)", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Imperva was acquired by Thales Group for approximately $3.6 billion in December 2023. "
    "Named an Overall Leader in the 2025 KuppingerCole Leadership Compass for API Security & Management. "
    "Strong WAF heritage with API security capabilities layered on top — advantageous for organizations "
    "already in the Imperva/Thales ecosystem."
)

imperva_params = [
    ("API Discovery", "Continuous discovery of public, private, and shadow APIs. Auto-profiling of traffic to establish behavioral baselines. ML-driven analysis identifies new and changed endpoints."),
    ("BOLA Detection", "Real-time Broken Object Level Authorization (BOLA) detection and response. ML-driven analysis spots deviations from normal object access patterns and blocks instantly. Industry-leading BOLA coverage given that BOLA represents ~40% of all API attacks."),
    ("Runtime Protection", "Schema enforcement against OpenAPI definitions. Inline blocking of malicious requests. ML-powered attack analytics. Behavioral baseline monitoring."),
    ("Bot Mitigation", "Built-in bot management integrated with API security. Behavioral modeling and fingerprinting. Credential stuffing protection."),
    ("GenAI/LLM Security", "Safeguards against prompt injection, data leaks, and model abuse for GenAI-powered APIs. Protects LLM-backed API endpoints from adversarial inputs."),
    ("Data Leak Detection", "Detects sensitive data in API responses leaving the origin. Per-endpoint alerting. PII/PHI/PCI pattern detection."),
    ("API Gateway Integration", "Native integration with Kong, MuleSoft, Azure APIM, Apigee, F5. Deepest API gateway integration in this comparison."),
    ("SIEM/SOAR Integration", "Integrates with Logsign, Maverix, and other SIEM platforms. Alert forwarding with full context."),
    ("CI/CD Integration", "CI/CD integration supported for shift-left testing."),
    ("Deployment", "Three deployment options: (1) API Security Add-On for existing Imperva Cloud WAF users (lowest cost entry); (2) Cloud-Managed via Imperva Cloud WAF console; (3) Self-Managed via local management console. Supports cloud, on-premises, and hybrid."),
    ("Cloud Support", "AWS, Azure, GCP, Red Hat Cloud Suite, Network Critical."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Imperva API Security — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "7030A0")
shade_cell(hdr[1], "7030A0")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in imperva_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Bundled with Imperva Cloud WAF or available as standalone subscription. "
    "Estimated range: $50,000–$200,000/year depending on deployment model and traffic volume. "
    "Significantly lower cost if already an existing Imperva/Thales WAF customer (add-on pricing). "
    "Enterprise customers with existing Imperva contracts can negotiate bundled pricing. "
    "No public pricing — requires vendor engagement."
)

doc.add_page_break()

# ── 4.6 CEQUENCE ──
doc.add_heading("4.6 Cequence Security", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Cequence offers a Unified Application Protection (UAP) platform combining API security and bot management. "
    "Named a Leader in the 2025 KuppingerCole Leadership Compass for API Security. "
    "Ranked #128 on the 2025 Deloitte Technology Fast 500 (fastest-growing companies in North America). "
    "Key differentiator: less than 15 minutes to onboard an API — no SDK, instrumentation, or JavaScript required."
)

cequence_params = [
    ("API Discovery & Compliance", "Discovers, monitors, and tests APIs. Full visibility of internal, external, and third-party APIs. OWASP API Security Top 10 vulnerability identification. Risk classification and prioritization."),
    ("Bot Management", "Integrated bot management (unique in this comparison — most competitors require a separate product). Behavioral modeling + fingerprinting to stop credential abuse, account takeover, and automated fraud."),
    ("Threat Detection", "ML-based analytics engine mitigates online fraud, business logic abuse, exploits, automated bot activity, and OWASP API Top 10 attacks. Real-time detection and blocking."),
    ("AI Gateway", "Industry-first security layer to govern and protect agentic AI. Securely connects AI agents to enterprise and SaaS applications. Integrated OAuth 2.0 access control. Enterprise scalability and performance."),
    ("Inline Blocking", "Native inline blocking capability. Real-time enforcement without separate WAF."),
    ("Onboarding", "Less than 15 minutes to onboard an API. No SDK, instrumentation, or JavaScript integration required. Lowest friction deployment in this comparison."),
    ("Deployment", "SaaS, on-premises, hybrid installations. Azure Marketplace availability."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Cequence Security — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "C00000")
shade_cell(hdr[1], "C00000")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in cequence_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Annual subscription based on number of APIs protected, deployment scale, and feature set. "
    "Pricing factors: traffic volume, security requirements, additional modules. "
    "No public pricing available — requires vendor engagement for custom quote."
)

doc.add_page_break()

# ── 4.7 CLOUDFLARE ──
doc.add_heading("4.7 Cloudflare API Shield", level=2)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Company Background: ")
run.bold = True
p.add_run(
    "Cloudflare (NYSE: NET) offers API Shield as part of its broader application security portfolio. "
    "Leverages Cloudflare's massive global edge network (310+ cities, 100+ countries). "
    "More accessible pricing than pure-play API security vendors but narrower in depth. "
    "Best suited for organizations already on the Cloudflare platform seeking integrated API protection."
)

cloudflare_params = [
    ("Authentication", "mTLS certificate validation, JWT validation, API key verification, OAuth 2.0 token validation. Blocks requests from illegitimate clients at the edge before they reach origin."),
    ("Schema Validation", "OpenAPI schema enforcement — accepts valid API requests and blocks malformed requests and HTTP anomalies. Complements Cloudflare WAF's negative security model for comprehensive coverage."),
    ("Rate Limiting", "Per-endpoint session-based rate limiting with auto-generated suggestions. Extends DDoS protection to GraphQL endpoints. Volume-based and sequential abuse prevention."),
    ("Data Leak Detection", "Detects sensitive data within API responses leaving the origin. Per-endpoint alerting for data exposure."),
    ("API Discovery", "Automated endpoint discovery via traffic analysis. ML-powered discovery of undocumented endpoints."),
    ("Bot Management", "Behavioral modeling + fingerprinting (available as a separate add-on). Not integrated into API Shield by default."),
    ("Inline Blocking", "Native inline blocking — part of Cloudflare's edge proxy architecture. All enforcement happens at the edge before traffic reaches origin servers."),
    ("Plan Availability", "Free/Pro/Business plans: Endpoint Management + Schema Validation ONLY. Enterprise plan: Full API Shield suite (mTLS, schema validation, discovery, rate limiting, data leak detection, sequence analysis)."),
    ("Deployment", "SaaS ONLY — runs on Cloudflare's edge network. Cannot be deployed on-premises or in private cloud. Requires DNS routing through Cloudflare."),
    ("Limitations", "Limited direct CI/CD integration for API testing. Does not integrate with third-party API gateways as deeply as purpose-built tools. No dedicated API DAST capability. Less comprehensive for complex multi-gateway enterprise environments."),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Parameter", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Cloudflare API Shield — Detailed Assessment", bold=True, size=Pt(9))
shade_cell(hdr[0], "F4811F")
shade_cell(hdr[1], "F4811F")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for param, detail in cloudflare_params:
    row = table.add_row()
    set_cell_text(row.cells[0], param, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], detail, size=Pt(8))
    row.cells[0].width = Cm(4)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Pricing & Licensing: ")
run.bold = True
p.add_run(
    "Enterprise plan add-on. Free/Pro/Business plans receive Endpoint Management + Schema Validation only. "
    "Full API Shield suite requires Enterprise plan. "
    "Estimated: $5,000–$50,000/year as add-on to Enterprise plan (varies by traffic and features). "
    "Significantly cheaper if already a Cloudflare Enterprise customer. "
    "Cloudflare offers a free tier with limited API security features — unique in this comparison."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. HEAD-TO-HEAD COMPARISON MATRIX
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Head-to-Head Comparison Matrix", level=1)

headers = ["Parameter", "Akamai", "Salt Security", "Traceable AI", "Wallarm", "Imperva", "Cequence", "Cloudflare"]

matrix_data = [
    ["API Discovery", "All types incl. legacy; 6+ protocols", "All types; no replay needed", "Risk-ranked + shadow/zombie via code analysis", "APIs + AI/agentic endpoints", "Public/private/shadow with behavioral baselines", "Internal/external/3rd party", "Traffic-based ML discovery"],
    ["Inline Blocking", "No — out-of-band; blocks via WAF/gateway integrations", "No — out-of-band; blocks via WAF/gateway integrations", "Yes — native inline blocking", "Yes — native inline (key differentiator)", "Yes — inline with schema enforcement", "Yes — native inline", "Yes — edge proxy blocking"],
    ["Active Testing / DAST", "150+ dynamic tests in CI/CD", "Pre-production via ACE engine", "Full DAST in CI pipeline (Harness-native)", "Vulnerability scanning", "Limited", "ML-based testing", "Schema validation only"],
    ["AI/GenAI Protection", "MCP server detection in code", "AI agent + MCP discovery platform", "Adaptive threat intelligence", "AI app inventory + discovery", "Prompt injection & LLM defense", "AI Gateway with OAuth 2.0", "Not specialized"],
    ["OWASP API Top 10", "Full coverage (150+ tests)", "Full coverage", "Full coverage", "Full coverage", "Full + industry-leading BOLA focus", "Full coverage", "Partial (schema-focused)"],
    ["Bot Management", "Separate product (Akamai Bot Manager)", "No native bot management", "No native bot management", "Not primary focus", "Integrated bot management", "Integrated (key differentiator)", "Separate add-on"],
    ["SIEM/SOAR", "Native SIEM API; ServiceNow (Q4 2025)", "SOAR integration; bidirectional WAF", "SIEM integration; built-in RBAC", "Included", "Multiple SIEM partners (Logsign, Maverix)", "Standard integration", "Cloudflare Logs / Logpush"],
    ["CI/CD Depth", "Deep — 150+ tests auto-run in pipeline", "Shift-left via ACE; less mature", "Deepest — native Harness integration", "Supported", "Supported", "Supported", "Limited"],
    ["Deployment Options", "SaaS / Hybrid / On-prem; 40+ traffic sources", "Inline or out-of-band; SaaS / hybrid", "SaaS / On-prem / Multi-cloud", "Cloud / On-prem / Hybrid", "Cloud / On-prem / Hybrid; 3 deployment models", "SaaS / On-prem / Hybrid", "SaaS only (edge)"],
    ["Onboarding Speed", "Moderate", "Moderate–High (complex mirroring)", "Moderate", "Fast", "Moderate", "<15 min per API (fastest)", "Fast (if on Cloudflare)"],
    ["Pricing Model", "Per endpoint + volume; annual", "Custom per org; annual", "Per module; annual", "Tiered subscription", "Bundled or standalone; annual", "Per API + traffic; annual", "Enterprise add-on"],
    ["Est. Annual Cost", "$30K–$150K+", "$100K–$300K+", "$20K–$70K+ per module", "$40K–$150K", "$50K–$200K", "Custom (contact vendor)", "$5K–$50K"],
    ["Multi-year Discount", "15–30% (2–3 yr)", "Available", "Available", "Available", "Available", "Available", "Standard"],
    ["Free Trial", "No", "No", "Yes (no CC required)", "Yes", "No", "No", "Free tier (limited features)"],
    ["Fortune 500 Customers", "20% of Fortune 500", "Major enterprises", "Large enterprises", "Growing enterprise base", "Broad enterprise base", "Deloitte Fast 500 (#128)", "Broad customer base"],
]

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
table.autofit = True

# Header row
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=Pt(7))
    shade_cell(table.rows[0].cells[i], "1F4E79")
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Data rows
for row_data in matrix_data:
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        set_cell_text(row.cells[i], cell_text, bold=(i == 0), size=Pt(7))
        if i == 0:
            shade_cell(row.cells[i], "D6E4F0")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. PRICING & LICENSING DEEP-DIVE
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Pricing & Licensing Deep-Dive", level=1)

doc.add_paragraph(
    "Note: API security vendor pricing is largely opaque. The estimates below are compiled from vendor marketplace "
    "listings (AWS/Azure), third-party benchmarking platforms (Vendr, TrustRadius), peer review sites (G2, Gartner Peer Insights), "
    "and industry analyst reports. All figures should be validated through direct vendor engagement and RFP/RFI processes."
)

pricing_headers = ["Vendor", "Pricing Model", "Entry-Level Estimate", "Mid-Market Estimate", "Enterprise Estimate", "Multi-Year Discount", "What's Included", "Hidden Cost Risks"]
pricing_data = [
    ["Akamai", "Per endpoint + request volume; annual subscription", "$25K–$35K/yr (10–25 endpoints)", "$35K–$70K/yr (25–50 endpoints)", "$75K–$150K+/yr (100+ endpoints)", "15–30% for 2–3 yr contracts", "Platform access, basic support", "Per-bundle API limits; overage fees; premium support charged separately"],
    ["Salt Security", "Custom per organization; annual subscription", "~$100K/yr (baseline)", "$150K–$200K/yr", "$200K–$300K+/yr", "Available on request", "Support + success teams, onboarding, deployment assistance, integration support", "High compute/storage costs; complex deployment; privacy exposure from traffic mirroring"],
    ["Traceable AI", "Per module (discovery, runtime, testing); annual", "~$20K/yr (discovery only, 250 endpoints)", "$50K–$90K/yr (discovery + protection)", "$100K–$200K+/yr (full suite)", "Available on request", "Module-specific features, support", "TCO unpredictable due to per-module licensing + integration effort + infrastructure overhead"],
    ["Wallarm", "Tiered subscription based on usage; annual", "$40K–$60K/yr", "$60K–$100K/yr", "$100K–$150K/yr", "Available on request", "Platform, basic support; higher tiers include API SOC-as-a-Service", "Fewer pre-built API tests than Akamai/Traceable"],
    ["Imperva", "Bundled with WAF or standalone; annual subscription", "$30K–$50K/yr (WAF add-on)", "$50K–$100K/yr", "$100K–$200K/yr", "Available on request", "Platform access, support", "Full value requires existing Imperva WAF investment; standalone is more expensive"],
    ["Cequence", "Per API + traffic volume; annual subscription", "Contact vendor", "Contact vendor", "Contact vendor", "Available on request", "Platform, support", "Pricing not publicly available"],
    ["Cloudflare", "Enterprise plan add-on; usage-based", "$5K–$15K/yr (add-on)", "$15K–$30K/yr", "$30K–$50K/yr", "Standard Cloudflare terms", "API Shield features on Enterprise plan; limited features free", "SaaS only; no on-prem; limited depth vs. pure-play tools"],
]

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
table.autofit = True

for i, h in enumerate(pricing_headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=Pt(7))
    shade_cell(table.rows[0].cells[i], "1F4E79")
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_data in pricing_data:
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        set_cell_text(row.cells[i], cell_text, bold=(i == 0), size=Pt(7))
        if i == 0:
            shade_cell(row.cells[i], "D6E4F0")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. INTEGRATION & DEPLOYMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Integration & Deployment Capabilities", level=1)

int_headers = ["Integration Point", "Akamai", "Salt Security", "Traceable AI", "Wallarm", "Imperva", "Cequence", "Cloudflare"]
int_data = [
    ["SIEM", "Native SIEM API (Splunk, QRadar, ArcSight)", "SIEM forwarding", "SIEM integration", "Included", "Logsign, Maverix, others", "Standard", "Logpush / Logs"],
    ["SOAR", "ServiceNow (Q4 2025), workflow automation", "Major SOAR platforms", "Supported", "Supported", "Supported", "Supported", "Limited"],
    ["WAF Integration", "Works with any WAF/CDN; 40+ traffic sources", "Bidirectional — pushes rules to WAFs, manages IP lists", "N/A (has own blocking)", "IS the WAF (NG WAF)", "Native with Imperva WAF; Kong, MuleSoft, F5", "Works alongside WAFs", "Native Cloudflare WAF"],
    ["API Gateways", "Kong, Apigee, MuleSoft, Azure APIM, AWS API GW", "Kong, Apigee, MuleSoft, AWS API GW", "Via Harness; standard gateways", "Standard gateways", "Kong, MuleSoft, Azure APIM, Apigee, F5", "Standard gateways", "Limited 3rd-party"],
    ["CI/CD Pipelines", "Jenkins, GitHub Actions, GitLab CI, Azure DevOps", "Shift-left via ACE", "Native Harness; Jenkins, GitHub, GitLab", "Supported", "Supported", "Supported", "Limited"],
    ["Ticketing", "Jira, ServiceNow", "Jira, Slack, PagerDuty, ServiceNow", "Via Harness platform", "Supported", "ServiceNow", "Supported", "Limited"],
    ["Cloud Platforms", "AWS, Azure, GCP", "AWS, Azure, GCP", "AWS, Azure, GCP", "AWS, Azure, GCP", "AWS, Azure, GCP, Red Hat", "AWS, Azure, GCP", "Cloudflare Edge (global)"],
    ["Deployment Model", "SaaS / Hybrid / On-prem", "SaaS / Hybrid", "SaaS / On-prem / Multi-cloud", "Cloud / On-prem / Hybrid", "3 models: Add-on, Cloud-Managed, Self-Managed", "SaaS / On-prem / Hybrid", "SaaS only"],
    ["On-Premises Support", "Yes", "Yes (complex)", "Yes", "Yes", "Yes (Self-Managed option)", "Yes", "No"],
]

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
table.autofit = True

for i, h in enumerate(int_headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=Pt(7))
    shade_cell(table.rows[0].cells[i], "1F4E79")
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_data in int_data:
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        set_cell_text(row.cells[i], cell_text, bold=(i == 0), size=Pt(7))
        if i == 0:
            shade_cell(row.cells[i], "D6E4F0")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. ANALYST RECOGNITION
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Analyst Recognition & Industry Awards (2025)", level=1)

analyst_data = [
    ["Akamai", "Overall Leader — 2025 KuppingerCole Leadership Compass (API Security & Management). Leader in Product, Innovation, and Market categories. Acquired Noname Security ($450M, June 2024)."],
    ["Salt Security", "Leader — 2025 KuppingerCole Leadership Compass (API Security & Management). Pioneer in API security category. $271M+ total funding raised."],
    ["Traceable AI", "Merged with Harness (March 2025) — combined valuation ~$5B. ~$250M expected 2025 annualized revenue. Strong position in API DAST."],
    ["Wallarm", "API Security Platform of the Year 2025 — CyberSecurity Breakthrough Awards. $55M Series C (July 2025, led by Toba Capital). Global InfoSec Awards recognition."],
    ["Imperva", "Overall Leader — 2025 KuppingerCole Leadership Compass (API Security & Management). Leader in WAAP category. Acquired by Thales for $3.6B (December 2023)."],
    ["Cequence", "Leader — 2025 KuppingerCole Leadership Compass (API Security). Leader in GigaOm API Security Radar. #128 on 2025 Deloitte Technology Fast 500. Global InfoSec Awards for API Security + Bot Management."],
    ["Cloudflare", "Leader in edge security and CDN. API Shield is part of broader platform — not evaluated as standalone API security in most analyst reports."],
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
set_cell_text(hdr[0], "Vendor", bold=True, size=Pt(9))
set_cell_text(hdr[1], "Analyst Recognition & Awards (2025)", bold=True, size=Pt(9))
shade_cell(hdr[0], "1F4E79")
shade_cell(hdr[1], "1F4E79")
hdr[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
hdr[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for vendor, recognition in analyst_data:
    row = table.add_row()
    set_cell_text(row.cells[0], vendor, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], recognition, size=Pt(8))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. RISK ASSESSMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading("9. Risk Assessment & Vendor Viability", level=1)

risk_data = [
    ["Akamai", "LOW", "Public company (NASDAQ: AKAM). $3.8B+ annual revenue. Noname integration adds API-native capability. Risk: Integration complexity during Noname absorption; potential product roadmap uncertainty."],
    ["Salt Security", "MEDIUM", "Private, well-funded ($271M+). Market pioneer. Risk: No profitability path disclosed; premium pricing may face pressure as competitors mature; complex deployment architecture."],
    ["Traceable AI", "LOW-MEDIUM", "Merged with Harness ($5B combined valuation). Strong revenue ($250M). Risk: Post-merger integration execution; potential feature consolidation; brand identity transition."],
    ["Wallarm", "MEDIUM", "Private, recently funded ($55M Series C). Growing but smaller than competitors. Risk: Smaller customer base; may face scaling challenges; limited analyst coverage compared to leaders."],
    ["Imperva", "LOW", "Owned by Thales Group (€18B+ revenue). Deep resources. Risk: Large-company bureaucracy may slow innovation; API security is one of many product lines."],
    ["Cequence", "MEDIUM", "Private. Fast-growing (Deloitte Fast 500). Risk: Less brand recognition than Akamai/Imperva; pricing transparency concerns."],
    ["Cloudflare", "LOW", "Public company (NYSE: NET). $1.6B+ annual revenue. Risk: API Shield is an add-on, not core focus; less depth than pure-play vendors; SaaS-only limits deployment flexibility."],
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(["Vendor", "Risk Level", "Assessment"]):
    set_cell_text(hdr[i], h, bold=True, size=Pt(9))
    shade_cell(hdr[i], "1F4E79")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

risk_colors = {"LOW": "C6EFCE", "LOW-MEDIUM": "FFEB9C", "MEDIUM": "FFEB9C", "HIGH": "FFC7CE"}
for vendor, risk, assessment in risk_data:
    row = table.add_row()
    set_cell_text(row.cells[0], vendor, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], risk, bold=True, size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(row.cells[1], risk_colors.get(risk, "FFFFFF"))
    set_cell_text(row.cells[2], assessment, size=Pt(8))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. RECOMMENDATION FRAMEWORK
# ══════════════════════════════════════════════════════════════
doc.add_heading("10. Recommendation Framework for SKF", level=1)

doc.add_heading("Selection by Priority", level=2)

reco_data = [
    ("Best for Inline Blocking + Cost Efficiency", "Wallarm or Cloudflare (if already on their edge)", "Wallarm provides native inline blocking with no separate WAF required. Cloudflare is the lowest-cost option if SKF is already on the Cloudflare platform."),
    ("Best for Deep API Testing in CI/CD", "Akamai (150+ tests) or Traceable AI (native Harness DAST)", "Akamai leads in pre-built test quantity. Traceable has the deepest CI/CD integration via the Harness merger."),
    ("Best for AI/GenAI API Security", "Salt Security (MCP discovery) or Cequence (AI Gateway)", "Salt is the market leader in agentic AI and MCP discovery. Cequence's AI Gateway with OAuth 2.0 is unique for AI agent governance."),
    ("Best for Existing WAF Customers", "Imperva (add-on) or Cloudflare (integrated)", "Both offer lower cost of entry for existing customers. Imperva's three deployment models provide the most flexibility."),
    ("Lowest Entry Cost", "Cloudflare API Shield ($5K–$15K/yr) or Traceable discovery tier (~$20K/yr)", "Cloudflare offers a free tier with basic features. Traceable's modular pricing allows starting with discovery only."),
    ("Most Comprehensive (Full Lifecycle)", "Salt Security or Akamai", "Both cover discovery, posture, runtime, and testing. Salt leads in behavioral modeling depth; Akamai leads in test automation."),
    ("Fastest Time to Value", "Cequence (<15 min onboarding) or Cloudflare", "Cequence's zero-instrumentation onboarding is unmatched. Cloudflare is instant for existing customers."),
    ("Best for Regulatory Compliance", "Imperva or Akamai", "Both have deep compliance mapping (PCI-DSS, HIPAA, GDPR, SOC 2) and audit trail capabilities."),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(["Selection Criteria", "Recommended Vendors", "Rationale"]):
    set_cell_text(hdr[i], h, bold=True, size=Pt(9))
    shade_cell(hdr[i], "1F4E79")
    hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for criteria, vendors, rationale in reco_data:
    row = table.add_row()
    set_cell_text(row.cells[0], criteria, bold=True, size=Pt(8))
    set_cell_text(row.cells[1], vendors, bold=True, size=Pt(8))
    set_cell_text(row.cells[2], rationale, size=Pt(8))

doc.add_paragraph("")

doc.add_heading("Suggested Next Steps for SKF", level=2)
next_steps = [
    "1. Define Requirements: Prioritize inline blocking vs. out-of-band, on-prem requirements, CI/CD depth, and AI/GenAI coverage based on SKF's API landscape.",
    "2. Shortlist 2–3 Vendors: Based on the selection framework above, narrow to 2–3 vendors for proof-of-concept evaluation.",
    "3. Request Custom Pricing: Engage shortlisted vendors with SKF's specific API count, traffic volume, and deployment requirements for accurate quotes.",
    "4. Conduct PoC (4–6 weeks): Deploy shortlisted tools against a representative subset of SKF's APIs to evaluate real-world performance, false positive rates, and integration effort.",
    "5. Evaluate TCO: Factor in not just license cost but integration effort, infrastructure requirements, training, and ongoing operational overhead.",
    "6. Negotiate Multi-Year: Once selected, negotiate a 2–3 year contract for 15–30% savings.",
]

for step in next_steps:
    doc.add_paragraph(step)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX A: OWASP COVERAGE MAP
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: OWASP API Security Top 10 (2023) Coverage Map", level=1)

owasp_headers = ["OWASP API Risk", "Akamai", "Salt", "Traceable", "Wallarm", "Imperva", "Cequence", "Cloudflare"]
owasp_data = [
    ["API1:2023 BOLA", "Detection via behavioral ML", "Detection via behavioral ML", "Detection via user-level transaction analysis", "Detection + inline blocking", "Industry-leading BOLA detection & auto-response", "Detection + blocking", "Partial (schema validation)"],
    ["API2:2023 Broken Authentication", "Posture assessment + runtime detection", "ACE engine analysis", "Authentication gap identification", "Detection + blocking", "mTLS, JWT validation", "Detection + blocking", "mTLS, JWT, OAuth 2.0 validation"],
    ["API3:2023 Broken Object Property Level Auth", "Sensitive data classification + detection", "Behavioral model detection", "Data flow mapping + detection", "Detection + blocking", "Schema enforcement", "Detection + blocking", "Schema validation"],
    ["API4:2023 Unrestricted Resource Consumption", "Detection via traffic analysis", "Behavioral anomaly detection", "Transaction-level monitoring", "Rate limiting + blocking", "Rate limiting", "ML-based detection", "Per-endpoint rate limiting"],
    ["API5:2023 Broken Function Level Auth", "Posture + runtime detection", "ACE conformance checks", "RBAC analysis", "Detection + blocking", "Posture management", "Detection + blocking", "Limited"],
    ["API6:2023 Unrestricted Access to Sensitive Flows", "Active testing (150+ tests)", "Behavioral model detection", "Business flow analysis", "Revenue Protection feature", "Bot mitigation", "Bot + fraud prevention", "Rate limiting"],
    ["API7:2023 Server-Side Request Forgery", "Active testing coverage", "Runtime detection", "DAST testing", "NG WAF detection", "WAF protection", "Detection", "WAF rules"],
    ["API8:2023 Security Misconfiguration", "Posture management + spec drift", "ACE conformance analysis", "Conformance analysis", "Vulnerability scanning", "Schema enforcement", "Compliance testing", "Schema validation"],
    ["API9:2023 Improper Inventory Mgmt", "Full discovery (6+ protocols)", "Continuous discovery", "Shadow/zombie API detection", "Full inventory + AI endpoints", "Continuous discovery", "Full discovery", "Traffic-based discovery"],
    ["API10:2023 Unsafe Consumption of APIs", "Third-party API monitoring", "Third-party API discovery", "Data flow + source analysis", "Third-party detection", "Third-party monitoring", "Third-party monitoring", "Limited"],
]

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'
table.autofit = True

for i, h in enumerate(owasp_headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=Pt(7))
    shade_cell(table.rows[0].cells[i], "1F4E79")
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for row_data in owasp_data:
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        set_cell_text(row.cells[i], cell_text, bold=(i == 0), size=Pt(7))
        if i == 0:
            shade_cell(row.cells[i], "D6E4F0")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# APPENDIX B: SOURCES
# ══════════════════════════════════════════════════════════════
doc.add_heading("Appendix B: Sources & References", level=1)

sources = [
    "Akamai — API Security Product Page (akamai.com/products/api-security)",
    "Akamai — Q4 2025 Enhancements Blog (akamai.com/blog/security/2026/jan/enhancements-akamai-api-security-q4-2025)",
    "Akamai — Noname Acquisition Press Release, June 2024 (akamai.com/newsroom)",
    "TechCrunch — Akamai confirms Noname acquisition for $450M, May 2024",
    "Salt Security — Platform Overview (salt.security/platform)",
    "Salt Security — Ecosystem Integrations (salt.security/blog)",
    "AWS Blog — Preventing API Breaches Using Salt Security with AWS WAF",
    "Vendr — Salt Security Pricing Benchmark (vendr.com/marketplace/salt-security)",
    "Traceable AI — Harness Merger Announcement, February 2025 (prnewswire.com)",
    "Traceable AI — Merger Completion, March 2025 (prnewswire.com)",
    "Akto — Traceable AI Product Analysis (akto.io/blog/traceable)",
    "Wallarm — Series C Announcement, $55M, July 2025 (wallarm.com/press-releases)",
    "CyberSecurity Breakthrough — Wallarm API Security Platform of Year 2025 (globenewswire.com)",
    "Wallarm — 2025 Year in Review (lab.wallarm.com)",
    "Imperva — API Security Product Page (imperva.com/products/api-security/)",
    "Thales — Imperva Acquisition Completion, December 2023 (thalesgroup.com)",
    "Imperva — KuppingerCole 2025 Leadership Compass Recognition",
    "Cequence Security — Platform Overview (cequence.ai/products/)",
    "Cequence — 2025 Deloitte Technology Fast 500 (#128) (globenewswire.com)",
    "KuppingerCole — Leadership Compass: API Security & Management 2025 (kuppingercole.com)",
    "Cloudflare — API Shield Documentation (developers.cloudflare.com/api-shield/)",
    "Cloudflare — API Shield Plans (developers.cloudflare.com/api-shield/plans/)",
    "OWASP — API Security Top 10 2023 (owasp.org/API-Security/)",
    "Gartner — Market Guide for API Protection (gartner.com)",
    "DevOps Radar — API Security Deep Dive: Salt, Traceable, Akamai (devops-radar.com)",
    "Mordor Intelligence — API Security Market Size Report 2025–2030",
    "Astra — API Security Trends 2026 (getastra.com)",
    "LevelBlue — Gartner API Attack Vector Prediction Validation (levelblue.com)",
]

for i, source in enumerate(sources, 1):
    p = doc.add_paragraph(f"{i}. {source}")
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(8)

# ── DISCLAIMER ──
doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Disclaimer: ")
run.bold = True
run.font.size = Pt(8)
run2 = p.add_run(
    "Pricing estimates are based on publicly available data from vendor marketplaces, third-party benchmarking platforms, "
    "and peer review sites as of March 2026. Actual pricing may vary based on negotiation, volume, contract terms, and "
    "specific deployment requirements. All figures should be validated through direct vendor engagement and formal RFP/RFI processes. "
    "This report is prepared for SKF internal evaluation purposes only."
)
run2.font.size = Pt(8)

# ── SAVE ──
output_path = "/Users/yashwanthgk/appsec-platform/API_Security_Tools_Comparison_SKF.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
